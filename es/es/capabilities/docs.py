"""Document reading: confine the path, dispatch on format, cache the result.

The only module the MCP layer imports for es_doc_*. Keeps doc_pdf free of
cache and confinement concerns, and doc_cache free of parsing concerns.

Dispatch is a table, CONVERTERS, keyed by lowercased extension; SUPPORTED is
derived from it so the two can't drift apart. CONVERTERS holds ".pdf" (via
doc_pdf, the reference/paginated converter), ".txt"/".md"/".csv"/".json"
(all four via doc_text, the flat/non-paginated converter), ".ics" (via
doc_ics, which synthesizes one "## " heading per VEVENT so a flat calendar
feed still pages well), and ".docx"/".xlsx" (both via doc_office, which reads
a Word document's own heading outline and gives each Excel sheet its own "##"
heading — neither has page_count/render, same as doc_text's formats) —
adding a new format is meant to be a pure addition (a new converter module +
one new table entry), not a change to extract()/render() themselves.
"""
import csv
import json
import os
import re
from pathlib import Path
from typing import List, Optional, Tuple
from xml.etree.ElementTree import ParseError
from zipfile import BadZipFile

from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError
from openpyxl.utils.exceptions import InvalidFileException
from pdfminer.pdfdocument import PDFEncryptionError
from pdfplumber.utils.exceptions import PdfminerException

from es import config, doc_cache, paths
from es.capabilities import doc_ics, doc_office, doc_pdf, doc_text

MAX_MARKDOWN_CHARS = 40_000
MAX_DOCUMENT_BYTES = 50 * 1024 * 1024
# Mirrors doc_pdf.MAX_AUTO_RENDER_PAGES rather than duplicating the number —
# render() and extract()'s auto-render both bound the same disk-fill risk
# (rasterizing into Hermes's shared upload cache), so one call site owning the
# limit and the other pointing at it keeps them from drifting apart.
MAX_RENDER_PAGES = doc_pdf.MAX_AUTO_RENDER_PAGES

# The stdlib csv module's default field_size_limit (128 KiB) is a parser
# safety valve, not a real document-shape limit — a single free-text CSV
# column (a pasted comment, a description field) can legitimately exceed it,
# and hitting the default raises `_csv.Error: field larger than field limit`,
# which used to surface to the agent as the literal, unexplainable es_code
# "Error". csv.field_size_limit() is a process-wide setting on the `_csv` C
# module, not a per-Reader option, so it can be raised HERE — once, for the
# life of the process — without doc_text.py (which owns the actual
# `csv.reader()` call) needing to change at all. 10 MiB comfortably covers any
# realistic single field while still being a real, bounded ceiling: a CSV
# with an unbalanced/unterminated quote makes csv.reader treat everything
# from that quote to EOF as ONE field (see doc_text._convert_csv's docstring
# on why it must use a real line iterator, not splitlines()), so an
# unterminated quote in a large file must still hit SOME limit rather than
# buffer the rest of a 50MB upload into one Python string.
CSV_FIELD_SIZE_LIMIT = 10 * 1024 * 1024
csv.field_size_limit(CSV_FIELD_SIZE_LIMIT)

# Maps a supported extension to its converter module. Each converter is
# expected to implement a shared per-format contract (doc_pdf.py is the
# reference): `convert(source, adir, pages=None) -> (markdown, images)` is
# REQUIRED; `page_count(source) -> int` and `render(source, adir, pages) ->
# images` are OPTIONAL — their presence is how the rest of this module tells
# a paginated/renderable format (PDF, via doc_pdf) apart from a flat one
# (txt/md/csv/json, via doc_text — none of the four have pages, so that
# module implements neither attribute) without hardcoding a format list
# anywhere else. SUPPORTED is DERIVED from this table rather than
# hand-maintained alongside it, so the two can never drift apart: an entry
# here is automatically "supported", and nothing can claim to be supported
# without a converter.
CONVERTERS = {
    ".pdf": doc_pdf,
    ".txt": doc_text,
    ".md": doc_text,
    ".csv": doc_text,
    ".json": doc_text,
    ".ics": doc_ics,
    ".docx": doc_office,
    ".xlsx": doc_office,
}
SUPPORTED = set(CONVERTERS)

# Cache filenames within a doc_id's artifact dir. Both are written ONLY for a
# full-document extract (pages=None) — see the comment in extract().
DOC_MD_NAME = "doc.md"
DOC_IMAGES_MANIFEST = "images.json"


class UnsupportedDocument(Exception):
    es_code = "doc_unsupported"


class DocumentTooLarge(Exception):
    es_code = "doc_too_large"


class InvalidPageRange(Exception):
    es_code = "doc_invalid_pages"


class EncryptedDocument(Exception):
    es_code = "doc_encrypted"


class UnreadableDocument(Exception):
    es_code = "doc_unreadable"


def parse_pages(spec: str, page_count: int) -> List[int]:
    """Parse "1-3,7" into a sorted, deduplicated, 1-indexed page list."""
    out = set()
    for part in str(spec).split(","):
        part = part.strip()
        if not part:
            continue
        try:
            if "-" in part:
                lo_s, hi_s = part.split("-", 1)
                lo, hi = int(lo_s), int(hi_s)
            else:
                lo = hi = int(part)
        except ValueError:
            raise InvalidPageRange(
                f"could not read page range {spec!r} — use forms like "
                '"3", "1-5", or "1-2,7"') from None
        if lo < 1 or hi < 1:
            raise InvalidPageRange(
                f"page range {part!r} must use page numbers starting at 1")
        if lo > hi:
            raise InvalidPageRange(
                f'page range {part!r} is reversed — did you mean "{hi}-{lo}"?')
        if page_count <= 0:
            raise InvalidPageRange(
                f"pages {part!r} are outside this document (it has no pages)")
        if hi > page_count:
            raise InvalidPageRange(
                f"pages {part!r} are outside this document (1-{page_count})")
        out.update(range(lo, hi + 1))
    if not out:
        raise InvalidPageRange(f"no pages selected by {spec!r}")
    return sorted(out)


def _reraise_pdf_error(real: Path, exc: PdfminerException):
    """pdfplumber wraps every parse failure in ONE exception class
    (PdfminerException), passing the real pdfminer exception as its sole arg
    — that is the only place "needs a password" is distinguishable from
    "not a readable PDF at all" (the encrypted case otherwise surfaces with
    an EMPTY message, telling the agent nothing). Verified empirically:
    pdfminer.pdfdocument.PDFPasswordIncorrect (raised for both a missing and
    a wrong password) is a subclass of PDFEncryptionError, so one isinstance
    check reliably separates the two — no guessing required."""
    cause = exc.args[0] if exc.args else None
    if isinstance(cause, PDFEncryptionError):
        raise EncryptedDocument(
            f"{real.name} is password-protected — es cannot open encrypted "
            "PDFs; ask the user for an unlocked copy") from exc
    raise UnreadableDocument(
        f"{real.name} could not be read as a PDF — it may be corrupt, "
        "truncated, or not actually a PDF; ask the user to resend it") from exc


# The OLE2/CFBF container signature (MS-CFB) — every legitimate, unencrypted
# .docx/.xlsx is a zip archive; a password-protected one is instead stored in
# this legacy container format (the same one .doc/.xls used), which is how
# real Office password-protection actually works, not a guess. Neither
# python-docx nor openpyxl exposes a distinct "needs a password" exception —
# both simply fail to open the file as a zip, indistinguishable by exception
# type alone from ordinary corruption (verified empirically against both
# libraries) — so this sniffs the file's own magic bytes instead, the same
# first-principles move _reraise_pdf_error makes by inspecting pdfminer's
# real cause rather than trusting a single generic exception type.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _is_ole2_container(real: Path) -> bool:
    try:
        with open(real, "rb") as f:
            return f.read(len(_OLE2_MAGIC)) == _OLE2_MAGIC
    except OSError:
        return False


def _reraise_office_error(real: Path, exc: Exception):
    """python-docx and openpyxl each raise their OWN exception type for a
    parse failure — PackageNotFoundError (python-docx, which itself
    normalizes a bad/missing zip into this one type), BadZipFile (openpyxl,
    which does NOT normalize it), and InvalidFileException (openpyxl's own
    filename-extension guard — included for defense in depth even though
    docs.py's dispatch already only ever routes a real .xlsx path here, so
    it should never actually fire in practice). All three collapse to the
    same UnreadableDocument, EXCEPT when the file is an OLE2 container (see
    _is_ole2_container above) — that one case is password-protection, not
    corruption, and gets its own EncryptedDocument message."""
    if _is_ole2_container(real):
        raise EncryptedDocument(
            f"{real.name} is password-protected — es cannot open encrypted "
            "Word/Excel documents; ask the user for an unlocked copy") from exc
    raise UnreadableDocument(
        f"{real.name} could not be read as a Word/Excel document — it may "
        "be corrupt, truncated, or not actually a .docx/.xlsx file; ask the "
        "user to resend it") from exc


def _reraise_text_error(real: Path, exc: Exception) -> None:
    """doc_text's own decoding (`errors="replace"`) only covers turning
    bytes into a str — it says nothing about csv.reader's or json.loads's
    own STRUCTURAL parsing of that str, and both can still raise past it:

    - csv.Error: either a field genuinely larger than CSV_FIELD_SIZE_LIMIT
      above (rare now that the limit is 10 MiB, not stdlib's 128 KiB
      default), or an unbalanced/unterminated quote — which makes
      csv.reader treat everything from that quote to EOF as one field,
      hitting the exact same limit from a different root cause. Either way
      the agent needs "narrower/cleaner CSV", not "Error".
    - RecursionError: json.loads recurses one Python stack frame per nesting
      level; a JSON document nested far deeper than any hand-authored file
      would be (~1000 levels) blows the interpreter's recursion limit. Not
      a bug in this module — a genuinely pathological/adversarial document.

    Both collapse to the same UnreadableDocument as every other converter's
    parse failures, distinguished only by exception type (never by
    guessing at message text) since which one fired already tells us which
    of .csv/.json triggered it."""
    if isinstance(exc, RecursionError):
        raise UnreadableDocument(
            f"{real.name} is nested too deeply to parse as JSON; ask the "
            "user to resend a flatter export") from exc
    raise UnreadableDocument(
        f"{real.name} could not be parsed as a CSV — a field is larger than "
        f"the {CSV_FIELD_SIZE_LIMIT // (1024*1024)}MB limit, or an "
        "unbalanced quote runs to the end of the file; ask the user to "
        "resend it or export a narrower/cleaner version") from exc


# Per-converter-module exception mapping: each module's OWN underlying
# library raises its own exception type on a parse failure, and this table
# names exactly those types so the two call sites in extract() below can
# catch precisely them — never a bare `except Exception`, which would just
# as happily swallow a genuine bug in OUR code (a TypeError from a mistake in
# this module, say) and misreport it to the agent as "your file is corrupt".
#
# doc_office's tuple is wider than "obviously corrupt" on purpose — verified
# empirically (not guessed) against real python-docx/openpyxl behavior for
# every case in the review that motivated this: a real .docx renamed .xlsx
# (openpyxl -> OSError), a real .xlsx renamed .docx (python-docx -> ValueError),
# a truncated worksheet XML (openpyxl's read_only streaming reader uses the
# stdlib xml.etree parser regardless of lxml being installed -> ParseError), a
# truncated word/document.xml (python-docx uses lxml -> XMLSyntaxError), a
# valid zip missing "[Content_Types].xml" entirely (both libraries do a plain
# dict-style archive lookup -> KeyError), and a valid zip whose
# "[Content_Types].xml" parses but has no default namespace so python-docx's
# lxml class lookup never upgrades it to its own CT_Types wrapper, and the
# next attribute access on it -> AttributeError. All nine are ordinary
# "wrong/damaged file" shapes (a renamed extension, a partial download), not
# adversarial input — a genuine bug in OUR OWN code around the `mod.convert()`
# call below would have to coincidentally raise one of these same types to be
# swallowed here, and nothing in this module's few lines around that call
# does.
#
# doc_ics has no entry: it already catches its own parse failures internally
# (see doc_ics._read_calendar's `except Exception: return None`, which
# convert() turns into a friendly in-band "could not be read" markdown
# instead of raising) — `.get(mod, ())` below yields an empty tuple, i.e.
# "catch nothing", for it, and that is still true after this feature.
#
# doc_text is NOT exception-free the way a stale version of this comment
# once claimed: `errors="replace"` only guards the bytes->str DECODE step,
# not csv.reader's or json.loads's own structural parsing, both of which can
# raise past it (csv.Error, RecursionError — see _reraise_text_error above).
#
# This lives here, keyed by module, rather than each converter module raising
# docs.py's own exception classes itself: docs.py already imports every
# converter module to populate CONVERTERS, so a converter importing back
# from docs.py to raise UnreadableDocument/EncryptedDocument would be a
# circular import. A converter module knowing its own library's exception
# TYPES is unavoidable (only that module imports that library) but deciding
# what those types MEAN in es's shared error catalogue is policy that
# belongs in one auditable place — the same reason _reraise_pdf_error already
# lived here before this feature added a second format family.
_CONVERSION_ERRORS = {
    doc_pdf: (PdfminerException,),
    doc_office: (PackageNotFoundError, InvalidFileException, BadZipFile,
                 OSError, ValueError, ParseError, XMLSyntaxError,
                 KeyError, AttributeError),
    doc_text: (csv.Error, RecursionError),
}


def _reraise_conversion_error(mod, real: Path, exc: Exception) -> None:
    """Always raises. Dispatches to the format-family-specific mapping for
    `mod` — the one place a converter's raw library exception is translated
    into the shared {EncryptedDocument, UnreadableDocument} catalogue."""
    if mod is doc_pdf:
        _reraise_pdf_error(real, exc)
    if mod is doc_office:
        _reraise_office_error(real, exc)
    if mod is doc_text:
        _reraise_text_error(real, exc)
    # Defensive fallback, not expected to fire today: every module currently
    # in _CONVERSION_ERRORS is handled by name above. Kept so a future module
    # added to that table without a matching branch here still fails safe
    # (a catalogue code) instead of leaking whatever it raised.
    raise UnreadableDocument(
        f"{real.name} could not be read; ask the user to resend it") from exc


def _page_count(mod, real: Path) -> Optional[int]:
    """None means "this format has no concept of pages" (a flat format like
    .csv/.txt/.json/.md, all served by doc_text) — kept a distinct value
    from 0 or 1:

    - NOT 0: for a format that DOES paginate, _page_count already treats a
      report of zero pages as an error below (a PDF that opens but claims no
      pages is corrupt/empty, not legitimately "paginated with nothing" — an
      indistinguishable value here would erase that distinction and make a
      flat format look broken).
    - NOT 1: a page count of 1 would suggest page-range/render semantics
      apply ("page 1 of 1"), which is false for something that isn't
      paginated at all.

    The dispatch signal is simply whether the converter module exposes
    `page_count` — a flat converter that never implements it is, by
    construction, a format with no pages.

    The exception mapping below (PdfminerException -> the PDF-specific
    encrypted/unreadable errors) is deliberately PDF-specific and today only
    ever fires for doc_pdf (the only converter with a page_count). A future
    paginated converter with its own failure modes gets its own mapping at
    its own call site — this function does not try to generalize that part
    ahead of need.
    """
    if not hasattr(mod, "page_count"):
        return None
    try:
        total = mod.page_count(real)
    except PdfminerException as e:
        _reraise_pdf_error(real, e)
    if total <= 0:
        # Opens fine but has no pages — pdfplumber does not consider this an
        # error, but a real document always has at least one page, and a
        # silent empty result is indistinguishable from "genuinely blank" for
        # the agent. Treat it as the same class of problem as corruption.
        raise UnreadableDocument(
            f"{real.name} has no readable pages — it may be corrupt or "
            "empty; ask the user to resend it")
    return total


_VAULT_PREFIX = "$vault/"


def _expand_source(source: str) -> str:
    """Rewrite the two accepted vault-relative `source` forms into a path
    string for paths.resolve_readable to confine — expansion never
    substitutes for confinement, it only decides what candidate string gets
    checked next:

    - an absolute path: returned untouched. Checked FIRST and unconditionally,
      so an absolute path is never reinterpreted as vault-relative. This is
      also how every Telegram upload arrives (Hermes injects an absolute
      path), so it stays the primary form for uploads — there is no
      cache-relative form to expand.
    - "$vault/..." (the prefix must include the slash, so a literal file
      named "$vault" with no trailing slash is never treated as the prefix —
      it just falls through to the bare-relative case below): an explicit,
      unambiguous synonym for the vault-relative form. Looked up by name here
      (docs.py owns "which root is 'the vault'"); paths.py never learns that
      name.
    - anything else (a bare relative path): joined onto the vault root — the
      same convention es_notes_read/es_notes_attach/es_notes_list already use
      (they hand back vault-relative paths like "Topics/Soccer/schedule.pdf"
      for exactly this source), so "$vault/X" and bare "X" are interchangeable.

    A "$vault/../../etc/passwd" (or bare "../../etc/passwd") traversal is NOT
    rejected here — Path joining does not normalize ".." — it is rejected by
    resolve_readable() next, exactly the same way any other out-of-root
    absolute path is, because that call site resolves (normalizes) the path
    THEN checks containment. Expansion only changes what string reaches that
    check, never whether the check runs.
    """
    s = str(source)
    if os.path.isabs(s):
        return s
    if s.startswith(_VAULT_PREFIX):
        return str(config.vault_root() / s[len(_VAULT_PREFIX):])
    return str(config.vault_root() / s)


def _prepare(source: str, roots, cache_root: Path):
    """Shared front half: purge, expand, confine, size-check, dispatch-check.

    Confinement runs before the size/extension checks below, but AFTER the
    stale-artifact purge — the purge only ever touches our own `.es/`
    namespace under `cache_root` (see doc_cache.purge), never the candidate
    `source` path, so an unauthorized caller cannot use it to affect anything
    outside the cache they already have no special claim over. Ordering the
    purge first just means every call — authorized or not — keeps the cache
    tidy; it does not weaken confinement.
    """
    doc_cache.purge(cache_root)
    expanded = _expand_source(source)
    try:
        real = paths.resolve_readable(expanded, roots)
    except paths.SourceNotFound as e:
        # paths.py is deliberately generic; this is where the document-domain
        # remedy belongs. Uploaded files are cache-evicted after 24h, and the
        # agent has no other way to learn that — so name the fix, not just
        # the symptom.
        raise paths.SourceNotFound(
            f"{e} — if this was a recently uploaded file, it may have aged "
            "out of the cache (uploads are only kept for 24 hours); ask the "
            "user to resend it") from e
    except paths.SourceForbidden as e:
        # Same reasoning as SourceNotFound above: paths.py stays generic (and
        # — critically — identical whether or not the path exists, closing a
        # probing oracle), so the document-specific remedy is added here, not
        # there. The remedy MUST NOT depend on anything paths.py doesn't
        # already expose (existence, real target, ...) or the oracle reopens
        # — everything appended below is static text plus the untouched `{e}`,
        # never anything that differs between an existing and a missing path.
        # Naming the accepted forms here (not just "uploads or the vault") is
        # the actual fix for the bug this feature exists for: a vault-relative
        # source used to land here — doc_forbidden — for what was really a
        # path-FORM mistake, not an authorization one.
        raise paths.SourceForbidden(
            f"{e} — source must be an absolute path inside the vault or "
            "uploads, \"$vault/...\", or a vault-relative path (e.g. "
            "\"Topics/Manual.pdf\"); ask the user to resend it as a Telegram "
            "upload, or save it into the vault first") from e
    except OSError as e:
        # resolve_readable's own existence check can still hit the filesystem
        # AFTER confinement already passed (e.g. ENAMETOOLONG on a path that
        # resolves lexically but blows past PATH_MAX/NAME_MAX). That is an
        # operational failure inside an already-allowed root, not a
        # confinement bypass — map it to a domain error instead of leaking a
        # raw OSError (and the full offending path) to the agent.
        raise UnreadableDocument(
            "that path is too long or otherwise unusable on this "
            "filesystem; ask the user to resend the file") from e

    ext = real.suffix.lower()
    if ext not in SUPPORTED:
        raise UnsupportedDocument(
            f"cannot read {ext or 'files without an extension'} — "
            f"supported: {', '.join(sorted(SUPPORTED))}")
    try:
        size = real.stat().st_size
    except OSError as e:
        raise UnreadableDocument(
            "could not read this file from disk; ask the user to resend it"
        ) from e
    if size > MAX_DOCUMENT_BYTES:
        raise DocumentTooLarge(
            f"{real.name} is larger than the {MAX_DOCUMENT_BYTES // (1024*1024)}MB "
            "limit; ask for a smaller file or a page range")
    did = doc_cache.doc_id(real, ext)
    adir = doc_cache.artifact_dir(cache_root, did)
    return real, did, adir


def _read_cached_markdown(adir: Path) -> Optional[str]:
    """None means "treat as a cache MISS" — both for a plain absence and for
    an unreadable/undecodable doc.md. doc_id is a CONTENT hash, so a doc.md
    truncated by a crash or ENOSPC would otherwise poison every future
    extract of this exact document for the full 24h TTL (identical content ->
    identical doc_id -> the same broken cache entry every time). Reconverting
    and overwriting on a bad read is strictly better than raising — mirrors
    how _read_images_manifest already treats a broken images.json."""
    md_path = adir / DOC_MD_NAME
    if not md_path.is_file():
        return None
    try:
        return md_path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError):
        return None


def _read_images_manifest(adir: Path) -> List[str]:
    manifest = adir / DOC_IMAGES_MANIFEST
    if not manifest.is_file():
        return []
    try:
        return json.loads(manifest.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return []


def _write_full_extract(adir: Path, markdown: str, images: List[Path]) -> None:
    (adir / DOC_MD_NAME).write_text(markdown, encoding="utf-8")
    (adir / DOC_IMAGES_MANIFEST).write_text(
        json.dumps([str(i) for i in images]), encoding="utf-8")


# Matches the "\n\n" that doc_pdf.convert's "\n\n".join(parts) puts before
# every page's "## Page N" heading except the very first — i.e. every safe
# block-boundary cut point in the markdown.
_PAGE_BOUNDARY_RE = re.compile(r"\n\n(?=## Page (\d+)\b)")
# The very first heading in the markdown — for a page-SUBSET extract (e.g.
# pages="5-8") this is NOT page 1, so the no-earlier-boundary fallback must
# read the real number here rather than assume 1.
_FIRST_PAGE_RE = re.compile(r"\A## Page (\d+)\b")
# A markdown image link as doc_pdf.convert emits it: "![page N](/path.png)".
# Used only to keep a HARD (no-boundary-available) cut from landing inside
# one — the boundary cut above never can, since links are their own block.
_IMAGE_LINK_RE = re.compile(r"!\[[^\]]*\]\([^)]*\)")


def _safe_hard_cut(text: str, limit: int) -> int:
    """A cut index <= limit that never lands inside a markdown image link —
    landing inside one would hand the agent a truncated, unusable path."""
    cut = min(limit, len(text))
    for m in _IMAGE_LINK_RE.finditer(text):
        if m.start() < cut < m.end():
            return m.start()
    return cut


# Every converter's own self-truncation marker — doc_text (CSV/JSON/txt/md),
# doc_office (docx/xlsx), and doc_ics all follow the same "\n\n*(truncated
# ...)*" convention this module's own PDF marker uses (see each module's own
# MAX_CHARS/MAX_ICS_CHARS comments), because none of them can rely on
# _truncate_markdown below to do it for them — that function only knows
# "## Page N" PDF-style boundaries. Anchored on the literal "*(...)* "
# wrapper plus the word "truncated" inside it (not just the bare word, which
# could coincidentally appear in real document content) to keep false
# positives effectively impossible without requiring every converter module
# to change its (markdown, images) return contract just to also hand back a
# flag — a change that would reach into doc_office.py/doc_ics.py, files this
# fix does not own.
_SELF_TRUNCATION_MARKER_RE = re.compile(r"\*\([^()]*\btruncated\b[^()]*\)\*")


def _converter_self_truncated(markdown: str) -> bool:
    """True if a converter already truncated ITS OWN output (before this
    module's outer MAX_MARKDOWN_CHARS cap ever ran) and said so in-band.
    Needed because every non-PDF converter self-truncates at its own,
    smaller budget (see doc_text.MAX_CHARS / doc_office.MAX_CHARS /
    doc_ics.MAX_ICS_CHARS, all well under this module's 40_000), so the
    result docs.extract() gets back is very often already under
    MAX_MARKDOWN_CHARS by the time _truncate_markdown looks at it — which
    otherwise reports `truncated: False` even though real content was
    genuinely cut, contradicting es_doc_extract's own docstring that
    `truncated` is the agent's signal to look for more."""
    return bool(_SELF_TRUNCATION_MARKER_RE.search(markdown))


def _truncate_markdown(markdown: str, total_pages: Optional[int]) -> Tuple[str, bool]:
    """Cut `markdown` to at most MAX_MARKDOWN_CHARS at a page-block boundary
    (the last "## Page N" heading before the limit), never mid-word or
    mid-image-link, and append a marker naming where extraction stopped and
    how to resume. Returns (markdown, truncated).

    If even the FIRST page's content alone exceeds the limit there is no
    earlier block boundary to cut at — a page-range resume can't help (the
    identical oversized page would just come back truncated the same way
    again), so that case falls back to a hard cut (still guarded against
    landing inside an image link) and says plainly why a resume marker isn't
    offered, rather than pretending a narrower range would fix anything.
    Reads the real page number off the first heading rather than assuming 1
    — a page-SUBSET extract (e.g. pages="5-8") starts with "## Page 5", not 1.

    `total_pages` is None for a non-paginated format (every converter except
    doc_pdf — see docs._page_count). The "page N" wording above only ever
    makes sense when a document actually HAS pages: a "## Page N" boundary
    can only occur in doc_pdf's own output (no other converter emits that
    literal heading text), so `candidates` is guaranteed empty whenever
    `total_pages is None` and only the hard-cut branch below can fire for a
    flat format — which is exactly why that branch, and only that branch,
    needs a second, page-free wording.
    """
    if len(markdown) <= MAX_MARKDOWN_CHARS:
        return markdown, False

    boundaries = [(m.start(), int(m.group(1)))
                  for m in _PAGE_BOUNDARY_RE.finditer(markdown)]
    candidates = [b for b in boundaries if b[0] <= MAX_MARKDOWN_CHARS]

    if candidates:
        cut_pos, next_page = max(candidates)
        stopped_after = next_page - 1
        marker = (f"\n\n*(truncated after page {stopped_after} of {total_pages} "
                   f"— call es_doc_extract again with pages=\"{stopped_after + 1}-"
                   f"{total_pages}\" to continue)*")
        return markdown[:cut_pos] + marker, True

    cut_pos = _safe_hard_cut(markdown, MAX_MARKDOWN_CHARS)
    if total_pages is None:
        # A flat format (csv/json/txt/md/ics/docx/xlsx): there are no pages
        # to narrow and no es_doc_render to fall back to (render() itself
        # refuses any format without page_count — see render()'s own
        # UnsupportedDocument check) — naming either would send the agent
        # chasing a remedy that provably doesn't exist for this document.
        # The honest answer is that one indivisible block (a row, a
        # paragraph, an event) is simply too large, with no narrower view
        # available at all.
        marker = (f"\n\n*(truncated — a single block of content exceeds the "
                   f"{MAX_MARKDOWN_CHARS}-character limit with no earlier "
                   "boundary to stop at, and this format has no narrower "
                   "view to fall back to; ask the user for a smaller/"
                   "narrower version of this document)*")
        return markdown[:cut_pos] + marker, True

    first_match = _FIRST_PAGE_RE.match(markdown)
    first_page = int(first_match.group(1)) if first_match else 1
    marker = (f"\n\n*(truncated inside page {first_page} — its content alone "
              f"exceeds the {MAX_MARKDOWN_CHARS}-character limit, so there is "
              "no earlier page boundary to stop at; narrowing pages won't "
              f"help since page {first_page} alone is already too large — try "
              "es_doc_render on this page instead)*")
    return markdown[:cut_pos] + marker, True


def extract(source: str, roots, cache_root: Path,
            pages: Optional[str] = None) -> dict:
    real, did, adir = _prepare(source, roots, cache_root)
    ext = real.suffix.lower()
    mod = CONVERTERS[ext]
    total = _page_count(mod, real)

    # `pages` on a format with no pages (total is None) is a loud error, not
    # a silent no-op — same philosophy as render()'s explicit-out-of-range
    # behavior below: an argument that cannot mean anything for this
    # document is more likely a mistaken assumption about its format than an
    # intentional "give me everything" request (which is already spelled by
    # omitting pages entirely).
    if pages is not None and total is None:
        raise InvalidPageRange(
            f"{ext} documents do not have pages — omit the pages argument "
            "to extract the whole document")
    wanted = parse_pages(pages, total) if pages is not None else None

    if wanted is None:
        # Full-document extract: doc_id is a content hash, so a previous
        # conversion of this exact content is still correct — check the
        # cache before paying for another convert().
        cached = _read_cached_markdown(adir)
        if cached is not None:
            doc_cache.touch(adir)  # TTL means "24h since last USE", not
                                    # "24h since conversion" — a cache hit is
                                    # a use.
            markdown = cached
            images = _read_images_manifest(adir)
        else:
            try:
                markdown, image_paths = mod.convert(real, adir, pages=None)
            except _CONVERSION_ERRORS.get(mod, ()) as e:
                _reraise_conversion_error(mod, real, e)
            _write_full_extract(adir, markdown, image_paths)
            doc_cache.touch(adir)
            images = [str(i) for i in image_paths]
    else:
        # A page-SUBSET extract is never written to doc.md/images.json: those
        # two files are the whole-document artifact that a future es_read
        # will address as `doc:<id>` (Phase 2). Writing a subset there would
        # silently replace the full document with a fragment for every
        # future reader of this doc_id — worse than not caching at all. So
        # subsets always convert fresh and are simply never persisted; only
        # the full-document result is cached.
        try:
            markdown, image_paths = mod.convert(real, adir, pages=wanted)
        except _CONVERSION_ERRORS.get(mod, ()) as e:
            _reraise_conversion_error(mod, real, e)
        doc_cache.touch(adir)
        images = [str(i) for i in image_paths]

    # Checked against the PRE-outer-cap markdown (whether it just came out of
    # convert() or out of the cache): a converter's own self-truncation
    # marker is the only signal docs.py gets that IT already cut real
    # content at its own, smaller budget — _truncate_markdown below only
    # ever sees "## Page N" boundaries, so it can't detect that on its own.
    self_truncated = _converter_self_truncated(markdown)
    markdown, truncated = _truncate_markdown(markdown, total)
    truncated = truncated or self_truncated
    # kind is trivially the extension without its dot ("pdf" for ".pdf") —
    # correct today and requires no per-format table of its own.
    return {"doc_id": did, "kind": ext.lstrip("."), "page_count": total,
            "markdown": markdown, "images": images, "truncated": truncated}


DEFAULT_RENDER_PAGES_HI = 10


def render(source: str, roots, cache_root: Path, pages: Optional[str] = None) -> dict:
    """pages=None means "no explicit request" — render the first
    DEFAULT_RENDER_PAGES_HI pages, clamped down to the document's actual
    length so a short document (the 1-3 page schedule this feature exists
    for) doesn't error just because it's shorter than the default window.

    An EXPLICIT pages argument is never clamped: if the agent asks for a
    range that runs past the document's end, that's a loud error (same as
    extract()'s explicit-range behavior) rather than a silent partial
    result — an explicit out-of-range ask is more likely a wrong page
    number than an intentional "give me what you can" request.
    """
    real, did, adir = _prepare(source, roots, cache_root)
    ext = real.suffix.lower()
    mod = CONVERTERS[ext]
    # Rasterizing pages only means something for a format that HAS pages to
    # rasterize — a converter without `render` (every flat format: txt/md/
    # csv/json, all via doc_text) can't support this call at all, independent
    # of whatever `pages` was passed, so the capability check comes first and
    # names the actual reason rather than failing confusingly deeper down
    # (e.g. inside parse_pages against a page_count that doesn't exist).
    # hasattr is enough: it mirrors the same "presence of the optional
    # attribute IS the capability" convention _page_count already uses for
    # page_count, so a reader only has to learn the pattern once.
    if not hasattr(mod, "render"):
        raise UnsupportedDocument(
            f"{ext} cannot be rendered to images — es_doc_render can only "
            "render pages of a PDF; try es_doc_extract instead")
    total = _page_count(mod, real)
    spec = pages if pages is not None else f"1-{min(DEFAULT_RENDER_PAGES_HI, total)}"
    wanted = parse_pages(spec, total)
    if len(wanted) > MAX_RENDER_PAGES:
        raise InvalidPageRange(
            f"cannot render {len(wanted)} pages in one call — the limit is "
            f"{MAX_RENDER_PAGES}; narrow the range (or call es_doc_render "
            "again for the rest)")
    images = mod.render(real, adir, wanted)
    doc_cache.touch(adir)
    return {"doc_id": did, "page_count": total,
            "images": [str(i) for i in images]}
