"""Word (.docx) and Excel (.xlsx) documents -> Markdown.

EMBEDDED IMAGES (.docx only — see the "XLSX IMAGES" note far below):
governed by the same rule doc_pdf.py's module docstring states for PDFs:
every kind of content in a document comes out as the thing it is, an image
extracted because it exists, not because it scored above a threshold. A
`.docx` photo used to be silently invisible to the agent (this module
returned `(markdown, [])` unconditionally); now every embedded raster image
is extracted to a sibling file and linked inline at its position in the
body's document order, mirroring how doc_pdf interleaves `_page_image_entries`
into `_assemble_page` — see `_iter_body_blocks` below for the .docx
equivalent.

ORIGINAL BYTES, NOT A RE-RENDER: a `.docx` package stores the image's own
original file (`word/media/imageN.<ext>`) — unlike a PDF, where an
embedded image's ON-PAGE appearance can be rotated/scaled/masked by a
separate placement matrix, so doc_pdf.py has to ask pdfium to RENDER the
object to reproduce that appearance faithfully (see doc_pdf's "CROP-RENDER
VS STREAM EXTRACTION" note). WordprocessingML has no equivalent placement
matrix for an inline picture's pixel content — `<w:drawing>` positions and
sizes the image as a whole (via `<wp:extent>`/`<a:xfrm>`), but never
resamples or re-encodes the underlying picture — so extracting
`ImagePart.blob` directly (python-docx's own decoded access to those exact
bytes) is both CHEAPER (no decode-then-render pass, just a bytes read
already sitting in memory from opening the package) and higher-fidelity
(the original compressed file, not a re-encoded copy) than rendering would
be. `ImagePart.image.ext` is kept as the output file's own extension
(png/jpeg/gif/bmp/...) instead of forcing everything to `.png`, for the
same reason: it costs nothing extra and avoids a lossy re-encode for a
lossy source format (a re-saved JPEG is not bit-identical to the original).

WHERE AN IMAGE LINK LANDS: a paragraph's own inline images are emitted
right after that paragraph's text block (or as their own block, if the
paragraph has no text at all — the common shape python-docx's own
`add_picture` produces, and the shape a Telegram-forwarded Word doc with
an inserted photo actually has). A TABLE CELL's image is handled
differently: emitted as its own block immediately AFTER the whole table,
never inside a pipe-table cell. A Markdown link inside a cell
(`| ![x](path) |`) is syntactically a table cell like any other, but if
the image sits alongside real cell text, or the path contains a `|`-free
but otherwise arbitrary string, the risk of subtly misaligning the row is
real and not worth it for a value doc_office.py already avoids elsewhere —
`format_cell` exists precisely because arbitrary text can break a pipe
table (escaping embedded `|`, collapsing embedded newlines). Simpler and
safer to keep pipe-table cells to their existing plain-text contract and
name the image's origin explicitly instead: `_table_block` emits
`![embedded image N — table row R, column C](path)` after the table, so
the agent still learns both that the image exists and which cell it came
from, without touching the table's own syntax at all.

DUPLICATE IMAGES: the SAME embedded image (the same `r:embed` relationship
id, e.g. a letterhead logo placed at both the top and bottom of a
template) is written to disk exactly ONCE, and every occurrence in the
document links to that same file — never re-extracted, and never counted a
second time against MAX_EXTRACTED_IMAGES. This is not a "same file appears
twice, is that one item or two" ambiguity — it is the SAME relationship id,
i.e. python-docx's own model already asserts these are one underlying image
part, not two independently-inserted ones (two independent insertions of
visually-identical bytes, e.g. pasting the same photo twice, get their OWN
relationship ids and are extracted/counted as two separate images, exactly
as their two separate appearances in the reading order warrant). Writing
the same bytes to disk twice would be pure waste with no benefit: the agent
gets the same picture either way, and the markdown still shows two links
(one at each position the image actually appears), each correctly pointing
at that one file — the "extracted because it exists" rule is about every
APPEARANCE in the reading order being represented, not about paying to
duplicate bytes that are provably identical by construction.

MAX_EXTRACTED_IMAGES (.docx's own, NOT doc_pdf.MAX_EXTRACTED_IMAGES): a
separate constant, deliberately not imported from doc_pdf, for the same
reason doc_pdf.MAX_EXTRACTED_DRAWINGS is kept separate from
doc_pdf.MAX_EXTRACTED_IMAGES — two independent modules, with genuinely
different per-item costs (a `.docx` extraction is a bytes copy already in
memory; a PDF extraction is pdfium rendering an object), coupling one
module's ceiling to another's module-level constant for no reason beyond
"the number happens to match today" is exactly the kind of unrelated
coupling this codebase avoids elsewhere. Set to the same value as
doc_pdf's (500) anyway, on the same reasoning: generous headroom above any
real Word document's realistic embedded-photo count (a report or log with
dozens to a couple hundred inline photos) while remaining a real, finite
bound. Hitting it is reported IN-BAND exactly once, at the end of the
document (see `_iter_body_blocks`'s final `if extractor.skipped` block) —
never a silent drop. Unlike doc_pdf, there is no per-page lookahead here (a
`.docx` body is one flat stream, not pages iterated with a known
remaining-content-on-this-page bound), so the count of skipped images is
only known once the whole body has actually been walked; a document cut
short first by MAX_CHARS never reaches that point, which is fine — the
MAX_CHARS truncation marker already explains why the rest of the document,
images included, wasn't processed.

XLSX IMAGES (investigated, NOT implemented — recorded here so the decision
doesn't have to be re-researched): an `.xlsx` absolutely can carry embedded
raster images too — a logo or photo floated over a sheet
(`xl/drawings/drawingN.xml`, backed by `xl/media/*`, the exact same
DrawingML anchoring model `.docx` uses for its own `<w:drawing>`) and a
chart object (`xl/charts/chartN.xml`, a nativechart definition, not a
raster image at all — nothing to "extract" as a file the way a photo is).
Deliberately out of scope here: a later plan converts spreadsheets to a
queryable database (`es_doc_query`, not markdown/es_read — see docs.py's
`TABLE_KINDS` note) rather than a markdown table, so any embedded-image
work done against `_convert_xlsx`'s current markdown-table output would be
thrown away once that lands. `_convert_xlsx` therefore still always
returns `(markdown, [])`, unchanged.

.docx: walked in DOCUMENT ORDER, not python-docx's separate `.paragraphs`/
`.tables` lists — those are two independent flat lists that do not preserve
how paragraphs and tables were actually interleaved in the source (a table
sitting between two paragraphs would come out AFTER both, since each list is
walked to completion on its own). The fix is to walk `document.element.body`
directly: the underlying WordprocessingML XML has one true order (`<w:p>` and
`<w:tbl>` children of `<w:body>`, in document order), and python-docx's
higher-level `Paragraph`/`Table` wrappers can be constructed straight from
those child elements. This module never asks `document.paragraphs` or
`document.tables` for order — only for nothing, since those attributes are
never read here at all.

Heading paragraph styles ("Heading 1".."Heading 6") map to `#`.."######"` so
the document's own outline becomes the `##`-boundary structure the later
paging reader depends on (see docs.py's CONVERTERS docstring); a non-heading
paragraph is emitted as plain text.

.xlsx: one `## <sheet name>` heading per worksheet — the natural (and only)
structural unit a spreadsheet has, which is exactly what makes a workbook
pageable by that same reader. Rows render as a pipe table via doc_support.

Every sheet converts IN FULL now — MAX_CHARS is a generous resource ceiling
(see below), not a context-window budget, so in practice every sheet's
heading AND every one of its rows survive. The fair-share machinery this
paragraph used to describe as the fix for "sheet 2 and 3 vanish entirely"
is kept anyway, as the BACKSTOP for the (now rare, but not impossible —
dozens of enormous sheets can still exhaust even a 50MB ceiling) case where
the ceiling genuinely is hit: every sheet's heading is emitted
UNCONDITIONALLY, and `_convert_xlsx` fair-shares whatever budget remains
across the CURRENT sheet and every sheet still to come (recomputed fresh
each iteration, so a sheet that uses less than its share leaves the surplus
for later ones, rather than a fixed up-front split wasting budget a short
sheet didn't need) — so even in that backstop case every sheet still gets a
heading, discoverable from es_read's outline. `_render_sheet_rows` also
always shows a sheet's first row regardless of remaining budget, which in
the overwhelmingly common case IS real content — but this is honestly a
"first row" guarantee, not a "real content" one: a sheet whose own row 1 is
entirely blank (no cells written on it at all — e.g. real data starting at
row 2) would show that blank row instead, if the ceiling were ever tight
enough to cut it off before row 2. A prior review flagged the module
docstring's older wording ("at least one row of REAL content") as
overstating this guarantee for exactly that reason; described accurately
here rather than restated. See
test_xlsx_later_sheets_reachable_when_first_sheet_exhausts_budget (now
exercising this backstop directly, via a monkeypatched MAX_CHARS, since the
real default ceiling is far too generous for a workbook of ordinary test
size to ever ration).

Formula cells are read as the FORMULA TEXT (`data_only=False`), never the
cached result (`data_only=True`). A workbook saved by any tool that never ran
Excel's calculation engine — including openpyxl itself, which is exactly the
shape of file a Telegram upload is likely to be — has NO cached value for a
formula cell at all: `data_only=True` against such a file silently renders
those cells as blank, not as "no cache available". The formula text is
always present and deterministic; a cache that only sometimes exists is not
a foundation this module can build predictable behavior on.

Both formats truncate at a whole-BLOCK boundary against a shared RESOURCE
ceiling (MAX_CHARS below) — a whole paragraph/table for .docx, a whole row
for .xlsx — mirroring doc_text.MAX_CHARS / doc_ics.MAX_ICS_CHARS: neither
format here has PDF-style "## Page N" boundaries to cut at, so each
truncates itself before returning. MAX_CHARS is deliberately NOT sized to
fit inside one MCP response — es_doc_extract's own response is now a small
receipt (a fixed-size preview plus a `doc:<id>` handle), never the document
itself, so there is no response-sized budget for this ceiling to clear at
all — it exists so that a genuinely pathological document (a zip-bomb-style
.docx, a workbook someone managed to inflate to gigabytes of text) still
can't make doc.md, and therefore es_read's outline, unbounded; see
MAX_CHARS's own comment for the sizing rationale. Every ordinary document —
including every deliberately oversized one this module's own test suite
builds — converts in full and never reaches it.

XLSX_MAX_ROWS/XLSX_MAX_COLS are a SEPARATE, structural cap on top of that
ceiling: a sheet's reported used range can be dramatically larger than its
real content (a single stray value at, say, ZZ10000 reports a 10000-row x
702-column used range for what is, in substance, two cells of data), and
that must be bounded before the character ceiling ever gets a chance to
look at it, or a single pathological sheet could force iterating tens of
thousands of all-blank rows for no reason. They are set to Excel's OWN
real per-sheet maximums (1,048,576 rows, 16,384 columns — XFD) rather than
a tuned smaller number: a legitimate sheet's real `max_row`/`max_column`
can never exceed them (openpyxl itself refuses to write past those
indices), so for any real, honestly-dimensioned sheet these caps are true
no-ops — every sheet converts fully — and they bind ONLY the fallback case
(dimension unknown, or a corrupt/adversarial `<dimension>` claiming more
than the format allows) where they take over as the outer iteration bound.
Measured (see doc_office's test suite): even that worst case — no known
dimension, iterating the full 1,048,576-row fallback on a narrow (2-column)
sheet — completes in ~1.3s; a wide worst case is bounded far sooner by the
character ceiling itself (a 50-column-wide, full-height sparse sheet hit
MAX_CHARS after ~329k of 1,048,576 rows in ~1.9s). DOCX_MAX_TABLE_ROWS is
the .docx analogue: it bounds the cost of rendering a SINGLE table block,
independent of MAX_CHARS, for the same reason (see `_table_block`) — kept
at its original tuned value (2000) since it protects against a genuinely
separate, real per-block cost (each table row means constructing real
python-docx cell/paragraph/run objects to read `.text`, proportional to
table size, not a redundant per-access lookup like the paragraph-style
issue below).

.docx conversion is deliberately LAZY end to end (`_iter_body_blocks` is a
generator; `_convert_docx` stops pulling from it the moment MAX_CHARS is
spent) rather than "collect every block, then truncate": this was a verified
unbounded-cost bug, not a style preference. Measured before the fix, walking
`document.element.body` fully and formatting every block before truncating
cost ~0.4ms/paragraph — 3.7s at 10,000 paragraphs, 111.65s at 300,000 (a
plain 0.84 MB .docx, well inside MAX_DOCUMENT_BYTES) — because
`Paragraph.style` does a linear style-collection lookup on every call
(re-walking and re-type-converting every `<w:style>` in the document's
styles part from scratch, per paragraph — see `_build_heading_levels`'s
docstring for the full profiling breakdown), and that cost was paid for
every block in the document. This made the lazy generator load-bearing:
stopping the walk the moment MAX_CHARS was spent was the only thing keeping
that cost from being paid for the WHOLE document.

That per-access style cost is now fixed at its root (`_build_heading_levels`
resolves every style id to a heading level ONCE, per document, not per
paragraph — see its docstring for the measured before/after). Profiling
full (unbudgeted) conversion after that fix turned up a SECOND per-call
cost of the same shape — `Paragraph.text`/`CT_P.text` internally calling
`self.xpath(...)` on every paragraph — fixed the same way: `_paragraph_text`
and `_paragraph_style_id` below re-walk the already-parsed lxml tree
directly (plain `iterchildren()`/`find()`/`get()` calls against
pre-resolved Clark-notation tag names) instead of going through
python-docx's own text/style element properties, which re-resolve their
own tag names and re-run an xpath expression on every single call. See
each function's own docstring for the exact profiling numbers and the
equivalence check against python-docx's own output.

With all three per-access costs removed, full (unbudgeted) conversion of a
300,000-paragraph, single-run-per-paragraph .docx measured (this machine,
same file for both sides of the comparison) at:

    paragraphs   before (unbudgeted, per-call style+text)   after
    10,000       4.062s                                     0.064s
    100,000      40.714s                                    0.624s
    300,000      124.301s                                   1.978s

(the "before" column re-measures, on this machine, the same per-call
`Paragraph.style`/`Paragraph.text` walk this module used before today's
fix — independently reproducing the 111.65s/300,000-paragraph figure a
prior profiling pass reported on different hardware, same order of
magnitude). ~63-65x faster across the board — i.e. full conversion is now
fast enough that MAX_CHARS (now a generous
resource ceiling rather than a small context budget — see above) is never
reached in practice for a document of this shape; see
test_docx_huge_paragraph_count_converts_quickly_and_bounded. The lazy
generator is kept regardless — it is still correct, still cheap, and still
the real defence against a still-plausible unbounded case: a `.docx` is a
zip, so MAX_DOCUMENT_BYTES only bounds the COMPRESSED size (a zip bomb can
inflate far past it), and fixing these per-access costs does not change
that a pathological document could still contain more blocks than any
resource ceiling should render in full — the lazy walk plus
DOCX_MAX_TABLE_ROWS bound the COST of touching such a document, not just
the characters it emits.

.xlsx: `ws.max_row`/`ws.max_column` in `read_only` mode come from the sheet
XML's `<dimension>` element — and are `None`, not `0`, when that element is
absent. This is not exotic: `openpyxl.Workbook(write_only=True)` (a
mainstream way tools generate large spreadsheets — pandas/ETL exports
included) never writes `<dimension>` at all. A prior version of this module
read `total_rows = ws.max_row or 0`, so `None or 0` produced `0`, which was
then read as "genuinely empty sheet" and short-circuited before a single row
was ever iterated — a verified DATA-LOSS bug: a 300,000-row write_only
workbook converted to a bare `"## Sheet"` heading with no error, no
truncation marker, and `{ok: true}`. `_render_sheet_rows` below never trusts
`ws.max_row`/`ws.max_column` as a presence signal any more — only actually
iterating and observing a row (`saw_any`) proves a sheet has content, dimension
or no dimension. `ws.calculate_dimension(force=True)` was considered as a fix
and measured at ~7.9s / a full read-only pass over a 300,000-row sheet (it
walks every row to find the true extent) — paying that cost just to print an
accurate row count would reintroduce the same O(sheet size) cost this module
otherwise avoids by rendering read_only rows lazily, so it is deliberately
never called here. When the sheet's true dimension is unknown, this module
reports what it can prove (rows actually rendered, and whether more exist
past the structural cap) without fabricating an exact total.
"""
from pathlib import Path
from typing import Dict, Iterator, List, Optional, Tuple
from xml.etree.ElementTree import ParseError
from zipfile import BadZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from lxml.etree import XMLSyntaxError
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException

from es import doc_cache
from es.capabilities.doc_support import (ParseFailed, format_cell, format_row,
                                          truncation_marker)

# A real RESOURCE ceiling, not a context-window budget — see the module
# docstring's "convert fully, bound resources not context" note. Shared by
# both formats, same as the old MAX_CHARS it replaces. Sized generously (the
# same order of magnitude as docs.MAX_DOCUMENT_BYTES, the 50MB upstream input
# cap) rather than to fit inside one MCP response: doc.md is a 24h-TTL cache
# entry es_read pages by section/line, not something returned whole, so the
# only real costs a full conversion imposes are disk (cheap at this size) and
# the size of es_read's outline (built from "## " headings — still trivially
# fast to scan at tens of megabytes). This is deliberately NOT sized against
# any MCP-response-sized number at all — es_doc_extract's own response is a
# small fixed-size receipt (a preview plus a handle) regardless of how big
# the source document is, so there is no response budget for this ceiling to
# clear; it protects only what gets cached and paged, never what gets
# returned. A .docx/.xlsx is a zip, so MAX_DOCUMENT_BYTES only bounds the
# COMPRESSED input size — a pathological file could still try to inflate far
# past this ceiling, which is exactly why this ceiling exists at all (see
# also the lazy .docx walk and the per-table/per-sheet structural caps
# below, which bound the COST of touching such a file, not just the
# characters it emits).
MAX_CHARS = 50_000_000

# .docx embedded-image extraction ceiling — deliberately its OWN constant,
# not doc_pdf.MAX_EXTRACTED_IMAGES, even though it is set to the same value.
# See the module docstring's "MAX_EXTRACTED_IMAGES" note for why coupling
# this module's ceiling to doc_pdf's would be the wrong kind of coupling.
MAX_EXTRACTED_IMAGES = 500

# --------------------------------------------------------------------------
# Parse-step error mapping (see doc_support.ParseFailed for the boundary
# rule this exists to enforce).
# --------------------------------------------------------------------------
#
# python-docx and openpyxl each raise their OWN exception type for a parse
# failure — PackageNotFoundError (python-docx, which itself normalizes a
# bad/missing zip into this one type), BadZipFile (openpyxl, which does NOT
# normalize it), InvalidFileException (openpyxl's own filename-extension
# guard — included for defense in depth even though docs.py's dispatch
# already only ever routes a real .xlsx path here, so it should never
# actually fire in practice), OSError (python-docx's own "no valid workbook/
# document part" case), ParseError (openpyxl's read_only streaming reader
# uses the stdlib xml.etree parser regardless of lxml being installed),
# XMLSyntaxError (python-docx uses lxml), KeyError (a valid zip missing
# "[Content_Types].xml" entirely — both libraries do a plain dict-style
# archive lookup), and AttributeError (a valid zip whose
# "[Content_Types].xml" parses but has no default namespace, so python-docx's
# lxml class lookup never upgrades it to its own CT_Types wrapper, and the
# next attribute access on it fails). All nine are verified empirically (not
# guessed) against real python-docx/openpyxl behavior for every ordinary
# "wrong/damaged file" shape (a renamed extension, a partial download) — see
# tests/test_docs.py's realistic-malformed-documents case.
#
# This tuple is used ONLY around the two open/parse call sites below
# (_open_docx, _open_xlsx, and the lazy per-row XML parse in
# _safe_row_iter) — never around this module's own rendering logic (body
# walking, sheet/table formatting, budget tracking). A bug in THAT code
# that happens to raise one of these same ordinary types (a typo'd dict
# key, an attribute on a None) must surface as itself, not get relabeled
# "corrupt file" — see doc_support.ParseFailed's docstring for why the
# BOUNDARY, not the exception type, is what keeps the two apart.
_PARSE_ERRORS = (PackageNotFoundError, InvalidFileException, BadZipFile,
                 OSError, ValueError, ParseError, XMLSyntaxError,
                 KeyError, AttributeError)

# The OLE2/CFBF container signature (MS-CFB) — every legitimate, unencrypted
# .docx/.xlsx is a zip archive; a password-protected one is instead stored in
# this legacy container format (the same one .doc/.xls used), which is how
# real Office password-protection actually works, not a guess. Neither
# python-docx nor openpyxl exposes a distinct "needs a password" exception —
# both simply fail to open the file as a zip, indistinguishable by exception
# type alone from ordinary corruption (verified empirically against both
# libraries) — so this sniffs the file's own magic bytes instead.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _is_ole2_container(source: Path) -> bool:
    try:
        with open(source, "rb") as f:
            return f.read(len(_OLE2_MAGIC)) == _OLE2_MAGIC
    except OSError:
        return False


def _raise_parse_failed(source: Path, exc: Exception) -> None:
    """Always raises ParseFailed. Shared by both formats: neither
    python-docx nor openpyxl distinguishes "needs a password" from ordinary
    corruption by exception type, so both route through the same OLE2 sniff
    (see _is_ole2_container above)."""
    if _is_ole2_container(source):
        raise ParseFailed(
            f"{source.name} is password-protected — es cannot open "
            "encrypted Word/Excel documents; ask the user for an unlocked "
            "copy", encrypted=True) from exc
    raise ParseFailed(
        f"{source.name} could not be read as a Word/Excel document — it "
        "may be corrupt, truncated, or not actually a .docx/.xlsx file; "
        "ask the user to resend it") from exc


def _open_docx(source: Path) -> Document:
    """The ONLY place .docx parsing can fail: verified empirically that
    python-docx's `Document()` eagerly parses the entire package (zip +
    every XML part it needs) at open time — unlike openpyxl's read_only
    mode (see _open_xlsx/_safe_row_iter below), nothing about the walk in
    _convert_docx re-enters the library's own parser."""
    try:
        return Document(str(source))
    except _PARSE_ERRORS as e:
        _raise_parse_failed(source, e)


def _open_xlsx(source: Path):
    """Opens the workbook-level parts (workbook.xml, styles, shared
    strings) eagerly — but NOT a sheet's own XML, which read_only mode
    streams lazily (see _safe_row_iter: a truncated sheet1.xml only raises
    while actually ITERATING that sheet's rows, verified empirically, never
    here)."""
    try:
        return load_workbook(str(source), data_only=False, read_only=True)
    except _PARSE_ERRORS as e:
        _raise_parse_failed(source, e)


def _safe_row_iter(row_iter, source: Path):
    """Wrap ONLY the underlying iterator's own `next()` call — the exact
    point where openpyxl's read_only reader lazily parses the next chunk of
    a sheet's XML (verified empirically: a truncated sheet1.xml raises
    `xml.etree.ElementTree.ParseError` mid-iteration, not at `load_workbook()`
    time — read_only mode defers a SHEET's own parse to iteration even
    though the workbook-level parts are already parsed by then). Every
    caller's own row-processing logic (format_cell/format_row, budget
    tracking, `saw_any`/`kept` bookkeeping) stays entirely outside this
    function and outside its try block, reached only via the values this
    generator yields — so a bug in that logic can never be caught here and
    relabeled "corrupt file"."""
    while True:
        try:
            row = next(row_iter)
        except StopIteration:
            return
        except _PARSE_ERRORS as e:
            _raise_parse_failed(source, e)
        yield row

# Per-sheet structural caps — see module docstring. Independent of MAX_CHARS.
# Set to Excel's OWN real per-sheet maximums (the last row/column index the
# file format itself can represent — "XFD" is column 16,384) rather than a
# smaller tuned number: any honestly-dimensioned real sheet's true max_row/
# max_column can never exceed these, so for full conversion of real-world
# spreadsheets they never actually bind — they exist purely as the fallback
# iteration bound for a dimension-less sheet (see _render_sheet_rows) or a
# corrupt/adversarial `<dimension>` claiming more rows/columns than the
# format allows.
XLSX_MAX_ROWS = 1_048_576
XLSX_MAX_COLS = 16_384

# Per-table structural cap for .docx — see module docstring. A single
# pathological table (e.g. a 20,000-row spreadsheet paste) would otherwise be
# rendered in full as ONE block before the character budget ever gets a
# chance to reject it wholesale; this bounds that cost independent of
# MAX_CHARS, mirroring XLSX_MAX_ROWS.
DOCX_MAX_TABLE_ROWS = 2000

_HEADING_LEVELS = {f"Heading {n}": n for n in range(1, 7)}

_W_P = qn("w:p")
_W_TBL = qn("w:tbl")
_W_R = qn("w:r")
_W_HYPERLINK = qn("w:hyperlink")
_W_T = qn("w:t")
_W_TAB = qn("w:tab")
_W_BR = qn("w:br")
_W_CR = qn("w:cr")
_W_NOBREAKHYPHEN = qn("w:noBreakHyphen")
_W_PTAB = qn("w:ptab")
_W_TYPE = qn("w:type")
_W_PPR = qn("w:pPr")
_W_PSTYLE = qn("w:pStyle")
_W_VAL = qn("w:val")
_W_DRAWING = qn("w:drawing")
_A_BLIP = qn("a:blip")
_R_EMBED = qn("r:embed")


# --------------------------------------------------------------------------
# .docx
# --------------------------------------------------------------------------

def _build_heading_levels(document: Document) -> Dict[Optional[str], int]:
    """Resolve every PARAGRAPH style's id to a heading level (1-6) ONCE, up
    front — the fix for the profiled cost of `Paragraph.style`.

    Measured (cProfile, 100,000 plain paragraphs with no explicit style —
    so every access falls through to the "resolve the document's DEFAULT
    paragraph style" branch): `Paragraph.style` costs ~82.7s of the walk's
    85.1s total. Almost none of that is python-docx re-parsing anything —
    it is `CT_Styles.default_for` re-walking and re-type-converting EVERY
    `<w:style>` element in the styles part, from scratch, on every single
    paragraph (`_iter_styles` -> `style.type` -> an enum `from_xml` lookup,
    per style, per paragraph — 20M+ calls for 100k paragraphs against a
    template with ~200 styles). The `get_by_id` path taken when a paragraph
    DOES carry an explicit style id is cheaper per call (an lxml `xpath()`
    call rather than a full re-scan) but still re-resolves from scratch
    every time, and still shows up in the profile (~8s of the same 100k-call
    run). Either way, the shared root cause is the same: python-docx treats
    "what heading level is this paragraph" as free to recompute per access,
    when the document's styles part does not change during this walk.

    The fix: read each paragraph's raw style id straight off its own XML
    (`_paragraph_style_id` below: `<w:p>` -> `./w:pPr/w:pStyle/@w:val` via
    direct `find()`/`get()` calls, no XPath, no re-walking the styles part)
    and look it up in THIS dict, built by walking `document.styles` exactly
    ONCE per document. `document.styles` and each `style.name`/`style_id`/
    `type` are cheap per-style property reads (no XPath either) — see the
    profiling above, where building this map for the whole document cost
    ~1ms. Re-measured after this change: the same 100,000-paragraph walk
    (id lookup + dict.get, no `Paragraph.style` at all) dropped from ~85s
    to ~0.26s.

    `levels[None]` is set only when the document's own DEFAULT paragraph
    style (`Styles.default(WD_STYLE_TYPE.PARAGRAPH)` — what a `<w:p>` with
    no explicit `<w:pStyle>` actually renders as, per the WordprocessingML
    spec) happens to itself be a heading style — mirroring exactly what
    `Paragraph.style`'s own fallback did, just resolved once instead of
    per paragraph. That is vanishingly rare in practice (nobody sets
    "Heading 1" as their document's default body style) but costs nothing
    extra to get right.
    """
    levels: Dict[Optional[str], int] = {}
    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        level = _HEADING_LEVELS.get(style.name)
        if level is not None:
            levels[style.style_id] = level
    default_style = document.styles.default(WD_STYLE_TYPE.PARAGRAPH)
    if default_style is not None:
        default_level = levels.get(default_style.style_id)
        if default_level is not None:
            levels[None] = default_level
    return levels


def _run_text(r, image_rids: Optional[List[str]] = None) -> str:
    """Reimplements `docx.oxml.text.run.CT_R.text` (join each inner-content
    child's text equivalent) WITHOUT its `self.xpath("w:br | w:cr | ...")`
    call — see `_paragraph_text`'s docstring for why that call is expensive
    enough per-paragraph to matter and how this was verified equivalent.

    `image_rids`, when given, collects any embedded-image relationship ids
    found on a `<w:drawing>` child of this run (see `_drawing_blip_rids`) —
    folded into this SAME walk rather than a second pass over `r`'s
    children, because a second full walk (one for text, one for images)
    measurably reintroduces the class of per-paragraph cost this function
    exists to avoid: measured directly, adding a SEPARATE
    `_paragraph_image_rids`-style function (its own `iterchildren()` pass
    per paragraph/run, called after `_paragraph_block`) added ~1s to full
    conversion of the 300,000-plain-paragraph fixture (~1.6s -> ~2.5s on
    this machine) purely from walking every run's children twice — even
    though not one of those paragraphs has a `<w:drawing>` at all. Folding
    the check into THIS walk instead (one extra `elif` per child, evaluated
    only when every earlier branch already missed) measured back down to
    ~1.4-1.9s, i.e. no measurable regression against the pre-image
    baseline. `image_rids=None` (the default) skips that `elif`'s condition
    outright rather than allocating/threading a list no caller asked for."""
    parts = []
    for e in r.iterchildren():
        tag = e.tag
        if tag == _W_T:
            parts.append(e.text or "")
        elif tag == _W_TAB or tag == _W_PTAB:
            parts.append("\t")
        elif tag == _W_BR:
            # CT_Br.type defaults to "textWrapping" (python-docx's
            # OptionalAttribute default) when `w:type` is absent from the
            # XML — a bare `e.get(...)` would instead default to None,
            # which is NOT "textWrapping" and would silently drop every
            # ordinary line break (the overwhelmingly common case: Word
            # only ever writes an explicit w:type for a page/column break).
            br_type = e.get(_W_TYPE)
            parts.append("\n" if br_type is None or br_type == "textWrapping" else "")
        elif tag == _W_CR:
            parts.append("\n")
        elif tag == _W_NOBREAKHYPHEN:
            parts.append("-")
        elif image_rids is not None and tag == _W_DRAWING:
            image_rids.extend(_drawing_blip_rids(e))
    return "".join(parts)


def _paragraph_text(p, image_rids: Optional[List[str]] = None) -> str:
    """Reimplements `Paragraph.text` (join each direct `w:r`/`w:hyperlink`
    child's text) by walking the already-parsed lxml tree directly, instead
    of through python-docx's `CT_P.text`/`CT_R.text`/`CT_Hyperlink.text`
    properties — each of which calls `self.xpath(...)` internally, and
    (like `Paragraph.style` — see `_build_heading_levels`) pays lxml's
    XPath-expression overhead on EVERY call rather than once.

    This was the SECOND hot spot found by profiling full conversion after
    the style-lookup fix (see the module docstring's profiling numbers):
    at 100,000 trivial paragraphs, `CT_P.xpath`/`CT_R.xpath` together cost
    ~2.7s of a ~4.1s total — larger than every other cost in the walk
    combined, including this module's own code. Unlike the style fix, this
    cost can't be hoisted "once per document" (each paragraph's text is
    genuinely unique), so the fix here is a cheaper implementation of the
    SAME lookup, not caching it — a plain `iterchildren()` + tag-equality
    walk, with no xpath expression to compile or evaluate, replicating the
    exact semantics of `CT_R.text`/`CT_Hyperlink.text`/`CT_P.text`:
    `w:t` -> its own text (or "" if empty), `w:tab`/`w:ptab` -> "\\t",
    `w:cr` -> "\\n", `w:br` -> "\\n" unless it declares a non-default
    (page/column) `w:type`, `w:noBreakHyphen` -> "-", every other run
    inner-content element (drawings, field codes, ...) contributes nothing
    — matching `CT_R.text`'s own xpath, which only ever selects those six
    tag names. Verified equivalent to `Paragraph.text` output-for-output
    against this module's own docx fixtures plus dedicated line-break/tab/
    hyperlink cases (see test_doc_office.py).

    `image_rids`, when given, is threaded through to `_run_text` so a
    paragraph's embedded images are collected in the SAME pass as its text
    — see `_run_text`'s docstring for why that matters at scale.
    """
    parts = []
    for e in p.iterchildren():
        tag = e.tag
        if tag == _W_R:
            parts.append(_run_text(e, image_rids))
        elif tag == _W_HYPERLINK:
            for r in e.iterchildren():
                if r.tag == _W_R:
                    parts.append(_run_text(r, image_rids))
    return "".join(parts)


def _paragraph_style_id(p) -> Optional[str]:
    """Reimplements `docx.oxml.text.paragraph.CT_P.style` (`./w:pPr/w:pStyle
    /@w:val`, or None if either element is absent) via direct `find()`/`get()`
    calls against pre-resolved Clark-notation tag/attribute names, instead
    of python-docx's own `.pPr`/`.style` element-property descriptors —
    each of which re-resolves its tag name (`qn(...)`, a string-format-and-
    dict-lookup) on every call rather than once at import time, same class
    of avoidable per-call cost as `Paragraph.style`/`Paragraph.text` (see
    `_build_heading_levels`/`_paragraph_text`). This was the third and
    final hot spot the profile turned up: at 300,000 paragraphs, the
    `pPr`/`style` property chain still cost ~1s even after the style-LEVEL
    lookup was hoisted out of the per-paragraph path — this function's job
    is only to fetch the raw id, cheaply, for `heading_levels.get()` to
    look up (that dict is still what maps an id to a heading level)."""
    pPr = p.find(_W_PPR)
    if pPr is None:
        return None
    pStyle = pPr.find(_W_PSTYLE)
    if pStyle is None:
        return None
    return pStyle.get(_W_VAL)


def _paragraph_block(p, heading_levels: Dict[Optional[str], int],
                      image_rids: Optional[List[str]] = None) -> Optional[str]:
    """`image_rids`, when given, is passed straight through to
    `_paragraph_text` so the caller gets this paragraph's embedded-image
    relationship ids (in document order) as a side effect of the SAME text
    walk `_paragraph_block` already has to do — see `_run_text`'s docstring
    for why a second, separate walk over the same paragraph is avoided."""
    text = _paragraph_text(p, image_rids).strip()
    if not text:
        return None
    level = heading_levels.get(_paragraph_style_id(p))
    if level:
        return f"{'#' * level} {text}"
    return text


# --------------------------------------------------------------------------
# .docx embedded images — see the module docstring's "EMBEDDED IMAGES" note.
# --------------------------------------------------------------------------

def _drawing_blip_rids(drawing) -> List[str]:
    """Every `r:embed` relationship id inside one `<w:drawing>` element, in
    document order. `drawing.iter(_A_BLIP)` is a native lxml walk (no xpath
    expression to compile) over just this one drawing's own subtree — cheap
    and only ever invoked once a `<w:drawing>` has already been found, never
    per paragraph/run that doesn't have one. Almost every real `<w:drawing>`
    holds exactly one `<a:blip>` (one inline/anchored picture), but this
    doesn't assume that — a drawing wrapping a picture-fill shape or a
    multi-image group is walked the same way, and an `<a:blip>` with no
    `r:embed` attribute at all (a blip that references a blob directly
    rather than a relationship — rare, but valid OOXML) is simply skipped
    rather than yielding a bogus None id."""
    return [rid for rid in (b.get(_R_EMBED) for b in drawing.iter(_A_BLIP))
            if rid]


def _cell_image_rids(cell) -> List[str]:
    """Every embedded-image relationship id inside one table cell, found by
    walking the cell's own underlying `<w:tc>` element (`cell._tc` —
    python-docx's private-but-stable handle onto it; this module already
    reaches into an underlying lxml element the same way for
    `document.element.body`). Bounded by the same DOCX_MAX_TABLE_ROWS cap as
    the rest of `_table_block`'s per-row work — a row beyond that cap is
    never reached, so this never runs for it either."""
    return [rid for drawing in cell._tc.iter(_W_DRAWING)
            for rid in _drawing_blip_rids(drawing)]


class _ImageExtractor:
    """Per-document state for `.docx` embedded-image extraction, threaded
    through `_iter_body_blocks`/`_table_block`: the running extraction-order
    counter (also the file-name index — see doc_cache.office_image_path),
    an `r:embed` id -> (already-written Path, index) cache (see the module
    docstring's "DUPLICATE IMAGES" note for why a repeated rId must never be
    re-extracted or re-counted), and how many images were skipped once
    MAX_EXTRACTED_IMAGES was reached (reported in-band once, at the end of
    the document — see `_iter_body_blocks`).
    """

    def __init__(self, document: Document, adir: Path):
        self._related_parts = document.part.related_parts
        self._adir = adir
        self._by_rid: Dict[str, Tuple[Path, int]] = {}
        self.count = 0
        self.skipped = 0
        self.saved: List[Path] = []

    def extract(self, rid: str) -> Optional[Tuple[Path, int]]:
        """Returns (path, 1-based extraction index) for `rid`'s image, or
        None if it can't/shouldn't be extracted (an unresolvable rid, a
        related part that isn't actually an image, or — only for a
        genuinely NEW image — the document-wide ceiling already reached).
        A repeated rid always succeeds regardless of the ceiling: it costs
        no new extraction, just another link to a file already on disk."""
        cached = self._by_rid.get(rid)
        if cached is not None:
            return cached
        if self.count >= MAX_EXTRACTED_IMAGES:
            self.skipped += 1
            return None
        part = self._related_parts.get(rid)
        image = getattr(part, "image", None)
        if image is None:
            return None
        ext = image.ext or "png"
        if not ext.isalnum() or len(ext) > 10:
            # Defense in depth: `ext` is derived from the package's own part
            # name (see doc_cache.office_image_path's docstring) and is
            # never expected to contain anything but ordinary extension
            # characters — but this is still a value read from an untrusted
            # input file, and it's about to become part of a filesystem
            # path, so an unexpected shape (e.g. a stray "/") falls back to
            # a safe, fixed extension rather than being trusted verbatim.
            ext = "bin"
        self.count += 1
        out = doc_cache.office_image_path(self._adir, self.count, ext)
        out.write_bytes(image.blob)
        result = (out, self.count)
        self._by_rid[rid] = result
        self.saved.append(out)
        return result


# Reserved headroom subtracted from the budget handed to `_table_block`, so
# that a table's own truncation marker (appended after its internal budget
# check) can never itself push the block's total length past what the
# OUTER loop in `_convert_docx` is willing to accept — see `_table_block`.
_TABLE_MARKER_RESERVE = 300


def _table_block(table: Table, budget: int,
                  extractor: _ImageExtractor) -> Tuple[Optional[str], List[str]]:
    """Render `table` as a pipe table, bounded by BOTH DOCX_MAX_TABLE_ROWS
    (a structural per-table row cap) AND `budget` (the character budget
    still available in the document at the point this table was reached) —
    mirroring `_render_sheet_rows`'s dual cap for the same reason: a single
    pathological table must never cost O(its own size) before the
    document-level budget ever gets a chance to reject it wholesale.
    Measured before this cap existed: a single 20,000x6 table rendered
    whole, unconditionally, in ~4.4s / +52MB RSS — regardless of MAX_CHARS.

    Built incrementally with `format_cell`/`format_row` (not
    `table_to_markdown`, which needs the WHOLE table materialized up front
    to compute one shared column width) so this function can stop the
    instant either cap is hit, without ever having extracted text from a
    row beyond that point — `row.cells` is what actually walks a row's XML,
    so skipping it (not just discarding its output) is what keeps this
    bounded. Column width is taken from the table's own FIRST row and
    applied to every row after: a python-docx `Table` is rectangular by
    construction (row-level cell merging aside), so this is exact — unlike
    the analogous first-row-width guess this module makes for a
    dimension-less .xlsx sheet, which is a heuristic because a spreadsheet
    has no such guarantee.

    Returns (markdown, image_blocks): `image_blocks` is a list of Markdown
    image-link lines for every embedded image found in a KEPT row's cells
    (never a row beyond either cap — an unlisted row's images are no more
    reachable than its text), meant to be emitted by the caller AFTER the
    table's own markdown, never inside a cell — see the module docstring's
    "WHERE AN IMAGE LINK LANDS" note for why a cell's own pipe-table syntax
    is never touched.
    """
    lines: List[str] = []
    image_blocks: List[str] = []
    used = 0
    kept = 0
    width: Optional[int] = None
    row_cap_hit = False
    budget_hit = False
    for i, row in enumerate(table.rows):
        if i >= DOCX_MAX_TABLE_ROWS:
            row_cap_hit = True
            break
        row_cells = list(row.cells)
        cells = [format_cell(cell.text) for cell in row_cells]
        if width is None:
            width = len(cells)
        pieces = [format_row(cells, width)]
        if kept == 0:
            pieces.append("|" + "|".join([" --- "] * width) + "|")
        cost = sum(len(p) + 1 for p in pieces)
        if kept > 0 and used + cost > budget:
            budget_hit = True
            break
        lines.extend(pieces)
        used += cost
        kept += 1
        for col, cell in enumerate(row_cells, start=1):
            for rid in _cell_image_rids(cell):
                extracted = extractor.extract(rid)
                if extracted is not None:
                    path, idx = extracted
                    image_blocks.append(
                        f"![embedded image {idx} — table row {i + 1}, "
                        f"column {col}]({path})")

    if not lines:
        return None, image_blocks

    md = "\n".join(lines)
    if row_cap_hit:
        md += "\n\n" + truncation_marker(
            f"after {kept} rows — this table "
            f"exceeds the {DOCX_MAX_TABLE_ROWS}-row-per-table limit")
    elif budget_hit:
        md += "\n\n" + truncation_marker(
            f"after {kept} rows — the character "
            "limit for this document was reached")
    return md, image_blocks


def _iter_body_blocks(document: Document, remaining_budget,
                       heading_levels: Dict[Optional[str], int],
                       extractor: _ImageExtractor) -> Iterator[str]:
    """Yield blocks from `document.element.body`'s direct children IN ORDER,
    LAZILY — handling each `<w:p>` via `_paragraph_block` (which also
    collects the paragraph's embedded images, in the SAME text walk — see
    `_run_text`'s docstring) and mapping each `<w:tbl>` to a python-docx
    `Table` — see the module docstring for why walking the body directly
    (and not `document.paragraphs`/`document.tables`) is the only way to
    keep the source's true interleaving. Anything else under `<w:body>`
    (e.g. the trailing `<w:sectPr>` section-properties element) is silently
    skipped — it carries no displayable content.

    A `<w:p>`'s own embedded images are yielded as separate blocks
    immediately AFTER its text block (or on their own, if the paragraph has
    no text — the shape `Document.add_picture` produces) — this is what
    puts an inline image at its position in the document, between whatever
    came before and after it, the same interleaving `_assemble_page` gives a
    PDF page's images (see the module docstring's "WHERE AN IMAGE LINK
    LANDS" note).

    A `<w:p>` is passed to `_paragraph_block` as its raw lxml element, NOT
    wrapped in a python-docx `Paragraph` — nothing here needs the wrapper
    any more (`_paragraph_block` reads the style id, the text, AND any
    embedded images straight off the XML; see `_build_heading_levels` and
    `_paragraph_text` for why going through `Paragraph.style`/
    `Paragraph.text` was the actual cost). `Table`/`Row`/`Cell` ARE still
    used for `<w:tbl>`, since only the heading/text paths were profiled as
    hot — see `_table_block`.

    This is a generator, not a list, on purpose: `_table_block` (real,
    proportional-to-content work — walking every kept row's cells) is the
    one genuinely expensive step left in this module that scales with
    DOCUMENT size rather than KEPT size — `_convert_docx` must still be
    able to stop asking this generator for more the moment its character
    budget is spent, without this function having formatted a single block
    beyond that point. `document.element.body.iterchildren()` is itself
    already a lazy walk over the (already-parsed-in-memory) XML tree, so
    nothing upstream of this loop is re-done by stopping early.

    `remaining_budget` is a zero-arg callable read fresh each time a
    `<w:tbl>` is reached (not a snapshot taken once up front), so a table's
    own internal cap always sees how much of MAX_CHARS is actually left at
    that point in the walk, not how much was left when iteration started.

    `heading_levels` is `_build_heading_levels(document)`, computed ONCE by
    the caller and threaded through rather than rebuilt here — this
    function runs once per `<w:p>`, so rebuilding it here would reintroduce
    exactly the "per paragraph, not per document" cost this fix removes.

    A final "images not extracted" note (if `extractor.skipped`) is yielded
    only AFTER the whole body has been walked — see the module docstring's
    "MAX_EXTRACTED_IMAGES" note for why the skipped count can only be known
    at that point, unlike doc_pdf's per-page lookahead.
    """
    for child in document.element.body.iterchildren():
        if child.tag == _W_P:
            image_rids: List[str] = []
            block = _paragraph_block(child, heading_levels, image_rids)
            if block:
                yield block
            for rid in image_rids:
                extracted = extractor.extract(rid)
                if extracted is not None:
                    path, idx = extracted
                    yield f"![embedded image {idx}]({path})"
        elif child.tag == _W_TBL:
            budget = max(0, remaining_budget() - _TABLE_MARKER_RESERVE)
            table_md, table_images = _table_block(
                Table(child, document), budget, extractor)
            if table_md:
                yield table_md
            for image_block in table_images:
                yield image_block
        else:
            continue
    if extractor.skipped:
        yield (f"*(this document also contains {extractor.skipped} further "
               f"embedded image{'s' if extractor.skipped != 1 else ''} not "
               f"extracted — the document-wide limit of "
               f"{MAX_EXTRACTED_IMAGES} images was reached.)*")


def _convert_docx(source: Path, adir: Path) -> Tuple[str, List[Path]]:
    document = _open_docx(source)
    heading_levels = _build_heading_levels(document)
    extractor = _ImageExtractor(document, adir)

    lines: List[str] = []
    used = 0
    kept = 0
    truncated = False

    def remaining_budget() -> int:
        return MAX_CHARS - used

    for block in _iter_body_blocks(document, remaining_budget, heading_levels, extractor):
        cost = len(block) + (2 if kept > 0 else 0)  # blank line before it
        if kept > 0 and used + cost > MAX_CHARS:
            truncated = True
            break
        if kept > 0:
            lines.append("")
        lines.append(block)
        used += cost
        kept += 1

    if kept == 0:
        return "*(this document has no readable text or tables)*", extractor.saved

    md = "\n".join(lines)
    if truncated:
        # The true total section count is deliberately NOT reported here:
        # knowing it would require walking the rest of the document after
        # all, which is exactly the O(document size) cost this function
        # exists to avoid (see the module docstring).
        md += "\n\n" + truncation_marker(
            f"after {kept} section"
            f"{'s' if kept != 1 else ''} — the {MAX_CHARS}-character "
            "limit was reached; this document has no page range to "
            "resume from, so ask for a narrower excerpt if more is "
            "needed")
    return md, extractor.saved


# --------------------------------------------------------------------------
# .xlsx
# --------------------------------------------------------------------------

def _render_sheet_rows(ws, budget: int, source: Path) -> Tuple[str, int, Optional[int], bool, bool]:
    """Render up to XLSX_MAX_ROWS x XLSX_MAX_COLS of `ws` as one pipe table
    (first row as the header), stopping early if the running character
    `budget` is exhausted first. Returns (markdown, kept_rows, total_rows,
    hit_row_cap, hit_budget).

    `hit_row_cap`/`hit_budget` say explicitly WHY iteration stopped short
    (mutually exclusive by construction: the loop below checks the
    structural cap before ever pricing a row against `budget`, so at most
    one of them is ever True for a single call) — `_sheet_truncation_note`
    used to reconstruct this by comparing `kept` against a returned
    "capped_rows" value that meant two different things depending on
    whether the sheet's dimension was known (min(known_rows,
    XLSX_MAX_ROWS) when known, but just the bare XLSX_MAX_ROWS fallback
    ceiling when not — NOT the same thing as "how many rows this sheet
    actually has"). That mismatch was a real bug: a tiny write_only sheet
    (1 real row, no <dimension>) was reported as "truncated after 1 of
    5000 rows" even though nothing was cut, because `kept < capped_rows`
    (1 < 5000) looked exactly like a budget cut. Returning the actual
    reason directly removes the need to infer it at all.

    `ws.max_row`/`ws.max_column` are `None` (not `0`) when the sheet's XML
    has no `<dimension>` element at all — e.g. every sheet written by
    `openpyxl.Workbook(write_only=True)`. This function NEVER treats that
    metadata as proof of anything: only actually iterating and observing a
    row (`saw_any` below) proves the sheet has content. A sheet that DOES
    declare a dimension but has zero real cells (openpyxl's default
    `max_row == max_col == 1` for a brand new empty sheet) still yields no
    rows when iterated — `saw_any` is what tells that case apart from "one
    row of real, if blank, content" too.

    `total_rows` in the return value is `None` when the sheet's true row
    count could not be determined without a full scan (dimension unknown,
    and iteration was cut short by either the structural row cap or the
    character budget before reaching the sheet's actual end) — see the
    module docstring for why this function deliberately never calls
    `ws.calculate_dimension(force=True)` to resolve that. Callers must
    handle `total_rows is None` rather than assume an int.
    """
    known_rows = ws.max_row
    known_cols = ws.max_column
    # Fall back to the row structural cap itself as the iteration bound when
    # the dimension is unknown (or, degenerately, reported as 0) — this
    # keeps the read_only parser's own early-stop (it never reads past
    # max_row) doing the bounding, instead of this function ever attempting
    # to iterate an unbounded sheet.
    row_cap = min(known_rows, XLSX_MAX_ROWS) if known_rows else XLSX_MAX_ROWS
    if known_cols:
        capped_cols = min(known_cols, XLSX_MAX_COLS)
    else:
        # Column count unknown too. Forcing every row to XLSX_MAX_COLS wide
        # (16,384 — Excel's own column maximum, "XFD") here would be its own
        # budget-exhausting bug: every row, including ones with two real
        # cells, would be padded out to 16,384 pipe-table columns, so the
        # FIRST such row alone could burn a large chunk of the character
        # ceiling and starve every row after it — that's not hypothetical,
        # it's what happened (against the much smaller original budget)
        # before this fallback existed.
        # Peeking at row 1's own actual width is a light, bounded probe (one
        # single-row read, not a full-sheet scan) and is representative for
        # the common case this guards against (a write_only/no-<dimension>
        # workbook whose rows are written with a consistent shape, e.g. via
        # repeated `ws.append([...])` calls). A later row genuinely wider
        # than row 1 will have its extra columns dropped, same as any sheet
        # whose real width exceeds XLSX_MAX_COLS today.
        first_row = next(_safe_row_iter(
            ws.iter_rows(min_row=1, max_row=1, values_only=True), source), None)
        capped_cols = min(len(first_row), XLSX_MAX_COLS) if first_row else XLSX_MAX_COLS

    lines: List[str] = []
    used = 0
    kept = 0
    saw_any = False
    idx = 0
    stopped_on_budget = False
    # Ask for one row PAST row_cap purely to detect "does more content exist
    # beyond the structural cap" — cheap, because the read_only parser stops
    # as soon as it sees a row index past this bound, so this never costs a
    # full-sheet scan even when the dimension is unknown.
    #
    # Wrapped through _safe_row_iter, not iterated directly: read_only mode
    # parses a sheet's own XML lazily, one row at a time, so a truncated/
    # malformed sheet1.xml raises HERE, mid-iteration — never at
    # load_workbook() time (see _open_xlsx/_safe_row_iter's docstrings).
    for idx, row in enumerate(_safe_row_iter(
            ws.iter_rows(min_row=1, max_row=row_cap + 1,
                         min_col=1, max_col=capped_cols,
                         values_only=True), source), start=1):
        saw_any = True
        if idx > row_cap:
            break
        cells = [format_cell("" if v is None else str(v)) for v in row]
        pieces = [format_row(cells, capped_cols)]
        if kept == 0:
            pieces.append("|" + "|".join([" --- "] * capped_cols) + "|")
        cost = sum(len(p) + 1 for p in pieces)
        if kept > 0 and used + cost > budget:
            stopped_on_budget = True
            break
        lines.extend(pieces)
        used += cost
        kept += 1

    if not saw_any:
        return "", 0, 0, False, False

    if known_rows is not None:
        total_rows: Optional[int] = known_rows
    elif stopped_on_budget:
        total_rows = None  # cut short by the budget; true extent still unknown
    elif idx > row_cap:
        total_rows = None  # more than row_cap rows exist; exact count unknown
    else:
        total_rows = idx  # generator ran to completion at/under the cap — that IS the true count

    hit_row_cap = idx > row_cap
    return "\n".join(lines), kept, total_rows, hit_row_cap, stopped_on_budget


def _sheet_truncation_note(kept: int, total_rows: Optional[int],
                            hit_row_cap: bool, hit_budget: bool) -> Optional[str]:
    """Returns the marker's DETAIL text only — no "truncated" prefix, no
    "*(...)*" wrapping — so _convert_xlsx can wrap it through
    doc_support.truncation_marker without duplicating the shared sentinel
    word by hand. None means nothing was truncated at all.

    `hit_row_cap`/`hit_budget` come straight from `_render_sheet_rows` (see
    its docstring) rather than being re-derived from `kept` vs. a "capped
    rows" count here — that reconstruction used to be wrong for a
    dimension-less sheet smaller than XLSX_MAX_ROWS (every real row
    rendered, nothing cut) because the value standing in for "this sheet's
    expected size" was actually just the structural cap's fallback
    ceiling, not this sheet's real size. `XLSX_MAX_ROWS` is read directly
    from the module constant rather than threaded through as a parameter,
    since it is the same fixed number in every case where `hit_row_cap` is
    True (see `_render_sheet_rows`'s docstring: reaching the row cap only
    ever means XLSX_MAX_ROWS was the binding limit)."""
    if not hit_row_cap and not hit_budget:
        return None  # nothing was cut short by either limit
    if hit_budget:
        if total_rows is not None and total_rows > XLSX_MAX_ROWS:
            # both limits are in play — say so, rather than reporting the
            # row-per-sheet cap as if it were this sheet's real size (that
            # would hide the true row count behind the structural cap).
            return (f"after {kept} of {total_rows} rows — the "
                    "character limit for this document was reached first; "
                    f"this sheet also exceeds the {XLSX_MAX_ROWS}-row-per-"
                    "sheet limit")
        if total_rows is None:
            return (f"after {kept} rows shown — the character "
                    "limit for this document was reached; this sheet's "
                    "total row count could not be determined (its XML has "
                    "no declared dimension)")
        return (f"after {kept} of {total_rows} rows shown — the "
                "character limit for this document was reached")
    # hit_row_cap and not hit_budget: the structural cap is always what was
    # actually hit here (XLSX_MAX_ROWS — see _render_sheet_rows's docstring).
    if total_rows is None:
        return (f"after {XLSX_MAX_ROWS} rows — this sheet exceeds "
                f"the {XLSX_MAX_ROWS}-row-per-sheet limit (its exact total "
                "is unknown: this sheet's XML has no declared dimension)")
    return (f"after {XLSX_MAX_ROWS} of {total_rows} rows — this "
            f"sheet exceeds the {XLSX_MAX_ROWS}-row-per-sheet limit")


def _convert_xlsx(source: Path) -> str:
    # read_only mode streams rows lazily straight out of the underlying zip
    # archive — the workbook (and archive) must stay open for the ENTIRE
    # sheet-by-sheet render loop below, not just long enough to list
    # `worksheets`, or the first `iter_rows` call on any sheet raises
    # "Attempt to use ZIP archive that was already closed".
    wb = _open_xlsx(source)
    try:
        worksheets = list(wb.worksheets)
        if not worksheets:
            return "*(this workbook has no sheets)*"

        n = len(worksheets)
        lines: List[str] = []
        used = 0
        for i, ws in enumerate(worksheets):
            # The heading is unconditional — see the module docstring's
            # "sheet reachability" note. No early exit here, unlike every
            # other budget check in this module: a document with no earlier
            # boundary to stop at is exactly the shape we cannot afford for
            # a MULTI-sheet workbook, since a missing heading isn't just
            # truncated content — it's a sheet the agent can never even
            # discover exists.
            header = f"## {ws.title}"
            header_cost = len(header) + (2 if lines else 0)
            if lines:
                lines.append("")
            lines.append(header)
            used += header_cost

            # Fair-share whatever budget remains across this sheet and every
            # sheet still to come (recomputed fresh each iteration, not a
            # fixed 1/n split decided up front) — a sheet that needs less
            # than its share leaves the surplus for the sheets after it. This
            # is what keeps a short "## Summary"/"## Notes" sheet from being
            # squeezed by an even split sized for a much bigger sheet.
            remaining_sheets = n - i
            sheet_budget = max(0, (MAX_CHARS - used) // remaining_sheets)

            table_md, kept, total_rows, hit_row_cap, hit_budget = _render_sheet_rows(
                ws, sheet_budget, source)
            if table_md:
                lines.append("")
                lines.append(table_md)
                used += len(table_md) + 2

            note = _sheet_truncation_note(kept, total_rows, hit_row_cap, hit_budget)
            if note:
                lines.append("")
                marker = truncation_marker(
                    f"{note} — this file has no page range to resume "
                    "from, so ask for a narrower export if the rest is "
                    "needed")
                lines.append(marker)
                used += len(marker) + 2
                # Deliberately NOT a `break`: every remaining sheet still
                # gets its own heading (and, via _render_sheet_rows's
                # first-row-always-shown rule, at least one row) even though
                # this sheet's own share is spent — see the module
                # docstring. The per-sheet fair-share above already prices
                # that in for the sheets still to come.
    finally:
        wb.close()

    return "\n".join(lines)


def convert(source: Path, adir: Path,
            pages: Optional[List[int]] = None, **_ignored) -> Tuple[str, List[Path]]:
    """`.docx` returns (markdown, extracted_image_paths) — see the module
    docstring's "EMBEDDED IMAGES" note: every embedded raster image is
    extracted to a sibling file under `adir` and linked inline at its
    position. `.xlsx` returns (markdown, []) unchanged — see the module
    docstring's "XLSX IMAGES" note for why that format's embedded images are
    deliberately out of scope for now.

    `pages` is accepted for signature parity with every other converter
    (docs.extract calls all of them uniformly) but unused: neither format
    implements `page_count`/`render`, so docs.py never lets an explicit
    `pages` argument reach here.
    """
    ext = source.suffix.lower()
    if ext == ".docx":
        return _convert_docx(source, adir)
    if ext == ".xlsx":
        return _convert_xlsx(source), []
    raise KeyError(ext)
