import re
import time
import zipfile

import pytest

from es.capabilities import doc_office


def test_docx_headings_become_markdown_headings(docx_file, tmp_path):
    md, images = doc_office.convert(docx_file, tmp_path)
    assert "# Season Overview" in md
    assert "## Fees" in md
    assert "Practices are Tuesdays" in md
    assert images == []


def test_docx_tables_render_as_pipe_tables(docx_file, tmp_path):
    md, _ = doc_office.convert(docx_file, tmp_path)
    assert "| Item | Cost |" in md
    assert "| Kit | $65 |" in md


# --------------------------------------------------------------------------
# Embedded images — every image comes out as the thing it is (a file, linked
# inline at its position), because it exists, not because it scored above a
# threshold. Mirrors doc_pdf's own image-extraction test shape.
# --------------------------------------------------------------------------

def test_docx_inline_image_extracted_and_linked_between_paragraphs(
        docx_with_inline_image, tmp_path):
    md, images = doc_office.convert(docx_with_inline_image, tmp_path)

    assert len(images) == 1
    assert images[0].exists()
    assert images[0].suffix == ".png"

    first_idx = md.index("FIRST paragraph")
    image_idx = md.index("![embedded image 1]")
    last_idx = md.index("LAST paragraph")
    assert first_idx < image_idx < last_idx


def test_docx_no_images_means_no_files_and_no_links(docx_file, tmp_path):
    """docx_file (the base fixture) has headings/paragraphs/a table but no
    embedded pictures at all — must not fabricate a file or a link."""
    md, images = doc_office.convert(docx_file, tmp_path)
    assert images == []
    assert "![" not in md


def test_docx_two_images_two_files_two_links_in_document_order(
        docx_with_two_images, tmp_path):
    md, images = doc_office.convert(docx_with_two_images, tmp_path)

    assert len(images) == 2
    assert images[0] != images[1]
    for img in images:
        assert img.exists()

    first_link_idx = md.index("![embedded image 1]")
    between_idx = md.index("BETWEEN THE TWO PHOTOS")
    second_link_idx = md.index("![embedded image 2]")
    assert first_link_idx < between_idx < second_link_idx


def test_docx_table_cell_image_emitted_after_table_not_inside_cell(
        docx_with_table_cell_image, tmp_path):
    """A Markdown link inside a pipe-table cell risks corrupting the table's
    own syntax (see the module docstring's "WHERE AN IMAGE LINK LANDS"
    note) — so a cell's image is reported as its own block AFTER the whole
    table, naming which row/column it came from, rather than embedded in
    the cell itself."""
    md, images = doc_office.convert(docx_with_table_cell_image, tmp_path)

    assert len(images) == 1
    assert images[0].exists()

    # The cell itself stays plain text — no link, no broken pipe syntax.
    table_line = next(line for line in md.splitlines()
                       if line.startswith("| Widget"))
    assert "![" not in table_line
    assert table_line.count("|") == 3  # "| Widget |  |" — unbroken

    # The image is reported after the table, naming its cell.
    table_idx = md.index("| Widget")
    image_idx = md.index("![embedded image 1")
    assert table_idx < image_idx
    assert "row 2, column 2" in md


def test_docx_duplicate_image_relationship_writes_one_file_two_links(
        docx_with_duplicated_image_relationship, tmp_path):
    """The SAME `r:embed` relationship id referenced from two paragraphs
    (e.g. a letterhead logo at the top and bottom of a template) is the SAME
    underlying image part by construction — python-docx's own model already
    asserts this, it isn't an ambiguous "same bytes, is that one image or
    two" judgment call. Writing it to disk twice would be pure waste with
    no benefit to the agent (same picture either way); each of the two
    APPEARANCES in the reading order still gets its own link, both pointing
    at the one file that was actually written."""
    md, images = doc_office.convert(
        docx_with_duplicated_image_relationship, tmp_path)

    assert len(images) == 1
    assert images[0].exists()
    assert md.count("![embedded image 1]") == 2

    para_a_idx = md.index("Para A")
    first_link_idx = md.index("![embedded image 1]")
    para_b_idx = md.index("Para B")
    second_link_idx = md.rindex("![embedded image 1]")
    assert para_a_idx < first_link_idx < para_b_idx < second_link_idx


def test_docx_image_extraction_ceiling_is_enforced_and_reported_in_band(
        tmp_path, monkeypatch):
    """A pathological document (far more embedded images than any real one
    would carry) must not write an unbounded number of files — and hitting
    the ceiling must be reported IN-BAND, never a silent drop. Monkeypatches
    the ceiling down (mirroring doc_pdf's own ceiling test) rather than
    building 500+ real images, which would make this test both slow and an
    unfaithful stand-in for the real limit's *shape*, not its exact value."""
    import io
    from docx import Document
    from PIL import Image

    monkeypatch.setattr(doc_office, "MAX_EXTRACTED_IMAGES", 3)

    p = tmp_path / "many_images.docx"
    d = Document()
    for i in range(5):
        photo = io.BytesIO()
        Image.new("RGB", (10, 10), (i, i, i)).save(photo, format="PNG")
        photo.seek(0)
        d.add_paragraph(f"before photo {i}")
        d.add_picture(photo)
    d.save(str(p))

    md, images = doc_office.convert(p, tmp_path)

    assert len(images) == 3  # only the ceiling's worth actually written
    assert "![embedded image 1]" in md
    assert "![embedded image 3]" in md
    assert "![embedded image 4]" not in md
    assert "2 further" in md
    assert "not extracted" in md
    assert "limit of 3 images" in md


def test_docx_preserves_document_order(tmp_path):
    """python-docx exposes paragraphs and tables as SEPARATE lists that do not
    preserve their interleaved order. A table between two paragraphs must come
    out between them, not appended at the end."""
    from docx import Document
    p = tmp_path / "ordered.docx"
    d = Document()
    d.add_paragraph("FIRST paragraph")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "MIDDLE cell"
    d.add_paragraph("LAST paragraph")
    d.save(str(p))
    md, _ = doc_office.convert(p, tmp_path)
    assert md.index("FIRST") < md.index("MIDDLE") < md.index("LAST")


def test_xlsx_each_sheet_is_a_section(xlsx_file, tmp_path):
    """Sheets are the natural heading unit — this is what makes a workbook
    pageable by the reader."""
    md, _ = doc_office.convert(xlsx_file, tmp_path)
    assert "## Roster" in md
    assert "## Fees" in md


def test_xlsx_rows_render_as_a_table(xlsx_file, tmp_path):
    md, _ = doc_office.convert(xlsx_file, tmp_path)
    assert "| Name | Number |" in md
    assert "| Alice | 9 |" in md


def test_xlsx_large_sheet_converts_fully(tmp_path):
    """Was test_xlsx_large_sheet_truncates_with_a_marker: under the old
    30,000-character MAX_CHARS, an 8000-row sheet was guaranteed to
    truncate. MAX_CHARS is now a generous resource ceiling (bound
    resources, not context — see the module docstring), so a sheet this
    size converts in full; the truncation MECHANISM itself is covered
    separately (test_xlsx_row_cap_truncates_before_the_character_budget,
    and the monkeypatched-MAX_CHARS tests below)."""
    from openpyxl import Workbook
    p = tmp_path / "big.xlsx"
    wb = Workbook(); ws = wb.active
    for i in range(8000):
        ws.append([f"row{i}", i, "padding value here"])
    wb.save(str(p))
    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated" not in md.lower()
    assert "| row0 | 0 | padding value here |" in md
    assert "| row7999 | 7999 | padding value here |" in md


def test_xlsx_later_sheets_reachable_when_first_sheet_exhausts_budget(tmp_path, monkeypatch):
    """Regression: before the per-sheet budget fix, a first sheet alone big
    enough to spend the whole MAX_CHARS budget made every LATER sheet vanish
    entirely -- not truncated, not noted, simply absent from both the
    Markdown and (since es_read's outline is built from `## ` headings) the
    outline itself. No parameter on es_doc_extract/es_read could reach a
    "## Summary"/"## Notes" sheet that was never emitted in the first place.
    Every sheet must now get at least a heading, and (per
    _render_sheet_rows's own first-row-always-shown rule) at least one row
    of real, independently-readable content.

    MAX_CHARS is now a generous resource ceiling (see the module
    docstring), so an 8000-row first sheet no longer exhausts it by
    default -- the fair-share machinery this test exists to cover would
    never actually engage without deliberately constraining the ceiling
    back down, which is exactly what it protects: the BACKSTOP for the
    rare case the ceiling genuinely is hit, not the common path. Monkeypatch
    MAX_CHARS back to its old value to reproduce that constrained scenario
    on a workbook of ordinary test size, rather than building a workbook
    large enough to exhaust the real default ceiling.

    Summary/Notes are deliberately MULTI-row (60 each), not the original
    single-row sheets: with only one row apiece, the separate "always show
    the first row regardless of budget" rule alone satisfies every
    assertion below even with fair-share removed entirely (that row IS the
    header, always shown) -- a prior task flagged this as a mutation the
    suite couldn't see. Asserting a row deep into each sheet (40 of 60) is
    only true when the fair-share split actually hands each later sheet a
    real slice of the remaining budget (~9-10k characters here, verified
    empirically) rather than the ~0 characters it would get if the first
    sheet's own share were computed as "everything left" instead of "an
    even split with the sheets still to come"."""
    monkeypatch.setattr(doc_office, "MAX_CHARS", 30_000)
    from openpyxl import Workbook
    from es.capabilities import read as read_cap

    p = tmp_path / "report.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for i in range(8000):
        ws.append([f"row{i}", i, "padding value here to burn the budget"])
    summary = wb.create_sheet("Summary")
    for i in range(60):
        summary.append([f"Metric {i}", i])
    notes = wb.create_sheet("Notes")
    for i in range(60):
        notes.append([f"Note line {i}: remember to reconcile next quarter."])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)

    # All three sheets are discoverable from the outline built over this one
    # markdown string -- exactly what es_read hands the agent on its FIRST,
    # argument-less call.
    outline = read_cap.outline(md)
    assert [s["title"] for s in outline] == ["Data", "Summary", "Notes"]

    # And each later sheet is independently readable well past its own first
    # row -- proving the fair-share budget, not just the "first row always
    # shown" rule, is what's doing the work here.
    summary_id = next(s["id"] for s in outline if s["title"] == "Summary")
    notes_id = next(s["id"] for s in outline if s["title"] == "Notes")
    summary_section = read_cap.section(md, summary_id)
    notes_section = read_cap.section(md, notes_id)
    assert "Metric 0" in summary_section
    assert "Metric 40" in summary_section
    assert "Note line 0" in notes_section
    assert "Note line 40" in notes_section


def test_xlsx_empty_workbook_does_not_raise(tmp_path):
    from openpyxl import Workbook
    p = tmp_path / "empty.xlsx"
    Workbook().save(str(p))
    md, _ = doc_office.convert(p, tmp_path)
    assert isinstance(md, str)


def test_docx_empty_document_does_not_raise(tmp_path):
    from docx import Document
    p = tmp_path / "empty.docx"
    Document().save(str(p))
    md, _ = doc_office.convert(p, tmp_path)
    assert isinstance(md, str)


def test_corrupt_office_file_raises(tmp_path):
    p = tmp_path / "bad.docx"
    p.write_bytes(b"not a zip at all")
    with pytest.raises(Exception):
        doc_office.convert(p, tmp_path)


def test_xlsx_formula_cells_render_their_value_or_formula_predictably(tmp_path):
    """openpyxl returns either the formula string or the cached value depending
    on data_only. Whichever you choose, be deliberate and consistent."""
    from openpyxl import Workbook
    p = tmp_path / "calc.xlsx"
    wb = Workbook(); ws = wb.active
    ws.append(["a", 2]); ws.append(["b", 3]); ws["B3"] = "=SUM(B1:B2)"
    wb.save(str(p))
    md, _ = doc_office.convert(p, tmp_path)
    assert isinstance(md, str) and "|" in md


# --------------------------------------------------------------------------
# Regression: verified DATA-LOSS bug — a write_only (or otherwise
# <dimension>-less) .xlsx sheet used to report `ws.max_row or 0` == 0 and
# get silently treated as "empty", discarding every row with no error and
# no truncation marker. See the doc_office module docstring for the full
# writeup and the measured ws.calculate_dimension(force=True) cost that
# ruled out the naive fix.
# --------------------------------------------------------------------------

def test_xlsx_write_only_workbook_rows_are_not_lost(tmp_path):
    """openpyxl.Workbook(write_only=True) never writes a <dimension> element,
    so ws.max_row/ws.max_column are None, not 0 — a mainstream way for tools
    (pandas/ETL exports included) to produce large spreadsheets, not an
    exotic edge case."""
    from openpyxl import Workbook
    p = tmp_path / "write_only.xlsx"
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Data")
    for i in range(50):
        ws.append([f"row{i}", i])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "## Data" in md
    assert "| row0 | 0 |" in md
    assert "| row49 | 49 |" in md
    # must not be mistaken for the genuinely-empty-sheet case
    assert "no data" not in md.lower()
    # nor mistaken for a truncated one — see
    # test_xlsx_small_write_only_sheet_is_not_falsely_reported_truncated for
    # the regression this guards (a fallback ceiling used in place of the
    # sheet's unknown true size used to look exactly like a budget cut).
    assert "truncated" not in md.lower()


def test_xlsx_small_write_only_sheet_is_not_falsely_reported_truncated(tmp_path):
    """Regression guard: `_sheet_truncation_note` used to infer "was this
    sheet cut short?" by comparing `kept` against a returned "capped_rows"
    value that meant "min(real total, XLSX_MAX_ROWS)" when the sheet's
    dimension was known, but just the bare XLSX_MAX_ROWS fallback ceiling
    (a safety bound for iteration, not a real count) when it was not — e.g.
    every `openpyxl.Workbook(write_only=True)` sheet. A tiny write_only
    sheet (2 real rows, well under both the character budget and
    XLSX_MAX_ROWS) satisfied `kept < capped_rows` (2 < 5000) purely because
    of that fallback value, and was reported as "truncated after 2 of 5000
    rows" even though nothing was cut at all. `_render_sheet_rows` now
    returns the actual reason (`hit_row_cap`/`hit_budget`) directly instead
    of leaving the caller to reconstruct it from an overloaded count."""
    from openpyxl import Workbook

    p = tmp_path / "tiny_write_only.xlsx"
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Data")
    ws.append(["a", "b"])
    ws.append(["c", "d"])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "| a | b |" in md
    assert "| c | d |" in md
    assert "truncated" not in md.lower()


def test_xlsx_dimension_stripped_sheet_rows_are_not_lost(tmp_path):
    """A sheet whose XML never had (or lost) its <dimension> element must
    render identically to one that has it. Build a normal workbook, then
    strip the <dimension .../> element directly from the zip member — this
    reproduces the bug report's exact repro without depending on how
    openpyxl happens to write write_only files today."""
    from openpyxl import Workbook
    p = tmp_path / "normal.xlsx"
    wb = Workbook(); ws = wb.active; ws.title = "Sheet1"
    for i in range(50):
        ws.append([f"r{i}", i])
    wb.save(str(p))

    with zipfile.ZipFile(str(p), "r") as zin:
        members = {n: zin.read(n) for n in zin.namelist()}
    sheet_name = next(n for n in members if n.startswith("xl/worksheets/sheet"))
    xml = members[sheet_name].decode("utf-8")
    assert "<dimension" in xml  # sanity: openpyxl did write one to strip
    members[sheet_name] = re.sub(r"<dimension[^/]*/>", "", xml).encode("utf-8")

    stripped = tmp_path / "stripped.xlsx"
    with zipfile.ZipFile(str(stripped), "w") as zout:
        for name, content in members.items():
            zout.writestr(name, content)

    md, _ = doc_office.convert(stripped, tmp_path)
    assert "| r0 | 0 |" in md
    assert "| r49 | 49 |" in md


def test_xlsx_large_write_only_sheet_converts_fully_and_quickly(tmp_path):
    """The bug report's own repro: a 300,000-row write_only workbook. Was
    test_..._converts_quickly_and_truncates: under the old 30,000-character
    MAX_CHARS this necessarily truncated almost immediately. MAX_CHARS is
    now a generous resource ceiling (bound resources, not context), so this
    now converts EVERY row, and must still do so quickly — bounded by
    reading/formatting cost, not by the sheet's size triggering something
    slower (e.g. a full-dimension scan)."""
    from openpyxl import Workbook
    p = tmp_path / "big_write_only.xlsx"
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Sheet")
    for i in range(300_000):
        ws.append([f"row{i}", i, "padding value here"])
    wb.save(str(p))

    t0 = time.time()
    md, _ = doc_office.convert(p, tmp_path)
    elapsed = time.time() - t0

    assert "| row0 |" in md
    assert "| row299999 |" in md
    assert "truncated" not in md.lower()
    # Not a tight perf bound: opening/parsing a 300,000-row read_only
    # workbook (shared strings etc.) plus formatting every one of those rows
    # as real content (unlike the .docx side, there is no per-access lookup
    # bug here to fix — this is genuine O(sheet size) work) costs several
    # seconds regardless of this module's own logic. Measured on this
    # machine: ~8s for this exact shape (300,000 rows x 3 real columns).
    # See test_docx_huge_paragraph_count_converts_quickly_and_bounded for
    # the tight sub-second-per-100k bound on the .docx side, where the fix
    # really did remove an O(document size) cost rather than just widening
    # the budget that used to hide it.
    assert elapsed < 30.0, f"conversion took {elapsed:.2f}s — should not scale with sheet size"


# --------------------------------------------------------------------------
# Item 3: truncation must report the sheet's TRUE row count, not the
# min(total_rows, XLSX_MAX_ROWS) structural cap masquerading as the total.
# --------------------------------------------------------------------------

def test_xlsx_truncation_note_reports_true_total_not_the_structural_cap(tmp_path, monkeypatch):
    """MAX_CHARS is now a generous resource ceiling that an ordinary 8000-row
    sheet never reaches on its own (see test_xlsx_large_sheet_converts_fully)
    — monkeypatch it back down so the CHARACTER budget, not the (now much
    larger, XLSX_MAX_ROWS = Excel's own 1,048,576-row maximum) structural
    cap, is what actually truncates this sheet, exercising the same
    known-dimension-but-budget-cut path the original regression covered."""
    monkeypatch.setattr(doc_office, "MAX_CHARS", 30_000)
    from openpyxl import Workbook
    p = tmp_path / "wide.xlsx"
    wb = Workbook(); ws = wb.active
    for i in range(8000):
        ws.append([f"row{i}", i, "padding value here"])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated" in md.lower()
    assert "of 8000 rows" in md  # the real total ...
    assert "of 1048576 rows" not in md  # ... not the structural row-per-sheet cap


# --------------------------------------------------------------------------
# Item 4: the row/column structural caps, exercised on their own — narrow
# rows so the character budget never fires first and can't mask a broken
# cap. monkeypatch keeps this independent of the module's real constants.
# --------------------------------------------------------------------------

def test_xlsx_row_cap_truncates_before_the_character_budget(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_office, "XLSX_MAX_ROWS", 10)
    from openpyxl import Workbook
    p = tmp_path / "narrow_rows.xlsx"
    wb = Workbook(); ws = wb.active
    for i in range(20):
        ws.append([i])  # single short column — negligible character cost
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated after 10 of 20 rows" in md
    assert "10-row-per-sheet limit" in md


def test_xlsx_col_cap_bounds_a_single_stray_wide_row(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_office, "XLSX_MAX_COLS", 5)
    from openpyxl import Workbook
    p = tmp_path / "narrow_cols.xlsx"
    wb = Workbook(); ws = wb.active
    ws.append(list(range(20)))  # one row, 20 short columns
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    header_line = next(line for line in md.splitlines() if line.startswith("|"))
    # capped to the first 5 of 20 columns => values 0..4 only, 6 pipes
    # ("| 0 | 1 | 2 | 3 | 4 |")
    assert header_line.count("|") == 6
    assert "4 |" in header_line
    assert "5 |" not in header_line


# --------------------------------------------------------------------------
# .docx conversion must be bounded in TIME and memory even though it is no
# longer bounded in OUTPUT SIZE — walking `document.element.body` fully via
# python-docx's own `Paragraph.style`/`Paragraph.text` properties (each of
# which re-resolves from scratch, via an xpath call, on every single access)
# was a verified O(document size) cost regardless of how much output
# survived any budget: measured ~0.4ms/paragraph, 111.65s at 300,000
# paragraphs in a plain 0.84MB file. The fix resolves both once (a style-id
# -> heading-level map, and a direct-lxml text/style-id reader) instead of
# racing a small budget against a slow per-access cost — so full,
# UNTRUNCATED conversion of the same 300,000-paragraph document is now the
# thing under test, not truncation. See the module docstring for the full
# writeup and the measured before/after at 10k/100k/300k paragraphs.
# --------------------------------------------------------------------------

def _build_docx_with_paragraphs(path, count):
    """Build `count` trivial paragraphs directly via lxml rather than
    `Document.add_paragraph` in a loop — the latter is itself slow enough at
    this scale (multiple minutes for 300,000 calls) to make a perf test
    built that way unusable; this stays under ~2.5s so the test itself
    measures doc_office, not python-docx's own paragraph-insertion cost."""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    d = Document()
    body = d.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for i in range(count):
        p = etree.SubElement(body, qn("w:p"))
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = f"paragraph number {i} with a bit of body text in it"
    if sect_pr is not None:
        body.remove(sect_pr)
        body.append(sect_pr)
    d.save(str(path))


def test_docx_huge_paragraph_count_converts_quickly_and_bounded(tmp_path):
    """Was truncation-at-30,000-characters: MAX_CHARS is now a generous
    resource ceiling (bound resources, not context), so this converts EVERY
    paragraph — the point of this test is that doing so is still fast,
    which is only true because `_build_heading_levels`/`_paragraph_text`/
    `_paragraph_style_id` removed the per-paragraph xpath costs that used
    to make full conversion of a document this size take 111.65s (see the
    module docstring). Measured on this machine after the fix: ~2.0s for
    this exact 300,000-paragraph document — the assertion below leaves
    generous headroom above that so the test itself doesn't flake, while
    still catching a real regression back toward the old per-access cost
    (which would fail this at well over a minute)."""
    p = tmp_path / "huge.docx"
    _build_docx_with_paragraphs(p, 300_000)

    t0 = time.time()
    md, images = doc_office.convert(p, tmp_path)
    elapsed = time.time() - t0

    assert images == []
    assert "paragraph number 0 " in md
    assert "paragraph number 299999 " in md  # every paragraph survives — full conversion
    assert "truncated" not in md.lower()
    assert elapsed < 15.0, f"conversion took {elapsed:.2f}s — should be a few seconds at most"


def test_docx_single_enormous_table_is_capped(tmp_path):
    """A single pathological table must not be rendered in full before
    DOCX_MAX_TABLE_ROWS (independent of MAX_CHARS — see the module
    docstring) ever gets a chance to reject it — previously a 20,000x6
    table cost ~4.4s / +52MB RSS regardless of any character budget.
    DOCX_MAX_TABLE_ROWS (2000) is unchanged by today's "convert fully"
    redesign — unlike MAX_CHARS, it bounds a genuinely separate,
    per-block cost (constructing real python-docx cell/paragraph/run
    objects to read `.text`, proportional to table size), not a
    context-window number — so a 3000-row table still truncates at 2000
    rows. The expected output size below is therefore sized to
    DOCX_MAX_TABLE_ROWS (~18 bytes/row for this fixture's narrow "| x | x |
    x | x |" rows), not to the old MAX_CHARS."""
    from docx import Document
    p = tmp_path / "big_table.docx"
    d = Document()
    t = d.add_table(rows=3000, cols=4)
    for r in t.rows:
        for c in r.cells:
            c.text = "x"
    d.save(str(p))

    t0 = time.time()
    md, _ = doc_office.convert(p, tmp_path)
    elapsed = time.time() - t0

    assert "truncated after 2000 rows" in md.lower()
    assert len(md) < 45_000  # ~2000 rows of "| x | x | x | x |", not 3000
    assert elapsed < 5.0, f"conversion took {elapsed:.2f}s — should be a small fraction of a second"


# --------------------------------------------------------------------------
# "Convert fully, bound resources not context" — real-shaped documents that
# would previously have lost most of their content to MAX_CHARS now convert
# in full against the DEFAULT (generous) ceiling, no monkeypatching needed.
# --------------------------------------------------------------------------

def test_xlsx_forty_thousand_rows_converts_fully(tmp_path):
    """Previously ~800 rows survived (30,000-character MAX_CHARS); a sheet
    this size is an entirely ordinary bulk export (a pandas/ETL dump, a
    CRM report) and must convert in full against the new resource
    ceiling."""
    from openpyxl import Workbook
    p = tmp_path / "forty_k.xlsx"
    wb = Workbook(); ws = wb.active
    for i in range(40_000):
        ws.append([f"row{i}", i, "some real content in this column"])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated" not in md.lower()
    assert "| row0 | 0 | some real content in this column |" in md
    assert "| row20000 | 20000 | some real content in this column |" in md
    assert "| row39999 | 39999 | some real content in this column |" in md


def test_xlsx_three_sheets_with_huge_first_sheet_all_present_in_full(tmp_path):
    """A workbook where the FIRST sheet alone is large must not starve the
    later sheets, nor itself be cut short, against the default ceiling —
    every sheet, including the big one, converts completely."""
    from openpyxl import Workbook
    from es.capabilities import read as read_cap

    p = tmp_path / "three_sheets.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for i in range(20_000):
        ws.append([f"row{i}", i, "padding content for this row"])
    summary = wb.create_sheet("Summary")
    for i in range(100):
        summary.append([f"Metric {i}", i])
    notes = wb.create_sheet("Notes")
    for i in range(100):
        notes.append([f"Note line {i}"])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated" not in md.lower()

    outline = read_cap.outline(md)
    assert [s["title"] for s in outline] == ["Data", "Summary", "Notes"]

    data_id = next(s["id"] for s in outline if s["title"] == "Data")
    summary_id = next(s["id"] for s in outline if s["title"] == "Summary")
    notes_id = next(s["id"] for s in outline if s["title"] == "Notes")

    data_section = read_cap.section(md, data_id)
    summary_section = read_cap.section(md, summary_id)
    notes_section = read_cap.section(md, notes_id)

    assert "row0" in data_section and "row19999" in data_section
    assert "Metric 0" in summary_section and "Metric 99" in summary_section
    assert "Note line 0" in notes_section and "Note line 99" in notes_section


def test_xlsx_sixty_sheet_workbook_every_sheet_has_real_data_rows(tmp_path):
    """A workbook with many small-to-modest sheets (a common shape for a
    per-region/per-month export) must not lose any sheet's real content —
    every one of the 60 sheets here must have actual data rows, not just a
    bare heading, against the default ceiling."""
    from openpyxl import Workbook
    from es.capabilities import read as read_cap

    p = tmp_path / "sixty_sheets.xlsx"
    wb = Workbook()
    ws0 = wb.active
    ws0.title = "Sheet0"
    for i in range(500):
        ws0.append([f"row{i}", i])
    for s in range(1, 60):
        ws = wb.create_sheet(f"Sheet{s}")
        for i in range(30):
            ws.append([f"r{s}_{i}", i])
    wb.save(str(p))

    md, _ = doc_office.convert(p, tmp_path)
    assert "truncated" not in md.lower()

    outline = read_cap.outline(md)
    assert len(outline) == 60
    for i, s in enumerate(outline):
        section = read_cap.section(md, s["id"])
        # A row DEEP into each sheet (not row 0) must be present — proving
        # every sheet converted fully, not just that the separate "first
        # row always shown regardless of budget" rule alone kept it afloat
        # (same rigor as
        # test_xlsx_later_sheets_reachable_when_first_sheet_exhausts_budget).
        if i == 0:
            assert "row499" in section  # Sheet0's last row (of 500)
        else:
            assert f"r{i}_29" in section  # SheetN's last row (of 30)
