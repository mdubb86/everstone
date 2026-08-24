import socket, subprocess, sys, time
from pathlib import Path

import pytest, requests

FIXTURES_DIR = Path(__file__).parent / "fixtures"
PDF_FIXTURES_DIR = FIXTURES_DIR / "pdf"


def _free_port():
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0)); return s.getsockname()[1]


@pytest.fixture
def radicale(tmp_path):
    storage = tmp_path / "collections"; storage.mkdir()
    port = _free_port()
    proc = subprocess.Popen(
        [sys.executable, "-m", "radicale", "--server-hosts", f"127.0.0.1:{port}",
         "--auth-type", "none", "--storage-filesystem-folder", str(storage),
         "--logging-level", "warning"],
        stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    base = f"http://127.0.0.1:{port}"
    for _ in range(100):
        try:
            requests.request("OPTIONS", base + "/", timeout=0.5); break
        except requests.exceptions.ConnectionError:
            time.sleep(0.1)
    else:
        proc.terminate(); raise RuntimeError("Radicale did not start")
    yield base
    proc.terminate()
    try: proc.wait(timeout=5)
    except subprocess.TimeoutExpired: proc.kill()


@pytest.fixture
def text_pdf(tmp_path):
    """Two-page PDF with a real text layer on both pages (plain drawString
    calls — no actual table; pdfplumber's extract_tables() finds none here)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 720, "Fall Season Schedule")
    c.drawString(72, 700, "Team: Thunder U10")
    c.showPage()
    c.drawString(72, 720, "Game 1  Sat Sep 5  9:00 AM  Field 3")
    c.drawString(72, 700, "Game 2  Sat Sep 12  11:00 AM  Field 1")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def table_pdf(tmp_path):
    """One page with prose plus a genuinely bordered (reportlab GRID-style)
    table that pdfplumber's find_tables()/extract_tables() actually detect —
    confirmed empirically, not assumed. One cell contains a literal '|' to
    exercise pipe-escaping in the emitted Markdown table."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Table, TableStyle

    p = tmp_path / "table.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Roster")
    data = [
        ["Name", "Position", "Number"],
        ["Alice", "Forward|Winger", "9"],
        ["Bob", "Goalie", "1"],
    ]
    t = Table(data, colWidths=[100, 100, 100])
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    t.wrap(0, 0)
    t.drawOn(c, 72, 600)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def table_only_pdf(tmp_path):
    """One page containing NOTHING but a bordered table — no heading, no
    other text anywhere on the page. Regression fixture for a blank-page
    test that (wrongly) decides on table-excluded text: that text would be
    empty here even though the page is far from blank."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Table, TableStyle

    p = tmp_path / "table_only.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    data = [
        ["Name", "Position", "Number"],
        ["Alice", "Forward", "9"],
        ["Bob", "Goalie", "1"],
    ]
    t = Table(data, colWidths=[100, 100, 100])
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    t.wrap(0, 0)
    t.drawOn(c, 72, 600)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def photo_page_pdf(tmp_path):
    """One page: a short heading plus ONE embedded photo — the plain "a page
    has an image" case, distinct from scanned_pdf (which has no text layer
    at all)."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    photo = tmp_path / "photo.png"
    Image.new("RGB", (300, 200), (100, 150, 200)).save(photo)

    p = tmp_path / "photo_page.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Team photo from the fall season banquet follows below.")
    c.drawImage(str(photo), 72, 500, width=300, height=200)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def text_photo_text_pdf(tmp_path):
    """One page with a photo sandwiched between two distinguishable blocks
    of text — used to assert the image's Markdown link lands at its
    POSITION in the reading order, not appended after all the text."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    photo = tmp_path / "photo.png"
    Image.new("RGB", (300, 150), (100, 150, 200)).save(photo)

    p = tmp_path / "text_photo_text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "INTRO TEXT ABOVE THE PHOTO")
    c.drawImage(str(photo), 72, 500, width=300, height=150)
    c.drawString(72, 400, "OUTRO TEXT BELOW THE PHOTO")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def two_images_pdf(tmp_path):
    """One page with two distinct embedded photos at different vertical
    positions — both must be extracted to separate files and linked."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    top_photo = tmp_path / "top.png"
    Image.new("RGB", (200, 100), (200, 0, 0)).save(top_photo)
    bottom_photo = tmp_path / "bottom.png"
    Image.new("RGB", (200, 100), (0, 0, 200)).save(bottom_photo)

    p = tmp_path / "two_images.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Two photos on one page.")
    c.drawImage(str(top_photo), 72, 600, width=200, height=100)
    c.drawImage(str(bottom_photo), 72, 400, width=200, height=100)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def image_on_second_page_pdf(tmp_path):
    """Two pages: page 1 is plain text with no image, page 2 has text plus
    one embedded photo — the image link must land in page 2's section only."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    photo = tmp_path / "photo.png"
    Image.new("RGB", (200, 150), (50, 200, 50)).save(photo)

    p = tmp_path / "image_on_second_page.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 720, "Page one has only text.")
    c.showPage()
    c.drawString(72, 720, "Page two has text and a photo.")
    c.drawImage(str(photo), 72, 500, width=200, height=150)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def scanned_pdf(tmp_path):
    """Single-page PDF with NO text layer — an image pasted onto the page."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    png = tmp_path / "scan.png"
    Image.new("RGB", (400, 300), (200, 200, 200)).save(png)
    p = tmp_path / "scanned.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawImage(str(png), 72, 400, width=400, height=300)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def vector_chart_pdf(tmp_path):
    """One page: a bar chart drawn ENTIRELY in vector paths (two axis lines
    plus three filled bars) — no embedded raster image at all. This is the
    headline case Task 2 closes: a chart pdfplumber's text/table extraction
    is blind to (no text, no embedded image) and Task 1's image extraction
    is equally blind to (nothing in page.get_objects(FPDF_PAGEOBJ_IMAGE))."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "vector_chart.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Revenue")
    c.line(100, 400, 100, 600)
    c.line(100, 400, 400, 400)
    c.rect(120, 400, 30, 80, fill=1)
    c.rect(170, 400, 30, 150, fill=1)
    c.rect(220, 400, 30, 60, fill=1)
    c.drawString(120, 390, "Q1")
    c.drawString(170, 390, "Q2")
    c.drawString(220, 390, "Q3")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def lone_rule_pdf(tmp_path):
    """One page: a single horizontal rule under a heading, and nothing else
    vector — the size-floor regression guard. A rule line is a real, common
    idiom (section separators, underlines) and must NOT be rasterized as if
    it were a drawing."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "lone_rule.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Section header")
    c.line(72, 730, 540, 730)
    c.drawString(72, 700, "Body text below the rule.")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def table_and_chart_pdf(tmp_path):
    """One page with BOTH a genuinely bordered table (find_tables() detects
    it, and its border lines must be subtracted from drawing detection) AND
    a separate vector bar chart lower on the page. Regression fixture for
    the subtraction step: without it, the table's own border lines would be
    clustered into a second, spurious drawing alongside the real chart."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Table, TableStyle

    p = tmp_path / "table_and_chart.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    data = [["Name", "Position", "Number"], ["Alice", "Forward", "9"], ["Bob", "Goalie", "1"]]
    t = Table(data, colWidths=[100, 100, 100])
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    t.wrap(0, 0)
    t.drawOn(c, 72, 650)
    c.line(100, 300, 100, 500)
    c.line(100, 300, 400, 300)
    c.rect(120, 300, 30, 80, fill=1)
    c.rect(170, 300, 30, 150, fill=1)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def two_charts_side_by_side_pdf(tmp_path):
    """One page with TWO SEPARATE bar charts, well apart horizontally — the
    evidence for the clustering rule: if everything remaining on a page were
    treated as a single drawing, this page would wrongly produce ONE image
    spanning (and mostly blank between) both charts instead of two."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "two_charts.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.line(100, 400, 100, 600)
    c.line(100, 400, 250, 400)
    c.rect(110, 400, 20, 60, fill=1)
    c.rect(140, 400, 20, 100, fill=1)
    c.rect(170, 400, 20, 40, fill=1)
    c.line(350, 400, 350, 600)
    c.line(350, 400, 500, 400)
    c.rect(360, 400, 20, 90, fill=1)
    c.rect(390, 400, 20, 30, fill=1)
    c.rect(420, 400, 20, 120, fill=1)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def chart_between_text_pdf(tmp_path):
    """A vector chart sandwiched between two distinguishable blocks of text —
    used to assert the drawing's Markdown link lands at its POSITION in the
    reading order, not appended after all the text (same guarantee Task 1
    gives embedded images; see test_image_position_is_interleaved_with_surrounding_text)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "chart_between_text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "INTRO TEXT ABOVE THE CHART")
    c.line(100, 500, 100, 650)
    c.line(100, 500, 300, 500)
    c.rect(120, 500, 30, 100, fill=1)
    c.rect(170, 500, 30, 60, fill=1)
    c.drawString(72, 400, "OUTRO TEXT BELOW THE CHART")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def photo_and_chart_pdf(tmp_path):
    """One page with BOTH an embedded raster photo AND a separate vector
    chart, at different vertical positions — both must be extracted, each
    with its own distinct filename (`-iMM` for the photo, `-dMM` for the
    chart)."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    photo = tmp_path / "photo.png"
    Image.new("RGB", (200, 100), (100, 150, 200)).save(photo)

    p = tmp_path / "photo_and_chart.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawImage(str(photo), 72, 620, width=200, height=100)
    c.line(100, 350, 100, 550)
    c.line(100, 350, 300, 350)
    c.rect(120, 350, 30, 100, fill=1)
    c.rect(170, 350, 30, 60, fill=1)
    c.showPage()
    c.save()
    return p


@pytest.fixture
def realistic_report_page_pdf(tmp_path):
    """A page built from common REAL, non-chart vector idioms — a header
    rule, a footer rule, a genuinely bordered table, and a filled background
    shading band — used to measure how much false-positive noise survives
    table subtraction and the size floor. Deliberately has no page-spanning
    border frame (see the task report's False Positives discussion for why
    that case is a known, accepted limitation rather than one this fixture
    manufactures to fail)."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Table, TableStyle

    p = tmp_path / "realistic_report_page.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawString(72, 740, "Quarterly Report")
    c.line(72, 725, 540, 725)  # header rule
    data = [["Item", "Q1", "Q2"], ["Widgets", "100", "120"], ["Gadgets", "80", "95"]]
    t = Table(data, colWidths=[150, 100, 100])
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    t.wrap(0, 0)
    t.drawOn(c, 72, 600)
    c.setFillColorRGB(0.9, 0.9, 0.9)
    c.rect(72, 460, 300, 40, fill=1, stroke=0)  # decorative shading band
    c.setFillColorRGB(0, 0, 0)
    c.drawString(80, 475, "Note: shaded band, decorative only")
    c.line(72, 100, 540, 100)  # footer rule
    c.drawString(72, 80, "Footer text")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def icml_numpapers_pdf():
    """A real, small (2.8KB) academic-conference PDF whose one page IS a
    vector bar chart (42 filled bars, drawn as page.rects, no embedded
    image) — copied from hermes-agent's own research-paper-writing skill
    templates (MIT-licensed, ships inside the same container this code
    runs in). The strongest available positive case: a real document, not
    a synthetic one."""
    return PDF_FIXTURES_DIR / "icml_numpapers.pdf"


@pytest.fixture
def colm2025_conference_pdf():
    """A real 5-page academic paper template (122KB, MIT-licensed, same
    source as icml_numpapers_pdf) with NO embedded images and NO real
    tables — used as a false-positive check on ordinary prose pages (rules,
    underlines, header separators)."""
    return PDF_FIXTURES_DIR / "colm2025_conference.pdf"


@pytest.fixture
def example_paper_pdf():
    """A real 7-page academic paper template (193KB, MIT-licensed, same
    source as icml_numpapers_pdf) containing a genuinely bordered table
    (detected and subtracted) alongside a real bar chart on the same page —
    used to confirm table subtraction holds on a real, not synthetic,
    document."""
    return PDF_FIXTURES_DIR / "example_paper.pdf"


@pytest.fixture
def csv_file(tmp_path):
    p = tmp_path / "roster.csv"
    p.write_text("Name,Position,Number\nAlice,Forward,9\nBob,Goalie,1\n", encoding="utf-8")
    return p


@pytest.fixture
def json_file(tmp_path):
    p = tmp_path / "data.json"
    p.write_text('{"team": "Thunder U10", "games": [{"opp": "Fury"}, {"opp": "SC"}]}',
                 encoding="utf-8")
    return p


@pytest.fixture
def txt_file(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("Practice moved to Thursday.\nBring the blue kit.\n", encoding="utf-8")
    return p


@pytest.fixture
def ics_file(tmp_path):
    p = tmp_path / "schedule.ics"
    p.write_text(
        "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//Test//EN\r\n"
        "BEGIN:VEVENT\r\nUID:1\r\nSUMMARY:Game 1 vs Cedar Park Fury\r\n"
        "DTSTART:20260905T140000Z\r\nDTEND:20260905T153000Z\r\n"
        "LOCATION:Kelly Reeves Athletic Complex - Field 3\r\nEND:VEVENT\r\n"
        "BEGIN:VEVENT\r\nUID:2\r\nSUMMARY:Game 2 vs Round Rock SC\r\n"
        "DTSTART:20260912T160000Z\r\nDTEND:20260912T173000Z\r\n"
        "LOCATION:Old Settlers Park - Field 1\r\nEND:VEVENT\r\n"
        "END:VCALENDAR\r\n", encoding="utf-8")
    return p


@pytest.fixture
def docx_file(tmp_path):
    from docx import Document
    p = tmp_path / "letter.docx"
    d = Document()
    d.add_heading("Season Overview", level=1)
    d.add_paragraph("Practices are Tuesdays and Thursdays.")
    d.add_heading("Fees", level=2)
    d.add_paragraph("Club dues are due Sep 1.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Item"; t.cell(0, 1).text = "Cost"
    t.cell(1, 0).text = "Kit";  t.cell(1, 1).text = "$65"
    d.save(str(p))
    return p


@pytest.fixture
def docx_with_inline_image(tmp_path):
    """Heading, paragraph, an inline embedded photo (its own paragraph — the
    shape `Document.add_picture` produces), paragraph — used to assert the
    image's link lands BETWEEN the two prose paragraphs, not appended after
    everything."""
    import io
    from docx import Document
    from PIL import Image

    photo = io.BytesIO()
    Image.new("RGB", (60, 40), (10, 20, 30)).save(photo, format="PNG")
    photo.seek(0)

    p = tmp_path / "with_image.docx"
    d = Document()
    d.add_heading("Season Overview", level=1)
    d.add_paragraph("FIRST paragraph before the photo.")
    d.add_picture(photo)
    d.add_paragraph("LAST paragraph after the photo.")
    d.save(str(p))
    return p


@pytest.fixture
def docx_with_two_images(tmp_path):
    """Two distinct embedded photos, each in its own paragraph, with a
    distinguishable marker paragraph between them — both must be extracted
    to separate files and linked in document order."""
    import io
    from docx import Document
    from PIL import Image

    photo1 = io.BytesIO()
    Image.new("RGB", (30, 30), (200, 0, 0)).save(photo1, format="PNG")
    photo1.seek(0)
    photo2 = io.BytesIO()
    Image.new("RGB", (30, 30), (0, 0, 200)).save(photo2, format="PNG")
    photo2.seek(0)

    p = tmp_path / "two_images.docx"
    d = Document()
    d.add_paragraph("before both photos")
    d.add_picture(photo1)
    d.add_paragraph("BETWEEN THE TWO PHOTOS")
    d.add_picture(photo2)
    d.add_paragraph("after both photos")
    d.save(str(p))
    return p


@pytest.fixture
def docx_with_table_cell_image(tmp_path):
    """A table with an inline photo embedded in one cell's own paragraph —
    used to confirm the image is reported AFTER the table (never inside a
    pipe-table cell, which would corrupt the table's own syntax)."""
    import io
    from docx import Document
    from PIL import Image

    photo = io.BytesIO()
    Image.new("RGB", (30, 30), (9, 9, 9)).save(photo, format="PNG")
    photo.seek(0)

    p = tmp_path / "table_image.docx"
    d = Document()
    d.add_paragraph("Roster")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Item"
    t.cell(0, 1).text = "Photo"
    t.cell(1, 0).text = "Widget"
    t.cell(1, 1).paragraphs[0].add_run().add_picture(photo)
    d.save(str(p))
    return p


@pytest.fixture
def docx_with_duplicated_image_relationship(tmp_path):
    """The SAME embedded-image relationship id referenced from TWO separate
    paragraphs (built by literally duplicating the `<w:drawing>`-bearing
    paragraph's XML onto a second paragraph, reusing its `r:embed` id) — the
    realistic shape of e.g. a letterhead logo placed at both the top and
    bottom of a template. Used to confirm this writes ONE file, not two, and
    that the SAME index is linked from both positions."""
    import copy
    import io
    from docx import Document
    from docx.oxml.ns import qn
    from PIL import Image

    photo = io.BytesIO()
    Image.new("RGB", (20, 20), (1, 2, 3)).save(photo, format="PNG")
    photo.seek(0)

    first = tmp_path / "single.docx"
    d = Document()
    d.add_paragraph("Para A")
    d.add_picture(photo)
    d.add_paragraph("Para B")
    d.save(str(first))

    d2 = Document(str(first))
    body = d2.element.body
    drawing_para = next(
        c for c in body.iterchildren()
        if c.tag == qn("w:p") and c.findall(".//" + qn("w:drawing")))
    dup_para = copy.deepcopy(drawing_para)
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.remove(sect_pr)
    body.append(dup_para)
    if sect_pr is not None:
        body.append(sect_pr)

    p = tmp_path / "duplicated.docx"
    d2.save(str(p))
    return p


@pytest.fixture
def xlsx_file(tmp_path):
    from openpyxl import Workbook
    p = tmp_path / "roster.xlsx"
    wb = Workbook()
    ws = wb.active; ws.title = "Roster"
    ws.append(["Name", "Number"]); ws.append(["Alice", 9]); ws.append(["Bob", 1])
    ws2 = wb.create_sheet("Fees")
    ws2.append(["Item", "Cost"]); ws2.append(["Kit", 65])
    wb.save(str(p))
    return p
