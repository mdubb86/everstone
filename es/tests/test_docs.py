import asyncio
import os
import re
import time

import pytest

from es.capabilities import doc_ics, doc_text, docs


# --- fixtures for malformed/unusual PDFs -----------------------------------
# Kept local to this module (not conftest.py) since these exercise docs.py's
# own error-mapping, not doc_pdf's conversion logic.

@pytest.fixture
def encrypted_pdf(tmp_path):
    """A real password-protected PDF (reportlab supports this natively —
    no extra dependency needed)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.pdfencrypt import StandardEncryption
    from reportlab.pdfgen import canvas

    p = tmp_path / "encrypted.pdf"
    enc = StandardEncryption("userpw", ownerPassword="ownerpw", canPrint=1)
    c = canvas.Canvas(str(p), pagesize=letter, encrypt=enc)
    c.drawString(72, 720, "secret")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def corrupt_pdf(tmp_path):
    """Has the %PDF magic bytes but no usable structure behind them."""
    p = tmp_path / "corrupt.pdf"
    p.write_bytes(b"%PDF-1.4\n%truncated, no xref table, no /Root")
    return p


@pytest.fixture
def empty_pdf(tmp_path):
    p = tmp_path / "empty.pdf"
    p.write_bytes(b"")
    return p


@pytest.fixture
def fake_pdf(tmp_path):
    """A plain text file wearing a .pdf extension."""
    p = tmp_path / "fake.pdf"
    p.write_text("this is just a text file renamed to .pdf\n" * 5)
    return p


@pytest.fixture
def one_scanned_one_text_pdf(tmp_path):
    """Page 1 is image-only (auto-rendered by extract); page 2 has a real
    text layer. Gives extract() a non-empty images.json (from page 1) that
    is distinguishable from a PNG a LATER image_pages call would drop into
    the same artifact dir for page 2."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    png = tmp_path / "scan.png"
    Image.new("RGB", (400, 300), (200, 200, 200)).save(png)
    p = tmp_path / "one_scanned_one_text.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    c.drawImage(str(png), 72, 400, width=400, height=300)
    c.showPage()
    c.drawString(72, 720, "Page two has a real text layer.")
    c.showPage()
    c.save()
    return p


@pytest.fixture
def zero_page_pdf(tmp_path):
    """Opens fine as a PDF but legitimately has zero pages."""
    from reportlab.pdfgen import canvas

    p = tmp_path / "zero.pdf"
    canvas.Canvas(str(p)).save()
    return p


# --- extract() ---------------------------------------------------------

def test_extract_returns_markdown_and_doc_id(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert len(out["doc_id"]) == 12
    assert out["kind"] == "pdf"
    assert out["page_count"] == 2
    assert "Fall Season Schedule" in out["preview"]
    assert out["complete"] is True


def test_extract_rejects_path_outside_roots(text_pdf, tmp_path):
    other = tmp_path / "elsewhere"
    other.mkdir()
    with pytest.raises(docs.paths.SourceForbidden):
        docs.extract(str(text_pdf), roots=[other], cache_root=tmp_path)


def test_extract_rejects_unsupported_extension(tmp_path):
    f = tmp_path / "notes.rtf"
    f.write_text("x")
    with pytest.raises(docs.UnsupportedDocument) as e:
        docs.extract(str(f), roots=[tmp_path], cache_root=tmp_path)
    assert ".rtf" in str(e.value)


def test_extract_writes_markdown_into_the_cache(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    cached = tmp_path / ".es" / out["doc_id"] / "doc.md"
    assert cached.is_file()
    assert "Fall Season Schedule" in cached.read_text(encoding="utf-8")


def _full_cached_markdown(out: dict, cache_root) -> str:
    """The full doc.md extract() cached."""
    return (cache_root / ".es" / out["doc_id"] / "doc.md").read_text(encoding="utf-8")


# --- a converter's own self-truncation must still be detectable in-band ----
# A converter that truncates ITSELF (doc_text/doc_office/doc_ics, each at its
# own resource-ceiling MAX_CHARS — see each module's own comment) says so with
# a "*(truncated ...)*" marker built through doc_support.truncation_marker,
# so a plain `"truncated after" in cached` substring check is enough to
# confirm it — no detector function needed on this side. (Older versions of
# these tests also asserted docs._converter_self_truncated(cached) directly;
# that helper only ever existed to feed extract()'s pre-receipt `truncated`
# flag, has no other caller now that the flag is gone, and is removed. The
# property it duplicated — self-truncation is honestly marked in-band — is
# still exercised by the plain substring check every test below already made
# independently of it.)
#
# Converters now convert in FULL (bounded by a resource ceiling in the tens of
# millions of characters, not a context-window budget), because doc.md is
# cached and es_read pages it — truncating at conversion time destroyed data
# nothing needed to destroy. So these tests monkeypatch the ceiling DOWN to
# force self-truncation, rather than relying on a real document crossing it.
# The property under test is unchanged; only how it is provoked and checked.

def test_self_truncation_is_reported_for_csv(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_text, "MAX_CHARS", 4_000)
    rows = "\n".join(f"{i},value-{i}" for i in range(6000))
    p = tmp_path / "many_rows.csv"
    p.write_text("id,value\n" + rows + "\n", encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_json(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_text, "MAX_CHARS", 4_000)
    import json as jsonlib
    data = [{"i": i, "note": "padding text to grow each entry a bit"} for i in range(1200)]
    p = tmp_path / "big.json"
    p.write_text(jsonlib.dumps(data), encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_txt(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_text, "MAX_CHARS", 4_000)
    line = "x" * 39 + "\n"
    p = tmp_path / "big.txt"
    p.write_text(line * 1000, encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_md(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_text, "MAX_CHARS", 4_000)
    line = "x" * 39 + "\n"
    p = tmp_path / "big.md"
    p.write_text(line * 1000, encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_ics(tmp_path, monkeypatch):
    monkeypatch.setattr(doc_ics, "MAX_ICS_CHARS", 4_000)
    events = []
    for i in range(500):
        events.append(
            "BEGIN:VEVENT\r\nUID:{n}\r\nSUMMARY:Game {n} vs Team {n}\r\n"
            "DTSTART:202609{d:02d}T140000Z\r\nLOCATION:Field {n}\r\n"
            "END:VEVENT\r\n".format(n=i, d=(i % 28) + 1))
    p = tmp_path / "many_events.ics"
    p.write_text(
        "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//Test//EN\r\n"
        + "".join(events) + "END:VCALENDAR\r\n", encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_docx(tmp_path, monkeypatch):
    from docx import Document
    from es.capabilities import doc_office
    monkeypatch.setattr(doc_office, "MAX_CHARS", 4_000)

    p = tmp_path / "many_paragraphs.docx"
    d = Document()
    for i in range(1200):
        d.add_paragraph(f"Paragraph number {i} with some filler text to pad it out.")
    d.save(str(p))
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_xlsx(tmp_path, monkeypatch):
    from openpyxl import Workbook
    from es.capabilities import doc_office
    monkeypatch.setattr(doc_office, "MAX_CHARS", 4_000)

    p = tmp_path / "many_rows.xlsx"
    wb = Workbook()
    ws = wb.active
    for i in range(3000):
        ws.append([f"row{i:05d}", f"val-{i:05d}", f"val-{i:05d}"])
    wb.save(str(p))
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached


def test_self_truncation_is_reported_for_write_only_xlsx_with_no_dimension(tmp_path, monkeypatch):
    """Regression guard for a LIVE bug: `openpyxl.Workbook(write_only=True)`
    never writes a `<dimension>` element to a sheet's XML (see doc_office's
    module docstring), so this sheet's true row count can never be
    determined without a full scan — the exact case
    doc_office._sheet_truncation_note's "total_rows is None" branch exists
    for. That branch's marker embeds a NESTED parenthetical aside ("...could
    not be determined (its XML has no declared dimension)"), and before this
    fix docs.py detected self-truncation with a regex that required NO
    parentheses at all between "*(" and the closing ")*" — so this exact
    marker was never detected, and `truncated` came back False.

    Verified live in the running container before this fix: a 5,000-row
    write_only workbook rendered only 856 rows (the rest silently dropped)
    and the es_doc_extract envelope reported `truncated: false` — the worst
    failure mode for this tool, a successful-looking result with wrong
    content. This is the test that would have caught it.

    doc_office.MAX_CHARS is monkeypatched down for the same reason as the
    docx/xlsx self-truncation tests above: it is now a generous resource
    ceiling a 5,000-row sheet never reaches on its own, so provoking
    self-truncation here means constraining the ceiling, not just building
    a bigger workbook.
    """
    from openpyxl import Workbook
    from es.capabilities import doc_office
    monkeypatch.setattr(doc_office, "MAX_CHARS", 4_000)

    p = tmp_path / "write_only_many_rows.xlsx"
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Sheet")
    for i in range(5000):
        ws.append([f"row{i:05d}", f"val-{i:05d}", f"val-{i:05d}"])
    wb.save(str(p))

    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated after" in cached
    assert "no declared dimension" in cached


def test_small_write_only_xlsx_with_no_dimension_is_not_falsely_truncated(tmp_path):
    """The other half of the write_only regression above: a small
    dimension-less sheet must not be falsely detected as self-truncated.
    Found while manually verifying this fix — `doc_office._sheet_truncation_
    note` used to infer "was this cut short?" from `kept < capped_rows`,
    where `capped_rows` was silently the bare XLSX_MAX_ROWS fallback ceiling
    (not this sheet's real size) whenever the dimension was unknown, so even
    a 2-row write_only sheet satisfied that comparison and was falsely
    reported as truncated. Fixed alongside the detection bug since both
    live in the same code path this task touches."""
    from openpyxl import Workbook

    p = tmp_path / "tiny_write_only.xlsx"
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Sheet")
    ws.append(["a", "b"])
    ws.append(["c", "d"])
    wb.save(str(p))

    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    cached = _full_cached_markdown(out, tmp_path)
    assert "truncated" not in cached.lower()
    assert out["complete"] is True


def test_extract_purges_stale_artifacts(text_pdf, tmp_path):
    from es import doc_cache
    stale = doc_cache.artifact_dir(tmp_path, "stale0000000")
    old = time.time() - (25 * 3600)
    os.utime(stale, (old, old))
    docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert not stale.exists()


def test_extract_second_call_is_a_cache_hit_not_a_reconvert(text_pdf, tmp_path, monkeypatch):
    calls = []
    original = docs.doc_pdf.convert

    def spy(*args, **kwargs):
        calls.append((args, kwargs))
        return original(*args, **kwargs)

    monkeypatch.setattr(docs.doc_pdf, "convert", spy)

    first = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    second = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)

    assert len(calls) == 1  # convert() ran once; the second call was a cache hit
    assert second["preview"] == first["preview"]
    assert second["doc_id"] == first["doc_id"]


def test_extract_recovers_from_a_corrupted_cached_doc_md(text_pdf, tmp_path):
    """doc_id is a CONTENT hash: if doc.md was left truncated by a crash or
    ENOSPC, re-sending the identical PDF lands in the same artifact dir and
    must NOT raise (and must not stay broken for the 24h TTL) — an
    undecodable/unreadable doc.md must be treated as a cache miss, the same
    way a broken images.json already is, and the entry overwritten with a
    good reconversion."""
    first = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    md_path = tmp_path / ".es" / first["doc_id"] / "doc.md"
    md_path.write_bytes(b"\xff\xfe\x00 not valid utf-8 \x80\x81")

    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert "Fall Season Schedule" in out["preview"]

    # The bad cache entry was overwritten with the good reconversion.
    assert "Fall Season Schedule" in md_path.read_text(encoding="utf-8")


def test_cache_hit_images_come_from_manifest_not_a_directory_scan(
        one_scanned_one_text_pdf, tmp_path):
    """A cache-hit extract must report only the images the ORIGINAL extract
    produced (from images.json) — not every PNG that happens to sit in the
    artifact dir, including ones a later image_pages call drops there for
    pages the extract itself never rendered.

    extract()'s own return no longer carries `images` at all (the receipt
    contract this task adds — see test_extract_returns_a_receipt_not_the_
    document); the manifest itself is still cached and still the thing a
    cache-hit must not let a later image_pages call corrupt, so this asserts
    against docs.read_cached() (the shared accessor for that cache entry)
    directly."""
    first = docs.extract(str(one_scanned_one_text_pdf),
                         roots=[one_scanned_one_text_pdf.parent], cache_root=tmp_path)
    adir = tmp_path / ".es" / first["doc_id"]
    first_images = docs.read_cached(adir)["images"]
    assert len(first_images) == 1  # only page 1, the image-only page

    # image_pages="2" into the SAME artifact dir — page 2 has real text and
    # was never rendered by the plain extract() above, so its PNG is new to
    # the dir.
    rendered = docs.extract(str(one_scanned_one_text_pdf),
                            roots=[one_scanned_one_text_pdf.parent],
                            cache_root=tmp_path, image_pages="2")
    assert rendered["page_images"]
    assert len(list(adir.glob("*.png"))) == 2  # both PNGs now physically present

    docs.extract(str(one_scanned_one_text_pdf),
                 roots=[one_scanned_one_text_pdf.parent], cache_root=tmp_path)  # cache hit
    second_images = docs.read_cached(adir)["images"]
    assert second_images == first_images  # unchanged by the image_pages call


def test_extract_cache_hit_still_touches_artifact_dir(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    adir = tmp_path / ".es" / out["doc_id"]
    old = time.time() - 1000
    os.utime(adir, (old, old))

    docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)

    assert os.stat(adir).st_mtime > old


def test_extract_rejects_oversized_document(text_pdf, tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "MAX_DOCUMENT_BYTES", 10)
    with pytest.raises(docs.DocumentTooLarge) as e:
        docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_too_large"


def test_extract_missing_file_mentions_resend(text_pdf, tmp_path):
    missing = text_pdf.parent / "gone.pdf"
    with pytest.raises(docs.paths.SourceNotFound) as e:
        docs.extract(str(missing), roots=[text_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_not_found"
    assert "resend" in str(e.value).lower() or "re-send" in str(e.value).lower()


def test_extract_encrypted_pdf_names_the_remedy(encrypted_pdf, tmp_path):
    with pytest.raises(docs.EncryptedDocument) as e:
        docs.extract(str(encrypted_pdf), roots=[encrypted_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_encrypted"
    assert "password" in str(e.value).lower()


def test_extract_corrupt_pdf_is_unreadable(corrupt_pdf, tmp_path):
    with pytest.raises(docs.UnreadableDocument) as e:
        docs.extract(str(corrupt_pdf), roots=[corrupt_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_unreadable"


def test_extract_zero_byte_pdf_is_unreadable(empty_pdf, tmp_path):
    with pytest.raises(docs.UnreadableDocument) as e:
        docs.extract(str(empty_pdf), roots=[empty_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_unreadable"


def test_extract_mislabeled_text_file_is_unreadable(fake_pdf, tmp_path):
    with pytest.raises(docs.UnreadableDocument) as e:
        docs.extract(str(fake_pdf), roots=[fake_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_unreadable"


def test_extract_zero_page_pdf_is_unreadable(zero_page_pdf, tmp_path):
    with pytest.raises(docs.UnreadableDocument) as e:
        docs.extract(str(zero_page_pdf), roots=[zero_page_pdf.parent], cache_root=tmp_path)
    assert e.value.es_code == "doc_unreadable"


def test_extract_rejects_over_long_path_without_leaking_oserror(tmp_path):
    """A path that resolves inside an allowed root but is too long for the
    filesystem must not surface as a raw OSError (which would leak the full
    path and an opaque 'OSError' code to the agent)."""
    root = tmp_path / "root"
    root.mkdir()
    source = str(root / ("a" * 5000 + ".pdf"))
    with pytest.raises(docs.UnreadableDocument) as e:
        docs.extract(source, roots=[root], cache_root=tmp_path)
    assert e.value.es_code == "doc_unreadable"


def test_extract_rejects_empty_image_pages_string(text_pdf, tmp_path):
    """image_pages="" is a malformed selector, not a synonym for "render
    nothing" — only omitting the argument entirely (None) means that."""
    with pytest.raises(docs.InvalidPageRange):
        docs.extract(str(text_pdf), roots=[text_pdf.parent],
                    cache_root=tmp_path, image_pages="")


# --- extract()'s image_pages parameter --------------------------------------

def test_extract_image_pages_returns_page_images(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent],
                       cache_root=tmp_path, image_pages="1-2")
    assert len(out["page_images"]) == 2
    assert out["page_count"] == 2


def test_extract_without_image_pages_returns_an_empty_page_images_list(text_pdf, tmp_path):
    """`page_images` is always present — an empty list, not a missing key or
    null, is the true statement when image_pages was never asked for."""
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert out["page_images"] == []


def _make_pdf(path, n_pages):
    """An n-page PDF with real text on every page (used to check image_pages
    against documents of varying length)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=letter)
    for i in range(n_pages):
        c.drawString(72, 720, f"Page {i + 1} content")
        c.showPage()
    c.save()
    return path


def test_extract_image_pages_rejects_page_out_of_range(text_pdf, tmp_path):
    with pytest.raises(docs.InvalidPageRange):
        docs.extract(str(text_pdf), roots=[text_pdf.parent],
                    cache_root=tmp_path, image_pages="9")


def test_extract_image_pages_rejects_more_pages_than_the_cap(text_pdf, tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "MAX_IMAGE_PAGES", 1)
    with pytest.raises(docs.InvalidPageRange) as e:
        docs.extract(str(text_pdf), roots=[text_pdf.parent],
                    cache_root=tmp_path, image_pages="1-2")
    msg = str(e.value)
    # Specific enough to fail if the cap or the requested-page count regress
    # (the old "1" in str(e.value) assertion would pass on almost any message).
    assert "cannot render 2 pages" in msg
    assert "limit is 1" in msg


def test_extract_image_pages_on_a_conversion_cache_hit_still_renders(text_pdf, tmp_path):
    """A second extract() call on the same source, now WITH image_pages, must
    be a conversion cache hit (same doc_id, same preview — convert() is not
    re-run) that still does the rendering work: image_pages is additive to an
    existing conversion, not part of what makes a call a cache hit."""
    first = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert first["page_images"] == []

    second = docs.extract(str(text_pdf), roots=[text_pdf.parent],
                          cache_root=tmp_path, image_pages="1")
    assert second["doc_id"] == first["doc_id"]
    assert second["preview"] == first["preview"]
    assert len(second["page_images"]) == 1


def test_extract_receipt_shape_is_stable_regardless_of_image_pages(text_pdf, tmp_path):
    """The receipt's key set must not change shape based on whether
    image_pages was given — only `page_images`'s CONTENTS (and `next`'s
    wording) differ."""
    without = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    with_pages = docs.extract(str(text_pdf), roots=[text_pdf.parent],
                              cache_root=tmp_path, image_pages="1")
    expected_keys = {"doc_id", "kind", "page_count", "preview", "complete",
                     "page_images", "next"}
    assert set(without) == expected_keys
    assert set(with_pages) == expected_keys


# --- parse_pages() ---------------------------------------------------------

@pytest.mark.parametrize("spec,expected", [
    ("1", [1]), ("1-3", [1, 2, 3]), ("1-2,5", [1, 2, 5]), ("3,1", [1, 3]),
])
def test_parse_pages(spec, expected):
    assert docs.parse_pages(spec, page_count=10) == expected


def test_parse_pages_rejects_garbage():
    with pytest.raises(docs.InvalidPageRange):
        docs.parse_pages("one", page_count=10)


def test_parse_pages_rejects_reversed_range():
    with pytest.raises(docs.InvalidPageRange, match="reversed"):
        docs.parse_pages("3-1", page_count=3)


def test_parse_pages_rejects_negative_component():
    with pytest.raises(docs.InvalidPageRange, match="starting at 1"):
        docs.parse_pages("1--3", page_count=10)


def test_parse_pages_on_zero_page_document_names_the_document_not_a_range():
    with pytest.raises(docs.InvalidPageRange, match="no pages"):
        docs.parse_pages("1", page_count=0)


# --- MCP layer --------------------------------------------------------

def test_mcp_tools_are_registered():
    """hasattr() would still pass if @mcp.tool() were removed (the plain
    function is still an attribute of the module) — assert real registration
    with the server instead."""
    from es import mcp_server
    names = {t.name for t in asyncio.run(mcp_server.mcp.list_tools())}
    assert "es_doc_extract" in names
    assert "es_doc_render" not in names


def test_mcp_extract_returns_envelope_on_bad_path(monkeypatch, tmp_path):
    from es import mcp_server
    monkeypatch.setattr(mcp_server, "_doc_roots", lambda: [tmp_path])
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: tmp_path)
    out = mcp_server.es_doc_extract("/etc/passwd")
    assert out["ok"] is False
    assert out["error"]["code"] in ("doc_not_found", "doc_forbidden")


def test_mcp_extract_returns_envelope_on_success(text_pdf, monkeypatch, tmp_path):
    from es import mcp_server
    monkeypatch.setattr(mcp_server, "_doc_roots", lambda: [text_pdf.parent])
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: tmp_path)
    out = mcp_server.es_doc_extract(str(text_pdf))
    assert out["ok"] is True
    assert "Fall Season Schedule" in out["data"]["preview"]


def test_mcp_extract_image_pages_returns_envelope_on_success(text_pdf, monkeypatch, tmp_path):
    from es import mcp_server
    monkeypatch.setattr(mcp_server, "_doc_roots", lambda: [text_pdf.parent])
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: tmp_path)
    out = mcp_server.es_doc_extract(str(text_pdf), image_pages="1")
    assert out["ok"] is True
    assert len(out["data"]["page_images"]) == 1


def test_mcp_extract_image_pages_returns_envelope_on_bad_path(monkeypatch, tmp_path):
    from es import mcp_server
    monkeypatch.setattr(mcp_server, "_doc_roots", lambda: [tmp_path])
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: tmp_path)
    out = mcp_server.es_doc_extract("/etc/passwd", image_pages="1")
    assert out["ok"] is False
    assert out["error"]["code"] in ("doc_not_found", "doc_forbidden")


def test_doc_cache_root_follows_hermes_home(monkeypatch):
    from es import mcp_server
    monkeypatch.setenv("HERMES_HOME", "/somewhere/else")
    assert str(mcp_server._doc_cache_root()) == \
        "/somewhere/else/profiles/everstone/cache/documents"


def test_doc_roots_returns_media_cache_and_vault_only(tmp_path, monkeypatch):
    """_doc_roots() is the security-relevant seam deciding which directories
    are readable at all; exercise it against real config instead of
    monkeypatching it away, so a regression here (e.g. returning ["/"])
    actually fails a test."""
    from es import mcp_server

    cache_dir = tmp_path / "media-cache"
    vault_dir = tmp_path / "vault"
    cfg_path = tmp_path / "config.yaml"
    cfg_path.write_text(
        "obsidian:\n"
        "  vault_name: Vault\n"
        "  attachments:\n"
        f"    sources: [{cache_dir}]\n"
    )
    monkeypatch.setenv("ES_CONFIG_PATH", str(cfg_path))
    monkeypatch.setenv("ES_VAULT_PATH", str(vault_dir))

    assert mcp_server._doc_roots() == [str(cache_dir), str(vault_dir)]


def test_forbidden_message_names_the_remedy(text_pdf, tmp_path):
    """docs._prepare wraps paths.SourceForbidden with a document-specific
    remedy, the same way it already wraps SourceNotFound — the generic
    paths.py message alone never tells the agent what IS readable."""
    outside = tmp_path / "outside"
    outside.mkdir()
    with pytest.raises(docs.paths.SourceForbidden) as e:
        docs.extract(str(text_pdf), roots=[outside], cache_root=tmp_path)
    msg = str(e.value).lower()
    assert "uploads" in msg or "vault" in msg
    assert e.value.es_code == "doc_forbidden"


def test_forbidden_message_identical_whether_file_exists_or_not(text_pdf, tmp_path):
    """Confinement must fail the same way regardless of whether the forbidden
    path exists — otherwise the error becomes a path-probing oracle."""
    allowed = tmp_path / "allowed"
    allowed.mkdir()
    outside = tmp_path / "outside"
    outside.mkdir()

    # Both candidates sit outside the allowed root; one exists, one doesn't,
    # but they share a basename so the messages are comparable once that
    # shared name is substituted back in.
    existing = outside / "does-not-exist.pdf"
    existing.write_bytes(text_pdf.read_bytes())
    missing = outside / "does-not-exist-2.pdf"

    with pytest.raises(docs.paths.SourceForbidden) as e_exists:
        docs.extract(str(existing), roots=[allowed], cache_root=tmp_path)
    with pytest.raises(docs.paths.SourceForbidden) as e_missing:
        docs.extract(str(missing), roots=[allowed], cache_root=tmp_path)

    msg_exists = str(e_exists.value).replace(str(existing), "<PATH>")
    msg_missing = str(e_missing.value).replace(str(missing), "<PATH>")
    assert msg_exists == msg_missing
    assert "exist" not in msg_missing.lower()


# --- dispatch table ---------------------------------------------------------

def test_dispatch_table_covers_every_supported_extension():
    """SUPPORTED and the dispatch table must not drift apart — a format listed
    as supported with no converter would fail at call time, not import time."""
    from es.capabilities import docs
    assert set(docs.CONVERTERS) == docs.SUPPORTED


def test_unsupported_extension_names_the_supported_list(tmp_path):
    from es.capabilities import docs
    f = tmp_path / "x.rtf"
    f.write_text("x")
    with pytest.raises(docs.UnsupportedDocument) as e:
        docs.extract(str(f), roots=[tmp_path], cache_root=tmp_path)
    assert ".pdf" in str(e.value)


def test_extract_dispatches_a_csv_to_doc_text_with_no_page_count(csv_file, tmp_path):
    """End-to-end through docs.extract (not doc_text.convert directly): a
    flat format must report kind + doc_id like any other converter, but
    page_count stays None — that's the signal (see docs._page_count) that
    this format has no pages at all, not "zero" or "one"."""
    out = docs.extract(str(csv_file), roots=[csv_file.parent], cache_root=tmp_path)
    assert out["kind"] == "csv"
    assert out["page_count"] is None
    assert "| Name | Position | Number |" in out["preview"]


# --- non-PDF error mapping --------------------------------------------------

def test_corrupt_docx_maps_to_doc_unreadable(tmp_path):
    from es.capabilities import docs
    p = tmp_path / "bad.docx"
    p.write_bytes(b"not a zip")
    with pytest.raises(docs.UnreadableDocument):
        docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)


def test_corrupt_xlsx_maps_to_doc_unreadable(tmp_path):
    from es.capabilities import docs
    p = tmp_path / "bad.xlsx"
    p.write_bytes(b"not a zip")
    with pytest.raises(docs.UnreadableDocument):
        docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)


def test_extract_image_pages_rejects_a_non_pdf_with_a_clear_reason(csv_file, tmp_path):
    from es.capabilities import docs
    with pytest.raises(docs.UnsupportedDocument) as e:
        docs.extract(str(csv_file), roots=[csv_file.parent], cache_root=tmp_path,
                     image_pages="1")
    assert "pdf" in str(e.value).lower()
    assert "image_pages" in str(e.value)


def test_every_format_returns_the_stable_shape(
        csv_file, json_file, txt_file, ics_file, docx_file, xlsx_file, text_pdf, tmp_path):
    from es.capabilities import docs
    expected = {"doc_id", "kind", "page_count", "preview", "complete",
                "page_images", "next"}
    for f in (csv_file, json_file, txt_file, ics_file, docx_file, xlsx_file, text_pdf):
        out = docs.extract(str(f), roots=[f.parent], cache_root=tmp_path)
        assert set(out) == expected, f.name
        assert out["preview"].strip(), f.name
        assert out["kind"] == f.suffix.lstrip("."), f.name


def test_no_converter_leaks_a_raw_library_exception(tmp_path):
    """Coarse fuzz pass: every supported extension, fed garbage bytes, must
    produce an es_code from our catalogue — never a library's exception
    class name. A converter that happens to SUCCEED on this input also
    passes trivially; see
    test_realistic_malformed_documents_do_not_leak_a_raw_library_exception
    below for inputs that are guaranteed to actually fail, built from real
    library behavior rather than 30 bytes of noise."""
    from es.capabilities import docs
    allowed = {"doc_unreadable", "doc_encrypted", "doc_unsupported"}
    for ext in sorted(docs.SUPPORTED):
        p = tmp_path / f"garbage{ext}"
        p.write_bytes(b"\x00\x01\x02 not a real document \xff\xfe")
        try:
            docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
        except Exception as e:
            code = getattr(e, "es_code", type(e).__name__)
            assert code in allowed, f"{ext} leaked {code}: {e}"


# --- realistic malformed inputs must not leak a raw library exception -------
# (item 3) Each case here is a real, verified-empirically failure mode of
# python-docx/openpyxl/csv/json against an ordinary-mistake-shaped input (a
# renamed file, a partial download) — not adversarial fuzzing. Every one of
# these previously surfaced to the agent as the literal library exception
# class name (OSError, ValueError, ParseError, XMLSyntaxError, "Error",
# RecursionError, KeyError, AttributeError).

def _rezip_replacing(src_path, dest_path, filename, new_content):
    """Copy the zip at `src_path` to `dest_path`, replacing one member's
    bytes. Used to build a real .docx/.xlsx with one internal XML part
    corrupted/truncated/removed, which is what a partial download or a
    renamed-extension file actually looks like on disk — not something a
    from-scratch synthetic fixture can approximate."""
    import zipfile
    with zipfile.ZipFile(src_path) as zin, zipfile.ZipFile(dest_path, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == filename:
                if new_content is None:
                    continue  # drop this member entirely
                data = new_content
            zout.writestr(item, data)
    return dest_path


def test_realistic_malformed_documents_do_not_leak_a_raw_library_exception(tmp_path):
    from docx import Document
    from openpyxl import Workbook

    from es.capabilities import docs

    real_docx = tmp_path / "_real.docx"
    d = Document()
    d.add_paragraph("hello world, this is a real paragraph of real text")
    d.save(str(real_docx))

    real_xlsx = tmp_path / "_real.xlsx"
    wb = Workbook()
    wb.active.append(["a", "b"])
    wb.save(str(real_xlsx))

    cases = {}

    # 1. real .docx renamed .xlsx -> OSError (openpyxl)
    p = tmp_path / "docx_as_xlsx.xlsx"
    p.write_bytes(real_docx.read_bytes())
    cases["docx renamed .xlsx"] = p

    # 2. real .xlsx renamed .docx -> ValueError (python-docx)
    p = tmp_path / "xlsx_as_docx.docx"
    p.write_bytes(real_xlsx.read_bytes())
    cases["xlsx renamed .docx"] = p

    # 3. .xlsx with truncated sheet xml -> xml.etree.ElementTree.ParseError
    p = tmp_path / "truncated_sheet.xlsx"
    import zipfile
    with zipfile.ZipFile(real_xlsx) as zin:
        sheet = zin.read("xl/worksheets/sheet1.xml")
    _rezip_replacing(real_xlsx, p, "xl/worksheets/sheet1.xml", sheet[: len(sheet) // 2])
    cases[".xlsx truncated sheet xml"] = p

    # 4. .docx with truncated document.xml -> lxml.etree.XMLSyntaxError
    p = tmp_path / "truncated_document.docx"
    with zipfile.ZipFile(real_docx) as zin:
        document_xml = zin.read("word/document.xml")
    _rezip_replacing(real_docx, p, "word/document.xml", document_xml[: len(document_xml) // 2])
    cases[".docx truncated document.xml"] = p

    # 5. .csv with a field > the (raised) field-size limit
    p = tmp_path / "long_field.csv"
    p.write_text("a,b\n" + ("x" * (docs.CSV_FIELD_SIZE_LIMIT + 1000)) + ",y\n", encoding="utf-8")
    cases[".csv field over the size limit"] = p

    # 6. .csv with one unbalanced quote that swallows the rest of a large file
    #    into a single field past the size limit — same underlying csv.Error
    #    as case 5, different root cause.
    p = tmp_path / "unbalanced_quote.csv"
    p.write_text('a,b\n"' + ("y" * (docs.CSV_FIELD_SIZE_LIMIT + 1000)) + "\n", encoding="utf-8")
    cases[".csv unbalanced quote"] = p

    # 7. .json nested deeper than ~1000 levels -> RecursionError
    p = tmp_path / "deep.json"
    p.write_text("[" * 3000 + "]" * 3000, encoding="utf-8")
    cases[".json nested too deeply"] = p

    # 8. valid zip, no [Content_Types].xml -> KeyError
    p = tmp_path / "no_content_types.xlsx"
    _rezip_replacing(real_xlsx, p, "[Content_Types].xml", None)
    cases["valid zip, no [Content_Types].xml"] = p

    # 9. valid zip, [Content_Types].xml with no default namespace (python-docx
    #    never upgrades it to its own CT_Types wrapper class, so the next
    #    attribute access on it raises) -> AttributeError
    p = tmp_path / "empty_content_types.docx"
    _rezip_replacing(real_docx, p, "[Content_Types].xml", b"<?xml version='1.0'?><Types></Types>")
    cases["docx with a namespace-less [Content_Types].xml"] = p

    allowed = {"doc_unreadable", "doc_encrypted"}
    leaked_class_names = {
        "OSError", "ValueError", "ParseError", "XMLSyntaxError", "Error",
        "RecursionError", "KeyError", "AttributeError",
    }
    for label, path in cases.items():
        with pytest.raises(Exception) as e:
            docs.extract(str(path), roots=[tmp_path], cache_root=tmp_path)
        code = getattr(e.value, "es_code", type(e.value).__name__)
        assert code not in leaked_class_names, f"{label}: leaked raw {code}: {e.value}"
        assert code in allowed, f"{label}: unexpected es_code {code}: {e.value}"


# --- encrypted .docx / .xlsx -------------------------------------------------
# Neither python-docx nor openpyxl exposes a distinct "needs a password"
# exception type — a real password-protected .docx/.xlsx is stored as an
# OLE2/CFBF container (the same legacy container .doc/.xls used), never as a
# zip, so both libraries just fail to open it as a zip — indistinguishable
# by exception type alone from ordinary corruption (verified empirically
# against real python-docx/openpyxl behavior while building this feature).
# There is no dependency in this project (e.g. msoffcrypto-tool) to build a
# byte-perfect real encrypted Office document for a fixture, so these use the
# OLE2 magic-byte header alone — the same synthetic-bytes style already used
# by corrupt_pdf/empty_pdf above — which is sufficient to prove the detector
# actually fires, since those magic bytes are the entire signal it uses.

_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def test_ole2_docx_is_reported_as_encrypted_not_generic_corruption(tmp_path):
    from es.capabilities import docs
    p = tmp_path / "protected.docx"
    p.write_bytes(_OLE2_MAGIC + b"\x00" * 500)
    with pytest.raises(docs.EncryptedDocument) as e:
        docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert "password" in str(e.value).lower()


def test_ole2_xlsx_is_reported_as_encrypted_not_generic_corruption(tmp_path):
    from es.capabilities import docs
    p = tmp_path / "protected.xlsx"
    p.write_bytes(_OLE2_MAGIC + b"\x00" * 500)
    with pytest.raises(docs.EncryptedDocument) as e:
        docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert "password" in str(e.value).lower()


# --- a bug in OUR OWN converter code must surface as itself, never get -----
# --- relabeled "corrupt file" (the whole point of this refactor) -----------
#
# _CONVERSION_ERRORS used to catch ValueError/KeyError/AttributeError over
# the ENTIRE mod.convert() call — not just the library's own open/parse
# step — so a real bug in doc_office.py/doc_text.py's own rendering logic
# that happened to raise one of those ordinary types would previously have
# been swallowed and misreported as doc_unreadable. These are regression
# guards for the fix: each monkeypatches a piece of a converter's OWN
# rendering logic (never the open/parse call itself) to raise, and asserts
# the raw exception propagates unchanged — proving docs.py's parse-error
# catch is scoped to the parse step, not the whole conversion.

def test_a_bug_in_doc_office_docx_rendering_is_not_masked_as_corrupt(
        docx_file, tmp_path, monkeypatch):
    from es.capabilities import docs, doc_office

    def boom(*args, **kwargs):
        raise ValueError("simulated bug in doc_office rendering")

    monkeypatch.setattr(doc_office, "_paragraph_block", boom)

    with pytest.raises(ValueError, match="simulated bug in doc_office rendering"):
        docs.extract(str(docx_file), roots=[docx_file.parent], cache_root=tmp_path)


def test_a_bug_in_doc_office_xlsx_rendering_is_not_masked_as_corrupt(
        xlsx_file, tmp_path, monkeypatch):
    from es.capabilities import docs, doc_office

    def boom(*args, **kwargs):
        raise ValueError("simulated bug in doc_office rendering")

    monkeypatch.setattr(doc_office, "_render_sheet_rows", boom)

    with pytest.raises(ValueError, match="simulated bug in doc_office rendering"):
        docs.extract(str(xlsx_file), roots=[xlsx_file.parent], cache_root=tmp_path)


def test_a_bug_in_doc_text_csv_rendering_is_not_masked_as_corrupt(
        csv_file, tmp_path, monkeypatch):
    from es.capabilities import docs, doc_text

    def boom(cells, width):
        raise ValueError("simulated bug in doc_text rendering")

    monkeypatch.setattr(doc_text, "format_row", boom)

    with pytest.raises(ValueError, match="simulated bug in doc_text rendering"):
        docs.extract(str(csv_file), roots=[csv_file.parent], cache_root=tmp_path)


def test_a_bug_in_doc_text_json_rendering_is_not_masked_as_corrupt(
        json_file, tmp_path, monkeypatch):
    """Same guard, but through the JSON handler's own truncation logic
    (rather than the parse step, which is json.loads/json.dumps — both
    inside the tight ParseFailed boundary, see doc_text._convert_json)."""
    from es.capabilities import docs, doc_text

    def boom(text, limit):
        raise ValueError("simulated bug in doc_text rendering")

    monkeypatch.setattr(doc_text, "_truncate_at_line_boundary", boom)

    with pytest.raises(ValueError, match="simulated bug in doc_text rendering"):
        docs.extract(str(json_file), roots=[json_file.parent], cache_root=tmp_path)


# --- cross-format cache collision (doc_id must include the format) ---------

def test_cross_format_cache_collision_is_fixed(text_pdf, tmp_path):
    """The exact repro from the review: a real PDF's bytes, saved once under
    a .csv extension and once under a .pdf extension, must NOT share a
    doc_id/artifact — each is read as its OWN format (a .csv reading of PDF
    bytes is garbage-but-real CSV output, not the PDF's actual content), so
    unifying their cache entries would silently hand back whichever format
    was converted first for the full 24h TTL."""
    from es.capabilities import docs
    content = text_pdf.read_bytes()
    as_csv = tmp_path / "same.csv"
    as_pdf = tmp_path / "same.pdf"
    as_csv.write_bytes(content)
    as_pdf.write_bytes(content)

    csv_out = docs.extract(str(as_csv), roots=[tmp_path], cache_root=tmp_path)
    pdf_out = docs.extract(str(as_pdf), roots=[tmp_path], cache_root=tmp_path)

    assert csv_out["doc_id"] != pdf_out["doc_id"]
    assert csv_out["kind"] == "csv"
    assert pdf_out["kind"] == "pdf"
    assert pdf_out["page_count"] == 2
    assert "Fall Season Schedule" in pdf_out["preview"]
    assert "Fall Season Schedule" not in csv_out["preview"]

    # Each format landed in its own artifact directory, keyed by its own id.
    assert (tmp_path / ".es" / csv_out["doc_id"] / "doc.md").is_file()
    assert (tmp_path / ".es" / pdf_out["doc_id"] / "doc.md").is_file()

    # Order independence: read .pdf first, THEN .csv, and re-check both are
    # still correct — the bug reproduced in either order.
    other_root = tmp_path / "reversed"
    other_root.mkdir()
    as_pdf2 = other_root / "same.pdf"
    as_csv2 = other_root / "same.csv"
    as_pdf2.write_bytes(content)
    as_csv2.write_bytes(content)
    pdf_out2 = docs.extract(str(as_pdf2), roots=[other_root], cache_root=tmp_path)
    csv_out2 = docs.extract(str(as_csv2), roots=[other_root], cache_root=tmp_path)
    assert pdf_out2["kind"] == "pdf" and "Fall Season Schedule" in pdf_out2["preview"]
    assert csv_out2["kind"] == "csv" and "Fall Season Schedule" not in csv_out2["preview"]


def test_cross_format_cache_collision_is_fixed_for_zero_byte_files(tmp_path):
    """Reproduced in the review specifically for zero-byte files too — an
    empty .csv and an empty .txt must not collide either."""
    from es.capabilities import docs
    as_csv = tmp_path / "empty.csv"
    as_txt = tmp_path / "empty.txt"
    as_csv.write_bytes(b"")
    as_txt.write_bytes(b"")

    csv_out = docs.extract(str(as_csv), roots=[tmp_path], cache_root=tmp_path)
    txt_out = docs.extract(str(as_txt), roots=[tmp_path], cache_root=tmp_path)
    assert csv_out["doc_id"] != txt_out["doc_id"]


def test_extract_same_extension_repeat_is_still_a_cache_hit(csv_file, tmp_path, monkeypatch):
    """Guard against the format-aware doc_id fix regressing the ORIGINAL
    cache-hit property for the common case: the same file, same extension,
    extracted twice must still convert only once."""
    from es.capabilities import docs
    calls = []
    original = docs.doc_text.convert

    def spy(*args, **kwargs):
        calls.append((args, kwargs))
        return original(*args, **kwargs)

    monkeypatch.setattr(docs.doc_text, "convert", spy)

    first = docs.extract(str(csv_file), roots=[csv_file.parent], cache_root=tmp_path)
    second = docs.extract(str(csv_file), roots=[csv_file.parent], cache_root=tmp_path)

    assert len(calls) == 1
    assert first["doc_id"] == second["doc_id"]
    assert first["preview"] == second["preview"]


# --- extract() is a receipt, not the document (Task 1) ---------------------

def test_extract_returns_a_receipt_not_the_document(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert set(out) == {"doc_id", "kind", "page_count", "preview", "complete",
                        "page_images", "next"}
    assert "markdown" not in out and "truncated" not in out and "images" not in out


def test_preview_is_capped(tmp_path):
    p = tmp_path / "long.txt"
    p.write_text("x" * 5000, encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert len(out["preview"]) <= docs.PREVIEW_CHARS
    assert out["complete"] is False


def test_preview_chars_is_pinned_at_800():
    """The docstring on PREVIEW_CHARS promises the agent "the first ~800
    characters" (see mcp_server.es_doc_extract's own docstring) — both
    boundary tests here are written RELATIVE to the constant, so nothing
    else in the suite would notice if it silently drifted (e.g. 800 -> 200).
    Pin the literal directly."""
    assert docs.PREVIEW_CHARS == 800


def test_preview_never_splits_a_markdown_image_link(tmp_path):
    """The now-deleted _safe_hard_cut existed exactly to stop a truncation
    cut from landing inside a "![page N](path)" link — that property moved
    to `preview` (the only place extract() still cuts text) but nothing
    guarded it there. Reproduced empirically: every scanned PDF of 7+ pages
    ends an 800-char raw slice mid-path once the cache path is
    production-length. A short tmp_path-rooted cache_root is NOT long
    enough to reproduce this (each "![page N](...)" link is too short to
    straddle the 800-char boundary) — this test deliberately nests the
    cache_root under a realistic prefix
    (.../hermes/profiles/everstone/cache/documents/.es/<id>/pNNN.png) to
    match cella's actual on-disk path depth."""
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    png = tmp_path / "scan.png"
    Image.new("RGB", (400, 300), (200, 200, 200)).save(png)
    pdf = tmp_path / "scanned9.pdf"
    c = canvas.Canvas(str(pdf), pagesize=letter)
    for _ in range(9):
        c.drawImage(str(png), 72, 400, width=400, height=300)
        c.showPage()
    c.save()

    cache_root = (tmp_path / "opt" / "data" / "hermes" / "profiles" /
                  "everstone" / "cache" / "documents")
    cache_root.mkdir(parents=True)

    out = docs.extract(str(pdf), roots=[tmp_path], cache_root=cache_root)
    assert out["complete"] is False  # otherwise the cut never fires at all
    assert not re.search(r"!\[[^\]]*\]\([^)]*$", out["preview"]), \
        "preview ended inside an unterminated markdown image link"


def test_complete_is_true_when_the_preview_is_the_whole_document(tmp_path):
    p = tmp_path / "short.txt"
    p.write_text("Practice moved to Thursday.\n", encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert out["complete"] is True
    assert "Practice moved to Thursday." in out["preview"]


def test_complete_is_exact_at_the_boundary(tmp_path):
    """complete must be an exact test (len(markdown) <= PREVIEW_CHARS), not a
    heuristic — that exactness is why preview is a character count rather than
    'the first section'."""
    for delta, expected in ((0, True), (1, False)):
        p = tmp_path / f"b{delta}.txt"
        p.write_text("y" * (docs.PREVIEW_CHARS + delta), encoding="utf-8")
        out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
        assert out["complete"] is expected, delta


def test_next_names_the_tool_and_the_handle(text_pdf, tmp_path):
    out = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=tmp_path)
    assert "es_read" in out["next"]
    assert out["doc_id"] in out["next"], "the agent should copy the handle, not build it"


def test_next_names_the_handle_on_the_incomplete_branch_too(tmp_path):
    """text_pdf (~130 chars) only ever exercises the complete=True branch of
    `next` — every long document (the common case) takes the OTHER branch,
    which was untested: a mutation dropping the handle there left every
    existing test passing."""
    p = tmp_path / "long.txt"
    p.write_text("x" * 5000, encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert out["complete"] is False
    assert "es_read" in out["next"]
    assert out["doc_id"] in out["next"], "the agent should copy the handle, not build it"


def test_the_full_markdown_is_still_cached(tmp_path):
    """The receipt is small; doc.md still caches the FULL document, unaffected
    by the preview cap — es_read pages that full cached copy.

    Uses a document whose full markdown genuinely exceeds PREVIEW_CHARS
    (rather than the small `text_pdf` fixture, whose ~130-character output is
    itself under PREVIEW_CHARS — `complete` would be True and preview would
    equal the cached copy exactly, making `len(cached) > len(preview)`
    impossible to satisfy by construction, not a real assertion)."""
    p = tmp_path / "long.txt"
    p.write_text("Fall Season Schedule\n" + ("x" * 5000), encoding="utf-8")
    out = docs.extract(str(p), roots=[tmp_path], cache_root=tmp_path)
    assert out["complete"] is False
    cached = (tmp_path / ".es" / out["doc_id"] / "doc.md").read_text(encoding="utf-8")
    assert len(cached) > len(out["preview"])
    assert cached.startswith(out["preview"][:50])


def test_es_read_still_reaches_the_whole_document_after_extract(text_pdf, tmp_path,
                                                               monkeypatch):
    """The contract change must not cost reach: everything the old dump returned
    is still retrievable through es_read."""
    from es import mcp_server
    monkeypatch.setattr(mcp_server, "_doc_roots", lambda: [text_pdf.parent])
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: tmp_path)
    did = mcp_server.es_doc_extract(str(text_pdf))["data"]["doc_id"]
    r = mcp_server.es_read("doc:" + did)["data"]
    assert "Fall Season Schedule" in (r["content"] or "") or r["outline"]
