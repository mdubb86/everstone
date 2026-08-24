"""es.capabilities.reader.resolve — turning an es_read `target` into
Markdown, for both vault notes and cached document conversions.

Exercises reader.resolve() directly against a real VaultClient/cache_root
(no MCP layer, no config.py) — the signature is explicit/injectable on
purpose so these tests don't need a container.
"""
import os
import time
from typing import List

import pytest

from es import doc_cache, mcp_server
from es.capabilities import doc_support, docs, reader
from es.vault_client import NoteNotFound, VaultClient


@pytest.fixture
def vault(tmp_path):
    root = tmp_path / "vault"
    root.mkdir()
    return VaultClient(root, "TestVault", categories=["Topics"])


@pytest.fixture
def cache_root(tmp_path):
    d = tmp_path / "cache"
    d.mkdir()
    return d


# --- vault notes -------------------------------------------------------

def test_resolve_note_by_path_returns_markdown_body(vault):
    vault.write_topic("Manual", body="# Manual\n\nRead me.")
    out = reader.resolve("Topics/Manual.md", vault=vault, cache_root=None)
    assert out["kind"] == "note"
    assert "Read me." in out["markdown"]
    assert out["path"] == "Topics/Manual.md"
    assert out["source"] == "Topics/Manual.md"


def test_resolve_note_by_topic_name_matches_by_path_resolution(vault):
    """Preserves the retired es_notes_read's target semantics exactly: a bare
    topic name resolves the same note as its vault-relative path."""
    vault.write_topic("Manual", body="# Manual\n\nRead me.")
    by_path = reader.resolve("Topics/Manual.md", vault=vault, cache_root=None)
    by_topic = reader.resolve("Manual", vault=vault, cache_root=None)
    assert by_topic["markdown"] == by_path["markdown"]
    assert by_topic["path"] == by_path["path"]
    assert by_topic["source"] == by_path["source"]


def test_resolve_note_returns_frontmatter(vault):
    """The retired es_notes_read returned {path, frontmatter, body} and the
    agent relies on frontmatter (topics, tags, created) — dropping it would
    be a regression now that es_read is the only read path for notes."""
    vault.write_journal("Practice moved", "Body text.", tags=["soccer"],
                        topics=["Thunder U10"])
    entries = vault.list_journal()
    assert len(entries) == 1
    out = reader.resolve(entries[0]["path"], vault=vault, cache_root=None)
    assert out["frontmatter"].get("tags") == ["soccer"]
    assert out["frontmatter"].get("topics") == ["[[Thunder U10]]"]
    assert "created" in out["frontmatter"]


def test_resolve_missing_note_path_raises_unchanged_not_found(vault):
    with pytest.raises(NoteNotFound):
        reader.resolve("Topics/Nope.md", vault=vault, cache_root=None)


# --- cached documents ("doc:<id>") -------------------------------------

def test_resolve_doc_handle_returns_cached_markdown(text_pdf, cache_root):
    """extract() itself only returns a receipt (preview, not the document —
    see docs.extract's new contract), so the full-markdown comparison below
    checks that `preview` is a prefix of what reader.resolve() (which reads
    the cached doc.md directly) returns, rather than comparing two full
    markdown strings extract() no longer produces."""
    extracted = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=cache_root)
    target = f"doc:{extracted['doc_id']}"

    out = reader.resolve(target, vault=None, cache_root=cache_root)
    assert out["kind"] == "doc"
    assert out["doc_id"] == extracted["doc_id"]
    assert out["source"] == target
    assert out["markdown"].startswith(extracted["preview"])
    assert "Fall Season Schedule" in out["markdown"]


def test_resolve_unknown_doc_handle_raises_expired_naming_the_remedy(cache_root):
    with pytest.raises(reader.DocHandleExpired) as e:
        reader.resolve("doc:deadbeef0000", vault=None, cache_root=cache_root)
    assert e.value.es_code == "doc_handle_expired"
    assert "es_doc_extract" in str(e.value)


def test_resolve_doc_handle_touches_the_artifact_dir(text_pdf, cache_root):
    """The cache TTL is 24h since last ACCESS (doc_cache.touch) — reading a
    doc: handle through es_read must count as an access, the same way a
    cache-hit inside es_doc_extract already does. Regression pattern mirrors
    test_touch_makes_a_stale_directory_fresh in test_doc_cache.py."""
    extracted = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=cache_root)
    doc_id = extracted["doc_id"]
    adir = doc_cache.artifact_dir(cache_root, doc_id)

    stale = time.time() - (25 * 3600)
    os.utime(adir, (stale, stale))

    reader.resolve(f"doc:{doc_id}", vault=None, cache_root=cache_root)

    assert adir.stat().st_mtime > stale
    assert doc_cache.purge(cache_root) == 0  # would have been evicted had
                                              # the read not touched it


def test_resolve_doc_handle_survives_a_missing_images_sidecar(text_pdf, cache_root):
    """A partially-purged artifact dir (images.json gone, doc.md intact)
    must still resolve — reader.py must inherit docs.py's existing
    missing-sidecar tolerance (_read_images_manifest) rather than working
    around it a second time."""
    extracted = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=cache_root)
    adir = doc_cache.artifact_dir(cache_root, extracted["doc_id"])
    (adir / docs.DOC_IMAGES_MANIFEST).unlink()

    out = reader.resolve(f"doc:{extracted['doc_id']}", vault=None, cache_root=cache_root)
    assert out["markdown"].startswith(extracted["preview"])


# --- table-kind handles are refused, not silently read as Markdown ------
#
# doc_table produces this kind for every .csv/.xlsx, so these tests drive a
# real conversion. The rejection matters more than it looks: a table artifact
# has no doc.md at all, so without it es_read would report a perfectly good
# spreadsheet as an expired handle.

def _seed_table_doc(cache_root, csv_text: str = "col_a,col_b\n1,2\n") -> str:
    """A REAL table document, converted by the real converter.

    Was a fabricated artifact (docs._write_full_extract(..., kind="table"))
    back when no converter produced the kind. Now one does, and a fabricated
    one would no longer even be found: a table artifact has no doc.md at all,
    so read_cached() gates it on tables.json + the database — exactly the
    kind of divergence between fixture and reality that makes a passing test
    meaningless."""
    src = cache_root / f"seed-table-{len(csv_text)}.csv"
    src.write_text(csv_text, encoding="utf-8")
    out = docs.extract(str(src), roots=[cache_root], cache_root=cache_root)
    assert out["kind"] == "table", "fixture must actually produce a table document"
    return out["doc_id"]


def test_resolve_table_kind_handle_raises_naming_es_doc_query(cache_root):
    did = _seed_table_doc(cache_root)
    with pytest.raises(reader.TableKindNotReadable) as e:
        reader.resolve(f"doc:{did}", vault=None, cache_root=cache_root)
    assert e.value.es_code == "doc_table_kind"
    assert "es_doc_query" in str(e.value)
    assert f"doc:{did}" in str(e.value)


def test_resolve_table_kind_handle_still_touches_the_artifact_dir(cache_root):
    """Even though the read is refused, the agent did just look this handle
    up — the same "a lookup is a use" reasoning
    test_resolve_doc_handle_touches_the_artifact_dir already applies to a
    successful read applies here too, so the artifact should not expire
    while the agent is actively (if unsuccessfully) trying to use it."""
    did = _seed_table_doc(cache_root)
    adir = doc_cache.artifact_dir(cache_root, did)
    stale = time.time() - (25 * 3600)
    os.utime(adir, (stale, stale))

    with pytest.raises(reader.TableKindNotReadable):
        reader.resolve(f"doc:{did}", vault=None, cache_root=cache_root)

    assert adir.stat().st_mtime > stale


def test_resolve_markdown_kind_handle_is_unaffected_by_the_table_guard(text_pdf, cache_root):
    """The guard must not fire on the common case: an ordinary (markdown-
    kind) extract still resolves exactly as before."""
    extracted = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=cache_root)
    out = reader.resolve(f"doc:{extracted['doc_id']}", vault=None, cache_root=cache_root)
    assert out["kind"] == "doc"
    assert "Fall Season Schedule" in out["markdown"]


# --- target cannot escape the vault or the cache ------------------------

def test_doc_handle_traversal_is_rejected_as_expired_not_a_path_error(cache_root):
    """A doc: id is never joined onto a path unless it's pure hex — a
    traversal attempt fails the hex check and comes back as the same
    DocHandleExpired an unknown id would, never touching the filesystem
    outside cache_root."""
    with pytest.raises(reader.DocHandleExpired):
        reader.resolve("doc:../../../etc/passwd", vault=None, cache_root=cache_root)


def test_doc_handle_hex_guard_is_what_rejects_traversal(tmp_path, cache_root):
    """test_doc_handle_traversal_is_rejected_as_expired_not_a_path_error
    (above) passes for a second reason that has nothing to do with the hex
    guard: `doc:../../../etc/passwd` also fails because no doc.md happens to
    live at that resolved path — so that test alone would keep passing even
    if reader._DOC_ID_RE were weakened to `^.+$`, silently reopening the
    escape. This test plants a REAL doc.md at a location a traversal would
    land on if the hex check ever let a non-hex id through, proving the
    regex itself — not luck — is what blocks it.
    """
    escaped = tmp_path / "escaped"
    escaped.mkdir()
    docs._write_full_extract(escaped, "# Secret\n\nShould never be reachable.", [])

    # A doc_id crafted so that cache_root/ES_NAMESPACE/<doc_id> resolves to
    # `escaped` once the OS follows the ".." components — exactly what an
    # UNGUARDED path join would do. ES_NAMESPACE itself must actually exist
    # on disk (real _resolve_doc always joins it in, real doc_id or not) —
    # the kernel needs every intermediate component, including this one, to
    # be a real directory before it can walk back out of it via "..".
    ns_dir = cache_root / doc_cache.ES_NAMESPACE
    ns_dir.mkdir(parents=True, exist_ok=True)
    traversal_id = os.path.relpath(escaped, ns_dir)
    assert not reader._DOC_ID_RE.match(traversal_id)  # contains '/' and '..'

    with pytest.raises(reader.DocHandleExpired):
        reader.resolve(f"doc:{traversal_id}", vault=None, cache_root=cache_root)

    # Prove the traversal would genuinely have worked without the guard:
    # joining the same id onto the cache root directly (bypassing resolve's
    # regex) does find the planted file.
    unguarded_path = ns_dir / traversal_id
    assert docs.read_cached(unguarded_path) is not None


def test_note_path_traversal_is_rejected_unchanged(vault):
    """Vault confinement (VaultClient._within_root) is untouched by this
    module — a traversal attempt still surfaces as NoteNotFound, exactly as
    the retired es_notes_read already behaved."""
    with pytest.raises(NoteNotFound):
        reader.resolve("../../../../etc/passwd", vault=vault, cache_root=None)


# --- the es_read MCP tool -----------------------------------------------
#
# These exercise es_read end to end (through mcp_server, through
# reader.resolve, through the read.py primitives) rather than any one layer
# in isolation — the point of the tool is how those layers compose, and the
# stable-envelope / whole-vs-outline behaviour only exists at this level.

ENVELOPE_KEYS = {"kind", "source", "path", "frontmatter", "content",
                 "outline", "more", "next_offset"}


@pytest.fixture
def wired_vault(tmp_path, monkeypatch):
    root = tmp_path / "vault"
    root.mkdir()
    v = VaultClient(root, "TestVault", categories=["Topics"])
    monkeypatch.setattr(mcp_server, "_notes_client", lambda: v)
    return v


@pytest.fixture
def wired_cache(tmp_path, monkeypatch):
    d = tmp_path / "cache"
    d.mkdir()
    monkeypatch.setattr(mcp_server, "_doc_cache_root", lambda: d)
    return d


def _big_markdown(n: int = 150) -> str:
    """A synthetic document comfortably over es_read's whole-vs-outline
    threshold, shaped like a converted calendar: a short preamble line, then
    many small `##` sections (deliberately small individually — the point of
    the 100+-event calendar case is that no SINGLE section is large)."""
    parts = [f"{n} events (0 recurring)"]
    for i in range(n):
        parts.append(f"\n## Event {i}\n\nDetails for event {i}.\n")
    return "\n".join(parts)


def _seed_doc(cache_root, markdown: str, ext: str = ".ics") -> str:
    """Write `markdown` directly into the doc cache as if es_doc_extract had
    already converted some external file — the fixture-of-choice for these
    tests' "large multi-section DOCUMENT" cases (the 117-event-calendar case
    they're modeled on is exactly a doc:<id>, never a vault note), without
    needing a real multi-hundred-event .ics file to drive through doc_ics.
    `ext` only seeds doc_cache.doc_id's hash namespace; the fake source
    file's own bytes are never read back by anything under test."""
    fake_source = cache_root / f"fake-{len(markdown)}{ext}"
    fake_source.write_bytes(os.urandom(8))
    did = doc_cache.doc_id(fake_source, ext)
    adir = doc_cache.artifact_dir(cache_root, did)
    docs._write_full_extract(adir, markdown, [])
    return f"doc:{did}"


def test_es_read_small_note_returns_whole_content(wired_vault):
    wired_vault.write_topic("Manual", body="# Manual\n\nRead me in full.")
    out = mcp_server.es_read("Manual")
    assert out["ok"] is True
    data = out["data"]
    assert data["content"].strip() == "# Manual\n\nRead me in full."
    assert data["more"] is False
    assert data["next_offset"] is None
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_large_document_returns_outline_not_everything(wired_cache):
    """The core behaviour: a large multi-section DOCUMENT (standing in for
    the 117-event-calendar case — hence a doc:<id>, not a vault note; see
    the note-vs-document threshold tests further down for why that distinction
    matters) must not dump every section — it comes back as an outline to
    choose from, plus a short preamble preview, with more=true telling the
    agent there's content it hasn't seen yet."""
    big = _big_markdown(150)
    target = _seed_doc(wired_cache, big)

    out = mcp_server.es_read(target)
    assert out["ok"] is True
    data = out["data"]
    assert len(data["outline"]) == 150
    assert data["outline"][0]["title"] == "Event 0"
    assert data["more"] is True
    assert data["content"] == "150 events (0 recurring)"
    # the full 150 event bodies must NOT be dumped into `content`
    assert "Details for event 149" not in (data["content"] or "")
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_section_returns_only_that_section(wired_cache):
    big = _big_markdown(150)
    target = _seed_doc(wired_cache, big)
    outline = mcp_server.es_read(target)["data"]["outline"]
    target_id = next(s["id"] for s in outline if s["title"] == "Event 7")

    out = mcp_server.es_read(target, section=target_id)
    assert out["ok"] is True
    data = out["data"]
    assert "Details for event 7." in data["content"]
    assert "Event 8" not in data["content"]
    assert data["more"] is True  # other sections remain
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_query_returns_matching_ids(wired_vault):
    wired_vault.write_topic(
        "Team", body="## Roster\n\nNo mention here.\n\n## Nickname\n\nThe Fury.\n")
    out = mcp_server.es_read("Team", search="fury")
    assert out["ok"] is True
    data = out["data"]
    assert [s["title"] for s in data["outline"]] == ["Nickname"]
    assert data["content"] is None  # query hands back ids to follow up with, not text
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_query_with_no_hits_explains_what_to_try(wired_vault):
    wired_vault.write_topic("Team", body="## Roster\n\nNothing relevant.\n")
    out = mcp_server.es_read("Team", search="zebra")
    assert out["ok"] is True
    data = out["data"]
    assert data["outline"] == []
    assert data["content"]  # not empty — tells the agent what to try next
    assert "zebra" in data["content"]
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_offset_on_headingless_document_reports_next_offset(wired_vault):
    # 500 rows: with the default 200-line window, offset=100 leaves rows still
    # unread past this page (100:300), so next_offset must be populated.
    csv_body = "\n".join(f"row{i},value{i}" for i in range(500))
    wired_vault.write_topic("Data", body=csv_body)

    out = mcp_server.es_read("Data", offset=100)
    assert out["ok"] is True
    data = out["data"]
    assert data["content"].splitlines()[0] == "row100,value100"
    assert data["next_offset"] == 300
    assert data["more"] is True
    assert data["outline"] is None
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_unknown_section_is_no_such_section(wired_vault):
    wired_vault.write_topic("Manual", body="## Real Section\n\nBody.\n")
    out = mcp_server.es_read("Manual", section="nope")
    assert out["ok"] is False
    assert out["error"]["code"] == "no_such_section"
    assert "real-section" in out["error"]["message"]


def test_es_read_expired_doc_handle(wired_cache):
    out = mcp_server.es_read("doc:deadbeef0000")
    assert out["ok"] is False
    assert out["error"]["code"] == "doc_handle_expired"
    assert "es_doc_extract" in out["error"]["message"]


def test_es_read_table_kind_handle_errors_not_a_null_envelope(wired_cache):
    """es_read on a table-kind handle must come back as a genuine error
    envelope (ok=False, a real es_code naming es_doc_query as the remedy) —
    never ok=True with content/outline left null, and never an empty read."""
    did = _seed_table_doc(wired_cache)
    out = mcp_server.es_read(f"doc:{did}")
    assert out["ok"] is False
    assert out["error"]["code"] == "doc_table_kind"
    assert "es_doc_query" in out["error"]["message"]


def test_es_read_note_by_path_and_topic_name_match_with_frontmatter(wired_vault):
    """A topic doc is addressable both by its vault-relative path and by its
    bare topic name (VaultClient._find_topic) — both must resolve to the
    same content, with frontmatter intact either way."""
    wired_vault.write_topic(
        "Manual", body="---\ntags: [important]\n---\n\n# Manual\n\nContent.\n")

    by_path = mcp_server.es_read("Topics/Manual.md")["data"]
    by_topic = mcp_server.es_read("Manual")["data"]

    assert by_path["content"] == by_topic["content"]
    assert by_path["path"] == by_topic["path"] == "Topics/Manual.md"
    assert by_path["frontmatter"]["tags"] == ["important"]
    assert by_topic["frontmatter"]["tags"] == ["important"]


def test_es_read_doc_handle_reads_a_converted_document(text_pdf, wired_cache):
    extracted = docs.extract(str(text_pdf), roots=[text_pdf.parent], cache_root=wired_cache)
    out = mcp_server.es_read(f"doc:{extracted['doc_id']}")
    assert out["ok"] is True
    data = out["data"]
    assert data["kind"] == "doc"
    assert data["path"] is None
    assert data["frontmatter"] is None
    assert "Fall Season Schedule" in data["content"]
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_envelope_key_set_is_identical_across_all_modes(wired_vault, wired_cache):
    """The agent must not have to reason about which keys exist depending on
    which mode it used — every mode returns exactly the same key set,
    across both kinds (note, doc) es_read can serve."""
    big = _big_markdown(150)
    big_cal = _seed_doc(wired_cache, big)
    wired_vault.write_topic("Small", body="# Small\n\nWhole thing.")
    section_id = mcp_server.es_read(big_cal)["data"]["outline"][0]["id"]

    modes = [
        mcp_server.es_read("Small"),
        mcp_server.es_read(big_cal),
        mcp_server.es_read(big_cal, section=section_id),
        mcp_server.es_read(big_cal, search="Event 3"),
        mcp_server.es_read(big_cal, search="no-such-text-anywhere"),
    ]
    for out in modes:
        assert out["ok"] is True
        assert set(out["data"].keys()) == ENVELOPE_KEYS


# --- note-vs-document whole-vs-outline threshold ------------------------
#
# A prior task flagged a gap: nothing exercised a genuinely long NOTE (as
# opposed to _big_markdown's synthetic, document-shaped fixture above) through
# the whole-vs-outline threshold — these tests close it, and pin down the
# actual decision: a note gets a much larger "whole" allowance than a
# document, because it's authored by the user (naturally bounded, and "what
# did I write about the tournament" is usually better answered directly than
# via a menu of section ids), while a document arrives from outside at
# whatever size its source format happened to be.

def _long_note_text(entries: int) -> str:
    """Models a topic/journal note a user has appended to over time: short
    dated paragraphs with NO `##` headings — the common shape of a running
    note nobody bothered to section, unlike _big_markdown's many-small-
    headings shape (which models a converted calendar, not a note). Each
    entry spans two lines (not one), so a large-enough `entries` count also
    crosses read.window's 200-*line* default page size, not just a character
    count — needed by test_es_read_document_at_the_same_size_is_still_paged
    below, which must observe actual partial content, not just a "large"
    flag with nothing behind it."""
    entry = ("Practice on day {i}: worked on passing drills and a short "
             "scrimmage.\nCoach noted improvement in first touch and "
             "off-ball movement. Weather was clear.")
    return "\n\n".join(entry.format(i=i) for i in range(entries))


def test_es_read_long_note_without_headings_still_returns_whole(wired_vault):
    """~10,500 characters — comfortably over the DOCUMENT threshold
    (_WHOLE_DOCUMENT_CHAR_LIMIT, 4,000) but under the NOTE threshold
    (_WHOLE_NOTE_CHAR_LIMIT, 16,000): a year of ordinary appended journal
    entries, not a converted document. Must come back whole — this is
    exactly the case the previous task flagged as a regression (a long note
    that used to return whole via es_notes_read now costing a second call)."""
    text = _long_note_text(70)
    assert 4_000 < len(text) < 16_000
    wired_vault.write_topic("Season Log", body=text)

    out = mcp_server.es_read("Season Log")
    assert out["ok"] is True
    data = out["data"]
    assert data["content"].strip() == text
    assert data["more"] is False
    assert data["next_offset"] is None
    assert data["outline"] is None  # no headings in this note at all
    assert set(data.keys()) == ENVELOPE_KEYS


def test_es_read_document_at_the_same_size_is_still_paged(wired_cache):
    """The SAME text that comes back whole for a note (previous test) must
    still be paged when read as a doc:<id> — proving the threshold actually
    differs by `kind`, not by some property of the content itself."""
    text = _long_note_text(70)
    target = _seed_doc(wired_cache, text, ext=".txt")

    out = mcp_server.es_read(target)
    assert out["ok"] is True
    data = out["data"]
    assert data["more"] is True
    assert data["next_offset"] is not None
    assert data["content"] != text  # only a window's worth, not the whole thing


def test_es_read_very_long_note_still_outlines(wired_vault):
    """A note big enough to exceed even the generous NOTE threshold (~19,500
    characters) must still page — the larger allowance is not "notes are
    never paged", just a higher bar."""
    text = _long_note_text(130)
    assert len(text) > 16_000
    wired_vault.write_topic("Long Season Log", body=text)

    out = mcp_server.es_read("Long Season Log")
    assert out["ok"] is True
    data = out["data"]
    assert data["more"] is True
    assert data["next_offset"] is not None
    assert data["content"] != text


# --- C1: content is bounded on every path that returns it ----------------
#
# es_read decided WHOLE-vs-OUTLINE using a threshold, but nothing capped the
# text any path actually returned — verified live against a realistic
# Obsidian topic note (150 paragraphs, one per line): note_chars=112348,
# content_chars=74899 (~18,700 tokens), unbounded. Past Hermes's own
# DEFAULT_MCP_RESULT_SIZE_CHARS house limit (50,000) a tool result is
# spilled to a file with only a preview kept in context — and this agent has
# no file tool to open the spillover. These tests pin `content` to
# mcp_server._CONTENT_CHAR_CAP on every path (preamble/first-section, an
# explicit `section`, and both `window` paths), each with an in-band marker
# and the rest still reachable via the documented `offset` escape hatch.

def _prose_lines(n: int, width: int = 500) -> str:
    """`n` headingless prose lines, each padded to `width` characters — the
    shape of an ordinary Obsidian topic note (one paragraph per line, no
    hard wrap), at whatever total size a test needs to cross a threshold."""
    filler = ("Practice went well today; the team worked on passing drills, "
              "set pieces, and a short scrimmage to close things out. ")
    line = (filler * (width // len(filler) + 1))[:width]
    return "\n".join(line for _ in range(n))


def test_es_read_long_note_preamble_is_capped_with_marker_and_remains_retrievable(wired_vault):
    """The primary regression: a realistic topic note (~112k characters, one
    paragraph per line, one heading partway through) used to return its
    ENTIRE ~75,000-character preamble verbatim. Must now be capped, marked
    in-band, and the rest still reachable by paging with `offset`."""
    preamble = _prose_lines(150, width=500)  # ~75,000 chars, no headings
    assert len(preamble) > mcp_server._CONTENT_CHAR_CAP
    body = preamble + "\n\n## Later Section\n\nNotes added after the fact.\n"
    wired_vault.write_topic("Season Log", body=body)

    out = mcp_server.es_read("Season Log")
    assert out["ok"] is True
    data = out["data"]
    assert data["more"] is True
    assert len(data["content"]) <= mcp_server._CONTENT_CHAR_CAP
    assert doc_support.TRUNCATION_SENTINEL in data["content"]
    assert "Later Section" not in data["content"]  # cut well before the heading

    # Fully retrievable: paging by line from the start reproduces the exact
    # underlying markdown (not just "some more text").
    resolved_md = reader.resolve("Season Log", vault=wired_vault, cache_root=None)["markdown"]
    collected: List[str] = []
    offset = 0
    while offset is not None:
        page = mcp_server.es_read("Season Log", offset=offset)["data"]
        collected.append(page["content"])
        offset = page["next_offset"]
    assert "\n".join(collected) == resolved_md


def test_es_read_five_megabyte_note_preamble_is_capped(wired_vault):
    """Verified live: a 5MB note whose only heading ("## Footnotes") sits at
    the very end returned its full 5,000,000-character preamble verbatim."""
    preamble = "Long-form notes about the season. " * 150_000  # ~5.25MB
    body = preamble + "\n\n## Footnotes\n\n[1] See appendix.\n"
    wired_vault.write_topic("Huge Log", body=body)

    out = mcp_server.es_read("Huge Log")
    assert out["ok"] is True
    data = out["data"]
    assert len(data["outline"]) == 1
    assert len(data["content"]) <= mcp_server._CONTENT_CHAR_CAP
    assert doc_support.TRUNCATION_SENTINEL in data["content"]
    assert data["more"] is True


def test_es_read_single_giant_line_is_capped_and_fully_pageable(wired_cache):
    """Verified live: an 800,000-character single-line document (a pasted
    transcript with no line breaks at all) came back whole (more=false).
    read_cap.window pages by LINE, so from its point of view this is
    exactly one line — no amount of line-based paging alone could return
    less than the whole thing. es_read must pre-split an oversized line
    into cap-sized synthetic lines so it's still boundable and fully
    retrievable by paging."""
    text = "0123456789" * 80_000  # 800,000 chars, one line, no headings
    target = _seed_doc(wired_cache, text, ext=".txt")

    out = mcp_server.es_read(target)
    assert out["ok"] is True
    data = out["data"]
    assert len(data["content"]) <= mcp_server._CONTENT_CHAR_CAP
    assert data["more"] is True
    assert data["next_offset"] is not None

    collected = data["content"]
    offset = data["next_offset"]
    while offset is not None:
        page = mcp_server.es_read(target, offset=offset)["data"]
        collected += page["content"]
        offset = page["next_offset"]
    assert collected == text


def test_es_read_docx_whose_only_heading_is_the_appendix_stays_within_cap(tmp_path, wired_cache):
    """Verified live: a .docx whose only heading is an appendix at the very
    end returned a 12,289-character preamble — 3x the 4,000-character
    document whole-vs-outline threshold, with nothing capping it. Pins the
    shape down as a regression guard (12,289 < the cap, so this doesn't
    itself exercise truncation — see the 5MB/112k tests above for that)."""
    from docx import Document

    d = Document()
    # Enough paragraphs to clear the 4,000-character document threshold but
    # comfortably under doc_office's own 30,000-character conversion budget
    # (MAX_CHARS) — the appendix heading must survive INTO the converted
    # markdown for this to exercise the right bug; too many paragraphs and
    # doc_office's own truncation cuts the document before ever reaching it.
    for i in range(100):
        d.add_paragraph(f"Paragraph {i}: ordinary body text about the season, "
                        "practices, and logistics that goes on for a while.")
    d.add_heading("Appendix", level=1)
    d.add_paragraph("Reference material goes here.")
    p = tmp_path / "manual.docx"
    d.save(str(p))

    extracted = docs.extract(str(p), roots=[p.parent], cache_root=wired_cache)
    out = mcp_server.es_read(f"doc:{extracted['doc_id']}")
    assert out["ok"] is True
    data = out["data"]
    assert len(data["outline"]) == 1
    assert data["outline"][0]["title"] == "Appendix"
    assert len(data["content"]) <= mcp_server._CONTENT_CHAR_CAP
    assert data["content"]


def test_es_read_section_content_is_capped_with_marker(wired_cache):
    """The top-level preamble/window paths aren't the only ones that need
    bounding — a single SECTION whose own body alone exceeds the cap (one
    outsized event description in an otherwise ordinary calendar) must be
    capped too."""
    huge_body = "x" * 50_000
    md = f"## Event 1\n\n{huge_body}\n\n## Event 2\n\nShort.\n"
    target = _seed_doc(wired_cache, md, ext=".ics")

    outline = mcp_server.es_read(target)["data"]["outline"]
    sid = next(s["id"] for s in outline if s["title"] == "Event 1")

    out = mcp_server.es_read(target, section=sid)
    assert out["ok"] is True
    data = out["data"]
    assert len(data["content"]) <= mcp_server._CONTENT_CHAR_CAP
    assert doc_support.TRUNCATION_SENTINEL in data["content"]


def test_es_read_offset_paging_visits_every_line_exactly_once(wired_cache):
    """The subtle part of bounding `window` by characters: when the char cap
    forces fewer lines per page than read_cap.window's own 200-line
    ceiling, `next_offset` must reflect what was ACTUALLY returned, or
    paging silently skips content. 400 lines of ~500 characters each
    (~200,000 total) — the character cap (32,000) binds well before the
    200-line default, so this only passes if next_offset is recomputed from
    the trimmed page rather than trusted from window()'s own idea of what
    it returned."""
    lines = [f"{i:05d} " + ("x" * 490) for i in range(400)]
    body = "\n".join(lines)
    target = _seed_doc(wired_cache, body, ext=".txt")

    collected: List[str] = []
    offset = 0
    pages = 0
    while offset is not None:
        page = mcp_server.es_read(target, offset=offset)["data"]
        assert len(page["content"]) <= mcp_server._CONTENT_CHAR_CAP
        collected.extend(page["content"].split("\n"))
        offset = page["next_offset"]
        pages += 1
        assert pages < 1000  # sanity bound against an infinite loop

    assert collected == lines
    assert pages > 1  # actually exercised more than one page


# --- item 2: content must not be null in outline mode ---------------------
#
# Verified live: a 40-page PDF and a .xlsx workbook both returned
# outline=N, content=None — the spec promised "an outline plus the first
# section", but the implementation only ever returned the PREAMBLE, which is
# empty whenever the document starts AT a heading (both formats do, always,
# for .xlsx — every sheet begins with its own "## <sheet>" heading).

def test_es_read_outline_mode_falls_back_to_first_section_when_preamble_empty(wired_cache):
    md = ("## Section A\n\n" + ("Body text for section A. " * 200) +
          "\n\n## Section B\n\n" + ("Body text for section B. " * 200))
    target = _seed_doc(wired_cache, md, ext=".txt")

    out = mcp_server.es_read(target)
    assert out["ok"] is True
    data = out["data"]
    assert len(data["outline"]) == 2
    assert data["content"] is not None
    assert "Body text for section A" in data["content"]
    assert "Section B" not in data["content"]  # first section only, not everything


def test_heading_first_document_outline_mode_returns_content_not_null(tmp_path, wired_cache):
    """A document whose text STARTS at a heading has no preamble at all, so
    outline mode's "return the preamble as a preview" path finds nothing —
    the exact shape verified live to regress to content=None.

    Written against an .xlsx originally (every sheet begins with its own
    "## <sheet>" heading). An .xlsx is a table document now with no markdown
    to page, so this uses a multi-page PDF: doc_pdf's output likewise starts
    at "## Page 1" with nothing before it. A .ics would NOT do — doc_ics
    emits an event-count preamble, so it takes the ordinary branch and never
    reaches the fallback this test exists for."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    p = tmp_path / "manual.pdf"
    c = canvas.Canvas(str(p), pagesize=letter)
    for page in range(1, 4):
        for line in range(40):
            c.drawString(72, 720 - line * 16,
                         f"Page {page} line {line}: padding text to grow this document.")
        c.showPage()
    c.save()

    extracted = docs.extract(str(p), roots=[p.parent], cache_root=wired_cache)
    out = mcp_server.es_read(f"doc:{extracted['doc_id']}")
    assert out["ok"] is True
    data = out["data"]
    assert len(data["outline"]) == 3
    assert data["content"] is not None
    assert "Page 1 line 0" in data["content"]


# --- section="page-N" is the promised replacement for the deleted `pages`
# argument (docs.extract's own comment: "section=\"page-37\" through es_read
# already expresses that intent") — its predecessor test (page-subset
# extraction) was deleted along with `pages`, and nothing else in the suite
# exercises reaching a PAGE well past the old front-of-document budgets.

def test_es_read_section_reaches_a_deep_pdf_page(wired_cache):
    """A page whose own content starts well past 40,000 characters into
    doc.md must still be reachable by `section="page-N"` — es_read pages the
    FULL cached document, not just some early prefix of it."""
    parts = []
    for i in range(1, 51):
        parts.append(f"## Page {i}\n\nFiller text unique to page {i}. " + ("x" * 900))
    md = "\n\n".join(parts)
    deep_start = md.index("## Page 45")
    assert deep_start > 40_000, "fixture must actually exercise a deep offset"
    target = _seed_doc(wired_cache, md, ext=".pdf")

    out = mcp_server.es_read(target, section="page-45")
    assert out["ok"] is True
    data = out["data"]
    assert "Filler text unique to page 45." in data["content"]
    assert "page 44" not in data["content"]
    assert "page 46" not in data["content"]
